#!/usr/bin/env python3
"""
Genera el Informe Técnico - Entrega 2 (contrato DNP-1025-2026) en formato .docx,
conforme a las Pautas para la elaboración, revisión y entrega de documentos
técnicos de la DENDD (versión 1, 2026).

Reglas de formato aplicadas: papel carta con márgenes de 2,54 cm; Arial 11 en el
cuerpo; interlineado 1,5; alineación izquierda sin justificar; sin sangría de
primera línea; espacio posterior de 6 puntos; numeración decimal hasta el tercer
nivel; paginado "Página X de Y" en el extremo superior derecho; tablas en Arial 9
o 10 con título arriba y fuente abajo; referencias en APA con sangría francesa.

Uso:
    python3 build_entrega2.py [--salida RUTA.docx]
"""

import argparse
import os
import re
import shutil
import subprocess
import tempfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FUENTE = "Arial"
NEGRO = RGBColor(0, 0, 0)
LOGO = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo_dnp.png")
SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

# Registro de entradas para la tabla de contenido y las listas estáticas.
# ENTRADAS_ACTUAL se llena durante la construcción; ENTRADAS_PREV y PAGINAS
# provienen de la pasada anterior del compilador de dos pasadas.
ENTRADAS_ACTUAL = []
ENTRADAS_PREV = []
PAGINAS = {}
TOC_EXCLUIR = {"Página legal y control documental", "Tabla de contenido",
               "Listas de tablas y figuras"}

SALIDA_DEFECTO = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "Bases de datos", "DNP", "productos",
    "Entrega2_InformeTecnico_RuizCastro_v2.docx",
)


# --------------------------------------------------------------------------
# Infraestructura de formato
# --------------------------------------------------------------------------

def configurar_estilos(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FUENTE
    normal.font.size = Pt(11)
    normal.font.color.rgb = NEGRO
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FUENTE)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)

    for nivel, tam in ((1, 14), (2, 12), (3, 11)):
        st = doc.styles[f"Heading {nivel}"]
        st.font.name = FUENTE
        st.font.size = Pt(tam)
        st.font.bold = True
        st.font.color.rgb = NEGRO
        st.font.italic = False
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FUENTE)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def margenes(seccion):
    for lado in ("top", "bottom", "left", "right"):
        setattr(seccion, f"{lado}_margin", Cm(2.54))
    seccion.header_distance = Cm(1.25)
    seccion.footer_distance = Cm(1.25)


def campo(parrafo, instruccion):
    """Inserta un campo de Word (PAGE, NUMPAGES, TOC)."""
    run = parrafo.add_run()
    ini = OxmlElement("w:fldChar")
    ini.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruccion
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    for el in (ini, instr, sep, fin):
        run._r.append(el)
    return run


def campo_abrir(doc, instruccion):
    """Abre un campo cuyo resultado son los párrafos siguientes.

    Permite entregar la tabla de contenido ya visible (el resultado que
    calcula el compilador de dos pasadas) y a la vez actualizable con F9 en
    Word, que es como la trae el resto de anexos de la Dirección."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run()
    ini = OxmlElement("w:fldChar")
    ini.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruccion
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    for el in (ini, instr, sep):
        r._r.append(el)
    return p


def campo_cerrar(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    p.add_run()._r.append(fin)
    return p


def encabezado_paginado(seccion):
    """Encabezado según el modelo de las pautas: logo institucional centrado
    y paginado 'Página X de Y' a la derecha, Arial 8."""
    seccion.header.is_linked_to_previous = False
    p0 = seccion.header.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO):
        p0.add_run().add_picture(LOGO, height=Cm(1.1))
    p = seccion.header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.add_run("Página ")
    campo(p, "PAGE")
    p.add_run(" de ")
    campo(p, "NUMPAGES")
    for parr in (p0, p):
        for run in parr.runs:
            run.font.name = FUENTE
            run.font.size = Pt(8)
            run.font.color.rgb = NEGRO


def parrafo(doc, texto, tam=11, negrita=False, cursiva=False, align=None,
            space_after=6, sangria_francesa=False, izquierda=None):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(tam)
    run.font.bold = negrita
    run.font.italic = cursiva
    run.font.color.rgb = NEGRO
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    if sangria_francesa:
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-1.27)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if izquierda is not None:
        pf.left_indent = Cm(izquierda)
    return p


def vineta(doc, texto):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(11)
    run.font.color.rgb = NEGRO
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def titulo(doc, texto, nivel=1):
    if texto not in TOC_EXCLUIR:
        ENTRADAS_ACTUAL.append(("h", nivel, texto))
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.name = FUENTE
        run.font.color.rgb = NEGRO
    return h


_CONTADOR_TABLAS = [0]


def tabla(doc, titulo_tabla, encabezados, filas, fuente_nota,
          anchos=None, tam=9):
    """Inserta una tabla numerada de forma consecutiva.

    El título usa el estilo Caption para que Word pueda generar la lista de
    tablas automáticamente con el campo correspondiente.
    """
    _CONTADOR_TABLAS[0] += 1
    ENTRADAS_ACTUAL.append(("t", 0, f"Tabla {_CONTADOR_TABLAS[0]}: {titulo_tabla}"))
    p = doc.add_paragraph(style="Caption")
    run = p.add_run(f"Tabla {_CONTADOR_TABLAS[0]}: {titulo_tabla}")
    run.font.name = FUENTE
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = NEGRO
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    # Que la fila de encabezado se repita cuando la tabla se parte entre páginas.
    tr = t.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)
    hdr = t.rows[0].cells
    for i, texto in enumerate(encabezados):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(texto)
        r.font.bold = True
        r.font.name = FUENTE
        r.font.size = Pt(tam)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for fila in filas:
        celdas = t.add_row().cells
        for i, texto in enumerate(fila):
            celdas[i].text = ""
            p = celdas[i].paragraphs[0]
            r = p.add_run(str(texto))
            r.font.name = FUENTE
            r.font.size = Pt(tam)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if anchos:
        # El ancho debe fijarse celda por celda y desactivando el ajuste
        # automático; de lo contrario Word y LibreOffice lo recalculan.
        tblPr = t._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        for fila in t.rows:
            for i, ancho in enumerate(anchos):
                celda = fila.cells[i]
                celda.width = Cm(ancho)
                tcW = celda._tc.get_or_add_tcPr().get_or_add_tcW()
                tcW.set(qn("w:w"), str(int(Cm(ancho).twips)))
                tcW.set(qn("w:type"), "dxa")
    parrafo(doc, f"Fuente: {fuente_nota}", tam=9, space_after=12)
    return t


def salto(doc):
    doc.add_page_break()


def _clave(texto):
    return re.sub(r"\s+", " ", texto).strip()[:30]


def _linea_toc(doc, texto, pagina, nivel=1, negrita=False):
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Cm(0.5 * (nivel - 1))
    pf.tab_stops.add_tab_stop(Cm(16.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(texto)
    r.font.name = FUENTE
    r.font.size = Pt(11 if nivel == 1 else 10)
    r.font.bold = negrita and nivel == 1
    r.font.color.rgb = NEGRO
    r2 = p.add_run("\t" + (str(pagina) if pagina else ""))
    r2.font.name = FUENTE
    r2.font.size = Pt(10)
    r2.font.color.rgb = NEGRO
    return p


def toc_estatica(doc):
    """Tabla de contenido y listas de tablas y figuras.

    Cada lista se entrega dentro de un campo de Word con su resultado ya
    calculado por el compilador de dos pasadas: se ve al abrir el archivo y
    se actualiza con F9 si el usuario reordena el documento."""
    titulo(doc, "Tabla de contenido", 1)
    campo_abrir(doc, r'TOC \o "1-3" \h \z \u')
    for tipo, nivel, texto in ENTRADAS_PREV:
        if tipo != "h":
            continue
        _linea_toc(doc, texto, PAGINAS.get(_clave(texto), ""), nivel,
                   negrita=(nivel == 1))
    campo_cerrar(doc)
    salto(doc)
    titulo(doc, "Listas de tablas y figuras", 1)
    parrafo(doc, "Lista de tablas", tam=11, negrita=True, space_after=3)
    campo_abrir(doc, r'TOC \h \z \c "Tabla"')
    for tipo, nivel, texto in ENTRADAS_PREV:
        if tipo == "t":
            _linea_toc(doc, texto, PAGINAS.get(_clave(texto), ""), 1)
    campo_cerrar(doc)
    if any(t == "f" for t, _, _ in ENTRADAS_PREV):
        parrafo(doc, "", space_after=3)
        parrafo(doc, "Lista de figuras", tam=11, negrita=True, space_after=3)
        campo_abrir(doc, r'TOC \h \z \c "Figura"')
        for tipo, nivel, texto in ENTRADAS_PREV:
            if tipo == "f":
                _linea_toc(doc, texto, PAGINAS.get(_clave(texto), ""), 1)
        campo_cerrar(doc)


def _paginas_del_pdf(pdf):
    """Mapea cada entrada registrada a la página del cuerpo donde aparece.

    Una entrada figura dos veces en el PDF: en la tabla de contenido y en el
    cuerpo. Las páginas de índice se detectan porque concentran muchas
    entradas (cuatro o más) y se excluyen; de lo que queda se toma la primera
    aparición al inicio de línea.
    """
    out = subprocess.run(["pdftotext", "-layout", pdf, "-"],
                         capture_output=True, check=True)
    paginas = out.stdout.decode("utf-8", "replace").split("\f")
    lineas = [[re.sub(r"\s+", " ", l).strip() for l in pg.split("\n")]
              for pg in paginas]
    claves = [_clave(t) for _, _, t in ENTRADAS_PREV]

    def apariciones(clave):
        return [i for i, ls in enumerate(lineas)
                if any(l.startswith(clave) for l in ls)]

    conteo = [sum(1 for c in claves if any(l.startswith(c) for l in ls))
              for ls in lineas]
    indices = {i for i, n in enumerate(conteo) if n >= 4}
    mapa = {}
    for clave in claves:
        cands = apariciones(clave)
        cuerpo = [p for p in cands if p not in indices]
        if cuerpo:
            mapa[clave] = cuerpo[0] + 1
        elif cands:
            mapa[clave] = cands[0] + 1
    return mapa


def compilar(construir_fn, salida):
    """Compilador de dos pasadas: construye, convierte a PDF para conocer la
    paginación real y reconstruye con los números de página definitivos."""
    global ENTRADAS_ACTUAL, ENTRADAS_PREV, PAGINAS
    ENTRADAS_PREV, PAGINAS = [], {}
    with tempfile.TemporaryDirectory() as tmp:
        borrador = os.path.join(tmp, "borrador.docx")
        # Pasada 0: registra las entradas (la TOC sale vacía).
        ENTRADAS_ACTUAL = []
        construir_fn(borrador)
        ENTRADAS_PREV = list(ENTRADAS_ACTUAL)
        # Pasada 1: TOC con el número correcto de líneas, páginas en blanco.
        ENTRADAS_ACTUAL = []
        construir_fn(borrador)
        subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf",
                        "--outdir", tmp, borrador],
                       capture_output=True, check=True)
        PAGINAS = _paginas_del_pdf(os.path.join(tmp, "borrador.pdf"))
        # Pasada 2: números de página definitivos.
        ENTRADAS_ACTUAL = []
        construir_fn(salida)
    return salida


# --------------------------------------------------------------------------
# Contenido
# --------------------------------------------------------------------------

def portada(doc, titulo_doc, subtitulo, version):
    """Portada según el modelo de la Ilustración 1 de las pautas.

    Los espaciados están calibrados para que todo el bloque quepa en una sola
    página carta con márgenes de 2,54 cm."""
    if os.path.exists(LOGO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(16)
        p.add_run().add_picture(LOGO, width=Cm(4.6))
    parrafo(doc, "Dirección de Economía Naranja y Desarrollo Digital", tam=12,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)
    parrafo(doc, titulo_doc, tam=18, negrita=True,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    parrafo(doc, subtitulo, tam=12.5,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)
    parrafo(doc, version, tam=11, negrita=True,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)
    parrafo(doc, "Elaborado por:", tam=11, negrita=True,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    parrafo(doc, "Ricardo Esteban Ruiz Castro", tam=11,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    parrafo(doc, "Contrato de prestación de servicios profesionales DNP-1025-2026",
            tam=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    parrafo(doc, "Centro de Pensamiento de Desarrollo Digital", tam=11,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    parrafo(doc, "Revisado y aprobado por:", tam=11, negrita=True,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    parrafo(doc, "Edwin Alejandro Buenhombre Moreno", tam=11,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    parrafo(doc, "Director Técnico y supervisor del contrato", tam=11,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)
    parrafo(doc, "Departamento Nacional de Planeación", tam=11, negrita=True,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    parrafo(doc, "Bogotá, D. C.", tam=11, align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=2)
    parrafo(doc, "2026", tam=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


_CONTADOR_FIGURAS = [0]
DIR_FIGURAS = ["."]


def figura(doc, nombre, titulo_fig, fuente_nota, ancho=15.5):
    """Inserta una figura numerada con título arriba y fuente abajo."""
    _CONTADOR_FIGURAS[0] += 1
    etiqueta = f"Figura {_CONTADOR_FIGURAS[0]}: {titulo_fig}"
    ENTRADAS_ACTUAL.append(("f", 0, etiqueta))
    p = doc.add_paragraph(style="Caption")
    run = p.add_run(etiqueta)
    run.font.name = FUENTE
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = NEGRO
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc.add_picture(os.path.join(DIR_FIGURAS[0], nombre), width=Cm(ancho))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo(doc, f"Fuente: {fuente_nota}", tam=9, space_after=12)


def pagina_legal(doc):
    titulo(doc, "Página legal y control documental", 1)
    parrafo(doc,
            "Este documento corresponde a la segunda entrega del contrato de "
            "prestación de servicios profesionales DNP-1025-2026 y se elabora "
            "para uso de la Dirección de Economía Naranja y Desarrollo Digital "
            "del Departamento Nacional de Planeación. Su contenido puede "
            "reproducirse total o parcialmente citando la fuente.")

    tabla(doc, "Control documental",
          ["Campo", "Contenido"],
          [["Título", "Marco conceptual, normativo y de fuentes para el análisis de "
                      "datos y la evaluación de impacto de las políticas públicas "
                      "del sector TIC"],
           ["Tipo documental", "Informe técnico"],
           ["Entrega contractual", "Entrega 2 de 4 – contrato DNP-1025-2026"],
           ["Versión", "2"],
           ["Dependencia responsable", "Dirección de Economía Naranja y Desarrollo Digital"],
           ["Elabora", "Ricardo Esteban Ruiz Castro, contratista"],
           ["Revisa y aprueba", "Edwin Alejandro Buenhombre Moreno, Director Técnico"],
           ["Fecha de la versión", "Agosto de 2026"],
           ["Fecha de corte de la información", "3 de agosto de 2026"]],
          "elaboración propia.", anchos=[4.5, 11.5])

    tabla(doc, "Historial de cambios",
          ["Versión", "Fecha", "Descripción del cambio"],
          [["1", "Junio de 2026", "Versión inicial radicada ante la supervisión del contrato."],
           ["2", "Agosto de 2026",
            "Ajuste del documento a las pautas para la elaboración, revisión y entrega "
            "de documentos técnicos de la DENDD. Incorporación de los repositorios de "
            "información sugeridos por la supervisión. Reproceso del ejercicio de "
            "analítica de texto sobre el corpus completo de los documentos CONPES 3975 "
            "y 4144, con actualización de los resultados reportados. Incorporación del "
            "análisis del Plan de Acción y Seguimiento y del reporte de avance de "
            "SisCONPES con corte al 21 de julio de 2026. Verificación de referencias y "
            "cifras contra fuente primaria. Incorporación del capítulo sobre situación "
            "deseada y punto de partida, y del anexo de declaración de uso de "
            "herramientas de inteligencia artificial."]],
          "elaboración propia.", anchos=[2.0, 3.0, 11.0])


def tabla_contenido(doc):
    toc_estatica(doc)


def siglas(doc):
    titulo(doc, "Siglas y abreviaturas", 1)
    tabla(doc, "Siglas y abreviaturas empleadas en el documento",
          ["Sigla", "Significado"],
          [["CONPES", "Consejo Nacional de Política Económica y Social"],
           ["DANE", "Departamento Administrativo Nacional de Estadística"],
           ["DENDD", "Dirección de Economía Naranja y Desarrollo Digital"],
           ["DNP", "Departamento Nacional de Planeación"],
           ["ECV", "Encuesta Nacional de Calidad de Vida"],
           ["END", "Estrategia Nacional Digital"],
           ["Entic", "Encuesta Nacional de Tecnologías de la Información y las Comunicaciones"],
           ["Furag", "Formulario Único de Reporte de Avance a la Gestión"],
           ["IA", "Inteligencia artificial"],
           ["IPD", "Índice de Pobreza Digital"],
           ["MinTIC", "Ministerio de Tecnologías de la Información y las Comunicaciones"],
           ["ONIA", "Observatorio Nacional de Inteligencia Artificial"],
           ["PAS", "Plan de Acción y Seguimiento"],
           ["PND", "Plan Nacional de Desarrollo"],
           ["Sinergia", "Sistema Nacional de Evaluación de Gestión y Resultados"],
           ["SisCONPES", "Sistema de Seguimiento a Documentos CONPES"],
           ["TIC", "Tecnologías de la información y las comunicaciones"],
           ["UIT", "Unión Internacional de Telecomunicaciones"]],
          "elaboración propia.", anchos=[3.0, 13.0])


def presentacion(doc):
    titulo(doc, "Presentación", 1)
    parrafo(doc,
            "La Dirección de Economía Naranja y Desarrollo Digital necesita "
            "evidencia propia para orientar sus decisiones de política. Buena "
            "parte de la información sobre el sector TIC existe, pero está "
            "dispersa entre sistemas de seguimiento, encuestas del sistema "
            "estadístico nacional, registros administrativos del sector y "
            "portales de datos abiertos, cada uno con su periodicidad, su nivel "
            "de desagregación y sus reglas de acceso. Reunirla y ordenarla es la "
            "condición previa para cualquier medición seria de resultados.")
    parrafo(doc,
            "El presente documento responde a esa necesidad, en la medida en que establece el marco "
            "conceptual y normativo con el que se analizarán las políticas del "
            "sector durante la vigencia 2026, inventaría las fuentes de "
            "información disponibles y su utilidad concreta, y propone el "
            "conjunto inicial de indicadores con los que se medirán resultados e "
            "impactos. También deja probado, sobre datos reales, que el enfoque "
            "es viable con la información pública hoy accesible.")
    parrafo(doc,
            "Su destinatario inmediato es la supervisión del contrato y el equipo "
            "del Centro de Pensamiento de Desarrollo Digital. El resultado "
            "esperado es doble: cerrar la discusión sobre qué se va a medir y con "
            "qué datos, de modo que las dos entregas restantes se concentren en "
            "producir hallazgos; y dejar un inventario de fuentes reutilizable por "
            "la Dirección más allá de este contrato, en particular por el "
            "Observatorio Nacional de Inteligencia Artificial.")


def introduccion(doc):
    titulo(doc, "1. Introducción", 1)
    parrafo(doc,
            "Colombia aprobó en febrero de 2025 su Política Nacional de "
            "Inteligencia Artificial mediante el documento CONPES 4144, que "
            "compromete 106 acciones y una inversión aproximada de 479.273 "
            "millones de pesos hasta 2030 (DNP, 2025). Es la política digital más "
            "ambiciosa que ha formulado el país y llega, además, con un sistema de "
            "seguimiento semestral. La pregunta que abre este trabajo es si esa "
            "arquitectura de seguimiento permite decir algo sobre los resultados "
            "de la política, y no solamente sobre el cumplimiento de sus tareas.")
    parrafo(doc,
            "Conviene precisar la diferencia, porque de ella depende todo lo que "
            "sigue: mientras el seguimiento registra si una acción se ejecutó, la "
            "evaluación establece qué cambió por haberla ejecutado. Un sistema "
            "puede, en consecuencia, reportar un avance del cien por ciento en "
            "sus acciones y ser al mismo tiempo incapaz de mostrar un efecto "
            "sobre la conectividad, las habilidades digitales o la adopción de "
            "tecnología en las empresas. Cerrar esa distancia supone definir de "
            "antemano qué se mide, con qué dato y contra qué referencia, y es "
            "justamente esa base la que el presente informe se propone construir.")
    parrafo(doc,
            "Con ese fin, la entrega persigue cuatro objetivos encadenados: fijar "
            "el marco conceptual y normativo del análisis; identificar y "
            "caracterizar las fuentes de información disponibles; proponer un "
            "conjunto inicial de indicadores de resultado e impacto alineados con "
            "los ejes del CONPES 4144 y con los sistemas de seguimiento del "
            "Departamento; y verificar, mediante un ejercicio acotado sobre datos "
            "reales, que el análisis previsto para las entregas siguientes es "
            "factible con la información hoy accesible.")
    parrafo(doc,
            "Estos objetivos se enmarcan dentro de los límites que fija el objeto "
            "contractual, de modo que el documento cubre las políticas del sector "
            "TIC vigentes en 2026 tomando el CONPES 4144 como instrumento central "
            "y el CONPES 3975 de 2019 como referencia de trayectoria. Quedan por "
            "fuera, en cambio, la valoración del desempeño de entidades "
            "particulares, la formulación de política nueva y la sustitución de "
            "los informes de seguimiento que producen Sinergia y SisCONPES. "
            "Tampoco se presentan todavía resultados de evaluación, pues estos "
            "corresponden a la Entrega 3 y a la entrega final.")
    parrafo(doc,
            "Metodológicamente, el trabajo combina revisión documental de "
            "instrumentos de política y literatura especializada, inventario "
            "sistemático de fuentes con verificación de acceso efectivo, y una "
            "aplicación exploratoria de analítica de texto sobre el corpus de los "
            "dos documentos CONPES. Todas las cifras citadas se contrastaron "
            "contra la fuente primaria; cuando un dato no pudo verificarse, se "
            "excluyó del documento y así se advierte.")
    parrafo(doc,
            "El informe se organiza de la siguiente manera. El capítulo 2 describe "
            "los antecedentes y el punto de partida medido del sector. El capítulo "
            "3 presenta el marco conceptual de la evaluación. El capítulo 4 "
            "expone el marco normativo e institucional y la articulación de este "
            "contrato con los compromisos de la Dirección. El capítulo 5 recoge la "
            "revisión bibliográfica. El capítulo 6 explica la metodología empleada "
            "en esta entrega. El capítulo 7 inventaría las fuentes de información. "
            "El capítulo 8 propone los indicadores. El capítulo 9 reporta la "
            "aplicación exploratoria de analítica de texto y la prueba de "
            "viabilidad. El capítulo 10 analiza la estructura del plan de acción "
            "de la política y su asignación de recursos. El capítulo 11 contrasta "
            "la situación deseada con el punto de partida. Los capítulos 12 y 13 "
            "presentan conclusiones y recomendaciones.")


def antecedentes(doc):
    titulo(doc, "2. Antecedentes y contexto", 1)

    titulo(doc, "2.1. Veinticinco años de política digital", 2)
    parrafo(doc,
            "Conviene empezar por deshacer una impresión frecuente, y es la de "
            "que la política digital colombiana nació con la ola reciente de "
            "instrumentos sobre datos e inteligencia artificial. Su punto de "
            "partida formal está bastante más atrás: en el documento CONPES 3072 "
            "de 2000, que adoptó la Agenda de Conectividad con el propósito de "
            "masificar el uso de las tecnologías de la información para elevar la "
            "competitividad del aparato productivo, modernizar las instituciones "
            "públicas y socializar el acceso a la información (DNP, 2000). Ese "
            "documento, a su vez, recogía un trabajo previo que venía desde la "
            "creación del Consejo Nacional de Informática en 1997, de modo que el "
            "país lleva ya cerca de tres décadas construyendo política pública "
            "sobre esta materia.")
    parrafo(doc,
            "Reconocer esa profundidad histórica no es un gesto de erudición: "
            "cambia la manera de leer los resultados. Si el CONPES 4144 fuera el "
            "primer instrumento del sector, sus indicadores partirían de cero y "
            "cualquier avance sería mérito suyo. Como no lo es, la evaluación "
            "tiene que distinguir entre lo que la política nueva produce y lo que "
            "hereda de una trayectoria de veinticinco años, tarea que exige "
            "conocer los eslabones de esa cadena.")
    tabla(doc, "Hitos de la política digital colombiana, 2000-2025",
          ["Año", "Instrumento", "Aporte principal"],
          [["2000", "CONPES 3072 – Agenda de Conectividad",
            "Primer marco de política para masificar el uso de las tecnologías de la información."],
           ["2009", "Ley 1341 – Ley TIC",
            "Define los principios de la sociedad de la información, ordena institucionalmente el sector y crea la Agencia Nacional del Espectro."],
           ["2010-2018", "Planes Vive Digital",
            "Dos ciclos de política de masificación de infraestructura y acceso, con foco en despliegue de red y puntos de acceso comunitario."],
           ["2016", "CONPES 3854 – Seguridad digital",
            "Traslada el enfoque de la seguridad informática a la gestión de riesgos digitales en la actividad económica y social."],
           ["2018", "CONPES 3920 – Explotación de datos",
            "Trata los datos como activo estratégico y abre el frente de gobernanza de la información pública."],
           ["2019", "CONPES 3975 – Transformación digital e IA",
            "Primer marco de política sobre inteligencia artificial, subordinado todavía al concepto articulador de transformación digital."],
           ["2019", "Ley 1978 – Modernización del sector TIC",
            "Reorganiza la institucionalidad, unifica la vigilancia y redefine el esquema de financiación del acceso universal."],
           ["2023", "Estrategia Nacional Digital 2023-2026",
            "Instrumento de implementación sectorial que articula conectividad, gobierno digital, economía digital y talento."],
           ["2023", "Ley 2294 – PND 2022-2026",
            "Incorpora la conectividad significativa y la alfabetización digital con enfoque diferencial entre los mandatos de desarrollo."],
           ["2025", "CONPES 4144 – Política Nacional de IA",
            "Convierte la inteligencia artificial en objeto autónomo de política, con seis ejes propios y 106 acciones a 2030."]],
          "elaboración propia con base en los instrumentos citados.",
          anchos=[1.6, 5.0, 9.4])
    parrafo(doc,
            "Vista en conjunto, la secuencia describe un desplazamiento en tres "
            "tiempos. Durante la primera década el problema público fue la "
            "conexión, y los instrumentos se ocuparon de desplegar "
            "infraestructura y crear puntos de acceso. A partir de 2016, con la "
            "seguridad digital y la explotación de datos, el foco se corrió hacia "
            "la calidad y el gobierno de la información que circula por esa "
            "infraestructura. Y desde 2019, con más nitidez en 2025, el objeto de "
            "la política dejó de ser la red para ser lo que se hace con ella: "
            "capacidades, adopción, riesgos y reglas de uso de una tecnología "
            "específica.")
    parrafo(doc,
            "El CONPES 4144 marca el punto de llegada de ese recorrido, porque a "
            "diferencia de su antecesor de 2019 ya no formula una política de "
            "digitalización con un capítulo dedicado a la inteligencia "
            "artificial, sino una política de inteligencia artificial con seis "
            "ejes propios, entre ellos uno consagrado a la ética y la gobernanza "
            "y otro a la mitigación de riesgos. Ese desplazamiento, que hasta "
            "aquí es una lectura cualitativa, resulta además medible sobre el "
            "texto de ambos documentos, y así se documenta en el capítulo 9.")

    titulo(doc, "2.2. El punto de partida medido", 2)
    parrafo(doc,
            "Puesto que toda evaluación requiere una fotografía inicial contra la "
            "cual contrastar los cambios posteriores, vale la pena detenerse en lo "
            "que dicen las mediciones oficiales más recientes, cuyo balance "
            "resulta a la vez alentador y exigente. El acceso a internet fijo o "
            "móvil alcanzó el 65,6% de los "
            "hogares colombianos en 2024, con un aumento de 1,7 puntos "
            "porcentuales frente a 2023 (DANE, 2025). El promedio, sin embargo, "
            "esconde el problema: en las cabeceras municipales la conectividad "
            "llega al 72,5% y en el área rural dispersa al 41,9%, una brecha de "
            "más de treinta puntos que se repite entre departamentos, con Bogotá "
            "en 82,7% y Vichada, Chocó y Vaupés por debajo del 30%.")
    parrafo(doc,
            "La medición más completa del rezago la produjo la propia Dirección. "
            "El Índice de Pobreza Digital, publicado en diciembre de 2025, "
            "establece que el 37,9% de la población de once años o más —cerca de "
            "16,79 millones de personas— vive en pobreza digital, entendida como "
            "la privación simultánea o parcial en conectividad, dispositivos y "
            "habilidades (DNP, 2025b). Al desagregar por dimensión, la más crítica "
            "no es la infraestructura sino la capacidad de uso: el 60% de la "
            "población presenta brechas en habilidades digitales básicas, frente "
            "al 39% con deficiencias de calidad o frecuencia de conexión y al "
            "33% con insuficiencia de dispositivos.")
    parrafo(doc,
            "A las cifras de acceso y de pobreza digital conviene sumar una "
            "tercera coordenada, que es la del punto de partida institucional. "
            "Colombia llega a esta política con un aparato de medición sectorial "
            "poco frecuente en la región: una encuesta de hogares con módulo de "
            "tecnología de periodicidad anual, un observatorio sectorial con "
            "series trimestrales, un portal de datos abiertos con registros "
            "administrativos consultables por interfaz de programación y, desde "
            "2025, una estadística multidimensional de pobreza digital. Ese "
            "acervo, construido a lo largo de la trayectoria descrita en el "
            "numeral anterior, es el que hace viable un contrato como este sin "
            "levantamientos de información propios, y constituye en sí mismo un "
            "activo de la política digital que rara vez se contabiliza entre sus "
            "resultados.")
    parrafo(doc,
            "Ese orden de magnitud reorienta la discusión. Si la privación "
            "dominante está en las habilidades y no en la red, el indicador de "
            "cobertura deja de ser suficiente como medida de resultado de la "
            "política digital, y la evaluación debe incorporar desde el diseño la "
            "dimensión de capacidades y el enfoque diferencial. Esta constatación "
            "orienta la propuesta de indicadores del capítulo 8.")


def marco_conceptual(doc):
    titulo(doc, "3. Marco conceptual", 1)
    parrafo(doc,
            "Evaluar una política de inteligencia artificial obliga a trabajar en "
            "la intersección de tres cuerpos teóricos que rara vez se citan "
            "juntos. El primero es el de la evaluación de políticas públicas, que "
            "aporta la gramática de la medición: qué se mide en cada eslabón de "
            "una intervención y bajo qué condiciones puede atribuirse un cambio a "
            "una causa. El segundo es el de la modernización del Estado, que "
            "explica por qué las administraciones contemporáneas se organizan como "
            "lo hacen y qué implica gestionar tecnología desde adentro. Y el "
            "tercero es el de la gobernanza anticipatoria y la gobernanza de la "
            "inteligencia artificial, que se ocupa del problema específico de "
            "regular una tecnología que cambia más rápido de lo que el ciclo "
            "normativo alcanza a procesar. Los tres se desarrollan a "
            "continuación, en ese orden, y confluyen en el numeral 3.5.")

    titulo(doc, "3.1. La cadena de valor de las intervenciones públicas", 2)
    parrafo(doc,
            "El primer anclaje del análisis es la cadena de valor de las "
            "intervenciones públicas, formulada en el ámbito latinoamericano por "
            "el Instituto Latinoamericano y del Caribe de Planificación Económica "
            "y Social, que distingue cinco eslabones sucesivos: insumos, "
            "actividades, productos, resultados e impactos (Bonnefoy y Armijo, "
            "2005). La secuencia parece obvia enunciada así, pero su utilidad "
            "práctica es considerable, porque ordena los indicadores según lo que "
            "efectivamente miden y previene el error más frecuente en los "
            "sistemas de seguimiento, que consiste en presentar un producto "
            "entregado como si fuera un resultado alcanzado.")
    parrafo(doc,
            "Para hacer tangible la diferencia conviene trasladarla al objeto de "
            "este contrato. Los recursos y el personal asignados a un programa de "
            "formación en inteligencia artificial son insumos; los cursos "
            "dictados son actividades; el número de personas formadas es un "
            "producto; la proporción de la población que efectivamente adquiere y "
            "conserva habilidades digitales es un resultado; y la reducción de la "
            "pobreza digital medida sobre los hogares es un impacto. Los cinco "
            "registros son legítimos y necesarios, pero responden preguntas "
            "distintas y se construyen con datos distintos, de manera que "
            "confundirlos produce evaluaciones que celebran actividad sin poder "
            "demostrar efecto.")
    parrafo(doc,
            "De esa distinción se desprende, además, la clasificación de los "
            "tipos de evaluación aplicables según el momento del ciclo de "
            "política, que es la que ordena la secuencia de las cuatro entregas "
            "de este contrato.")
    tabla(doc, "Tipos de evaluación según el momento del ciclo de política",
          ["Tipo de evaluación", "Momento", "Pregunta que responde", "Requisito de información"],
          [["De diseño", "Ex ante",
            "¿La intervención responde al problema identificado y su lógica causal es plausible?",
            "Documentos de política, diagnóstico y plan de acción"],
           ["De proceso", "Durante la implementación",
            "¿La ejecución avanza según lo planificado y dónde se producen los cuellos de botella?",
            "Reportes de avance y ejecución presupuestal"],
           ["De resultados", "Intermedia o ex post",
            "¿Se alcanzaron los productos y resultados previstos?",
            "Indicadores con línea base y meta"],
           ["De impacto", "Ex post, con datos suficientes",
            "¿Qué parte del cambio observado es atribuible a la intervención?",
            "Datos de tratamiento y comparación, o variación exógena"]],
          "elaboración propia con base en Bonnefoy y Armijo (2005) y Gertler et al. (2017).",
          anchos=[2.6, 2.8, 6.2, 4.4])
    parrafo(doc,
            "Aplicada al cronograma contractual, esta clasificación explica por "
            "qué el presente documento constituye una evaluación de diseño, por "
            "qué la Entrega 3 aborda procesos y resultados intermedios, y por qué "
            "la aproximación de impacto queda reservada para la entrega final, "
            "cuando exista un cierre de vigencia sobre el cual medir.")

    titulo(doc, "3.2. El problema del contrafactual y sus límites en el sector TIC", 2)
    parrafo(doc,
            "El eslabón más exigente de la cadena es el último. La evaluación de "
            "impacto se define, en la literatura de referencia, por la "
            "identificación del efecto causal de una intervención sobre un "
            "resultado de interés, lo que supone construir un contrafactual: una "
            "estimación de qué habría ocurrido en ausencia de la política "
            "(Gertler et al., 2017). Sin ese contrafactual, cualquier cambio "
            "observado después de una intervención puede deberse a la "
            "intervención, a la tendencia previa, a un factor externo simultáneo "
            "o a la composición de la población observada.")
    parrafo(doc,
            "Ahora bien, aunque el problema del contrafactual es general, en el "
            "sector de las tecnologías de la información adquiere una dificultad "
            "particular por tres razones que conviene enunciar con precisión, "
            "porque delimitan lo que este contrato puede y no puede afirmar. La "
            "primera es la superposición de intervenciones: un mismo municipio "
            "recibe simultáneamente despliegue de infraestructura, programas de "
            "formación, iniciativas de adopción empresarial y proyectos de "
            "gobierno digital, de modo que aislar el efecto de uno solo exige una "
            "variación que en la práctica no existe. La segunda es el cambio "
            "tecnológico autónomo, pues buena parte de la mejora en conectividad "
            "y en capacidad de cómputo ocurre por dinámica de mercado y "
            "abaratamiento de equipos, con independencia de la política, tal como "
            "se documenta en el capítulo 9 para el caso de la migración a fibra "
            "óptica. La tercera es la heterogeneidad territorial, ya que una "
            "misma intervención rinde de forma distinta en una cabecera urbana "
            "con mercado consolidado que en un municipio de la Amazonia, lo que "
            "convierte el promedio nacional en un estadístico poco informativo.")
    parrafo(doc,
            "Reconocer estas limitaciones no invalida el ejercicio ni justifica "
            "renunciar a la evaluación; lo que exige es disciplina en el lenguaje. "
            "En consecuencia, cuando los datos permitan estimar un efecto causal "
            "el documento lo afirmará como tal, y cuando solo permitan observar "
            "asociaciones o trayectorias las nombrará por su nombre, sin "
            "deslizarse hacia afirmaciones causales que la evidencia no sostenga. "
            "Esa disciplina, más que un escrúpulo metodológico, protege la "
            "credibilidad de las recomendaciones que el producto final entregue a "
            "la Dirección.")

    titulo(doc, "3.3. Del gobierno electrónico a la gobernanza de la era digital", 2)
    parrafo(doc,
            "El segundo cuerpo teórico que informa este trabajo explica el "
            "contexto administrativo en el que la política se ejecuta. Durante las "
            "décadas de 1980 y 1990, la Nueva Gestión Pública propuso reorganizar "
            "el Estado alrededor de la eficiencia, la desagregación de funciones y "
            "la introducción de mecanismos de mercado en la provisión de servicios "
            "(Barzelay, 2001). Ese paradigma dejó una huella duradera en la "
            "arquitectura institucional colombiana, visible en la orientación a "
            "resultados de los sistemas de seguimiento y en la lógica de "
            "indicadores de gestión que el propio Plan de Acción y Seguimiento del "
            "CONPES 4144 reproduce.")
    parrafo(doc,
            "Sin embargo, hacia mediados de los años dos mil comenzó a "
            "documentarse el agotamiento de ese modelo. En un trabajo que se "
            "volvió referencia obligada, Dunleavy y sus coautores sostuvieron que "
            "la Nueva Gestión Pública había llegado a su fin y que el paradigma "
            "emergente era el de la gobernanza de la era digital, caracterizado "
            "por la reintegración de funciones antes fragmentadas, la "
            "organización de los servicios en torno a las necesidades del "
            "ciudadano y la digitalización de los procesos administrativos como "
            "eje y no como accesorio (Dunleavy et al., 2006). La distinción es "
            "pertinente para este análisis porque la Política Nacional de "
            "Inteligencia Artificial pertenece de lleno al segundo paradigma —su "
            "eje de uso y adopción supone precisamente reintegrar capacidades "
            "tecnológicas dentro del Estado—, mientras que su instrumento de "
            "medición sigue construido con la gramática del primero.")
    parrafo(doc,
            "A esa discusión se suma la lectura contemporánea de la cuarta "
            "revolución industrial, que subraya la velocidad y el carácter "
            "sistémico de la transformación tecnológica actual y su exigencia "
            "sobre las capacidades estatales (Schwab, 2016). En el ámbito "
            "iberoamericano, esa exigencia se ha traducido en modelos operativos "
            "para transformar instituciones jerárquicas en organizaciones en red, "
            "entre los cuales el Hexágono de Innovación Pública ocupa un lugar "
            "destacado por su adopción institucional en la región.")

    titulo(doc, "3.4. Innovación pública y gestión del conocimiento", 2)
    parrafo(doc,
            "El modelo del Hexágono de Innovación Pública fue formulado por Raúl "
            "Oliván a partir del análisis de ciento cinco metodologías de "
            "innovación aplicadas en agencias públicas, y adoptado por la "
            "Secretaría General Iberoamericana como marco de referencia regional "
            "(Oliván, 2020). Su propuesta consiste en seis vectores que operan "
            "simultáneamente —apertura, transdisciplinariedad, agilidad, "
            "prototipado, colaboración y tecnología— y cuyo propósito conjunto es "
            "acelerar la transición desde instituciones cerradas y "
            "compartimentadas hacia organizaciones abiertas y en red.")
    parrafo(doc,
            "La pertinencia del modelo para este contrato no es meramente "
            "conceptual, dado que es el marco con el que la Dirección estructuró "
            "la propuesta de creación del Observatorio Nacional de Inteligencia "
            "Artificial. En consecuencia, el producto contractual se articula con "
            "esos mismos vectores, y esa correspondencia se detalla en el numeral "
            "4.4. Conviene, con todo, señalar una precaución: el Hexágono es un "
            "modelo de diseño organizacional, no un marco de evaluación, de "
            "manera que orienta cómo se organiza el trabajo del observatorio pero "
            "no sustituye los criterios de medición desarrollados en los "
            "numerales anteriores.")
    parrafo(doc,
            "Un segundo aporte de esta familia teórica proviene de la gestión del "
            "conocimiento organizacional. El ciclo formulado por Nonaka y "
            "Takeuchi describe cómo el conocimiento circula dentro de una "
            "organización mediante cuatro conversiones sucesivas entre lo tácito "
            "y lo explícito: socialización, exteriorización, combinación e "
            "interiorización (Nonaka y Takeuchi, 1995). Ese ciclo, que la "
            "propuesta del observatorio adopta como esquema de tratamiento de sus "
            "contenidos, tiene una implicación directa sobre el diseño de este "
            "producto: un informe técnico que no se convierta en práctica "
            "institucional se queda en el eslabón de la exteriorización y no "
            "completa el ciclo. De ahí que el inventario de fuentes y las rutinas "
            "de procesamiento se entreguen como instrumentos reutilizables y no "
            "solo como resultados descritos en prosa.")
    parrafo(doc,
            "En el terreno específico de la innovación en el sector público, "
            "finalmente, la distinción entre innovación incremental y disruptiva "
            "permite calibrar expectativas sobre lo que una política como esta "
            "puede lograr en un cuatrienio. La innovación incremental crea valor "
            "sobre lo existente mediante mejoras graduales de procesos y "
            "servicios, mientras la disruptiva altera el modelo mismo de "
            "provisión (Bason, 2010). Buena parte de las acciones del CONPES 4144 "
            "pertenece a la primera categoría, lo que es razonable y esperable, "
            "pero también significa que sus efectos serán acumulativos y de "
            "manifestación lenta, condición que la evaluación debe incorporar en "
            "sus expectativas de medición.")

    titulo(doc, "3.5. Gobernanza anticipatoria y gobernanza de la inteligencia artificial", 2)
    parrafo(doc,
            "El tercer cuerpo teórico responde a una dificultad que los dos "
            "anteriores no resuelven, y es la de gobernar una tecnología cuyo "
            "desarrollo va sistemáticamente por delante de la capacidad "
            "regulatoria del Estado. La literatura especializada ha caracterizado "
            "ese desfase como el rasgo definitorio del problema: la gobernanza de "
            "la inteligencia artificial enfrenta simultáneamente incertidumbre "
            "sobre los efectos de la tecnología, asimetría de información entre "
            "reguladores y desarrolladores, y una velocidad de cambio que vuelve "
            "obsoletos los instrumentos antes de que terminen de implementarse "
            "(Taeihagh, 2021).")
    parrafo(doc,
            "Frente a ese desfase, el enfoque de gobernanza anticipatoria propone "
            "un desplazamiento de fondo: en lugar de reaccionar cuando los "
            "efectos se materializan, las instituciones deben desarrollar "
            "capacidades para operar bajo incertidumbre y considerar de manera "
            "sistemática futuros alternativos, usando esa consideración para "
            "orientar decisiones presentes (Miller, 2018). No se trata de "
            "predecir, que sería una pretensión insostenible, sino de ampliar el "
            "repertorio de futuros que la institución es capaz de imaginar y para "
            "los cuales está preparada.")
    parrafo(doc,
            "Este es, precisamente, el enfoque que el CONPES 4144 incorporó al "
            "encargar al Departamento Nacional de Planeación una estrategia "
            "anticipatoria de prospectiva y vigilancia tecnológica, y el que da "
            "fundamento al Observatorio Nacional de Inteligencia Artificial como "
            "instrumento de esa estrategia, según se detalla en el capítulo 4.")
    parrafo(doc,
            "De la convergencia de los tres cuerpos teóricos se sigue una "
            "consecuencia operativa que atraviesa todo el contrato: la "
            "inteligencia artificial es a la vez el objeto de la política que se "
            "evalúa y un instrumento para evaluarla. Las técnicas de "
            "procesamiento de lenguaje natural permiten analizar el corpus "
            "documental de la política con una sistematicidad que la lectura "
            "manual no alcanza, y las herramientas de analítica permiten cruzar "
            "registros administrativos con desagregaciones que los informes "
            "agregados ocultan. El capítulo 9 documenta la primera aplicación de "
            "ese instrumental y, coherente con el principio de trazabilidad de "
            "las pautas documentales de la Dirección, entrega el código que "
            "permite reproducirlo.")


def marco_normativo(doc):
    titulo(doc, "4. Marco normativo e institucional", 1)

    titulo(doc, "4.1. La jerarquía normativa de la política digital", 2)
    parrafo(doc,
            "Antes de inventariar los instrumentos aplicables conviene situarlos "
            "unos respecto de otros, porque no todos obligan de la misma manera "
            "ni pueden invocarse indistintamente. La teoría clásica del "
            "ordenamiento jurídico, formulada por Hans Kelsen, sostiene que las "
            "normas se organizan en una estructura escalonada donde cada nivel "
            "deriva su validez del inmediatamente superior, de modo que ninguna "
            "disposición puede contradecir aquella de la que depende (Kelsen, "
            "1934/2009). Trasladada al caso que ocupa a este contrato, esa "
            "estructura permite entender por qué un documento CONPES orienta la "
            "acción del Ejecutivo pero no puede crear obligaciones reservadas a "
            "la ley, y por qué un plan de acción detalla compromisos sin alterar "
            "el marco que lo habilita.")
    figura(doc, "fig_kelsen.png",
           "Jerarquía normativa aplicable a la política digital y de "
           "inteligencia artificial en Colombia",
           "elaboración propia con base en la estructura escalonada del "
           "ordenamiento jurídico (Kelsen, 1934/2009) y en los instrumentos "
           "vigentes.")
    parrafo(doc,
            "La lectura de la figura deja tres consecuencias prácticas para el "
            "análisis. En primer lugar, los mandatos constitucionales de habeas "
            "data y acceso a la información no son un adorno introductorio: son "
            "el fundamento de validez de todo el andamiaje de gobernanza de datos "
            "sobre el que la política de inteligencia artificial se apoya, y "
            "cualquier acción que los comprometa es inválida por más justificada "
            "que esté técnicamente. En segundo lugar, el nivel legal es el que "
            "fija los límites de lo que la política puede hacer sin trámite "
            "adicional ante el Congreso, distinción decisiva para el eje de ética "
            "y gobernanza, cuyas acciones normativas dependen de si el instrumento "
            "elegido es un lineamiento administrativo o requiere reforma legal. Y "
            "en tercer lugar, el documento CONPES, pese a ser el instrumento "
            "central de este análisis, ocupa un nivel intermedio: es una decisión "
            "de política del Consejo Nacional de Política Económica y Social, con "
            "fuerza vinculante para las entidades que asumen sus acciones, pero "
            "sin capacidad de modificar leyes ni de crear tributos o "
            "restricciones de derechos.")
    parrafo(doc,
            "Esta última precisión importa especialmente al evaluar el "
            "cumplimiento, ya que varias de las acciones rezagadas del eje de "
            "gobernanza consisten en producir instrumentos que, para ser "
            "exigibles, deberán adoptar la forma de decreto o de proyecto de ley. "
            "El rezago, en esos casos, no es únicamente un problema de gestión "
            "sino también de trámite normativo, y la evaluación debe distinguir "
            "entre ambas causas antes de atribuir responsabilidades.")

    titulo(doc, "4.2. Instrumentos aplicables", 2)
    tabla(doc, "Instrumentos de política considerados en el análisis",
          ["Instrumento", "Año", "Pertinencia para el análisis"],
          [["Plan Nacional de Desarrollo 2022-2026",
            "2023",
            "Define los mandatos de transformación digital y de alfabetización digital con enfoque diferencial a los que deben alinearse los indicadores del sector."],
           ["CONPES 3920 – Explotación de datos",
            "2018",
            "Marco de gobernanza de datos públicos. Antecedente de la dimensión de datos e infraestructura del CONPES 4144."],
           ["CONPES 3975 – Transformación digital e IA",
            "2019",
            "Primer marco de política de IA en el país. Sirve de línea de comparación para medir el desplazamiento del enfoque de política."],
           ["Estrategia Nacional Digital 2023-2026",
            "2023",
            "Instrumento de implementación sectorial. Articula conectividad, gobierno digital, economía digital y talento digital."],
           ["CONPES 4144 – Política Nacional de IA",
            "2025",
            "Instrumento central. Define 106 acciones en seis ejes con inversión aproximada de 479.273 millones de pesos y horizonte 2030."]],
          "elaboración propia con base en los documentos citados.",
          anchos=[4.0, 1.6, 10.4])

    titulo(doc, "4.3. Los seis ejes del CONPES 4144 como unidades de análisis", 2)
    parrafo(doc,
            "La política se estructura en seis ejes estratégicos que operan como "
            "unidades de análisis del estudio (DNP, 2025). Su denominación oficial "
            "es la siguiente: ética y gobernanza; datos e infraestructura; "
            "investigación, desarrollo e innovación; desarrollo de capacidades y "
            "talento digital; mitigación de riesgos; y uso y adopción de la "
            "inteligencia artificial.")
    parrafo(doc,
            "El análisis de la vigencia 2026 se concentrará en los ejes primero, "
            "segundo, cuarto y sexto. El criterio de priorización es la "
            "disponibilidad de información pública con desagregación territorial "
            "suficiente y la correspondencia con las funciones misionales de la "
            "Dirección. Los ejes de investigación y de mitigación de riesgos se "
            "abordarán de forma cualitativa, porque sus resultados dependen de "
            "registros administrativos sectoriales que no ofrecen todavía series "
            "consultables.")

    titulo(doc, "4.4. Marco institucional y articulación del contrato", 2)
    parrafo(doc,
            "El CONPES 4144 asigna al DNP, en coordinación con el Ministerio de "
            "Ciencia, Tecnología e Innovación, el Ministerio de Tecnologías de la "
            "Información y las Comunicaciones y el Departamento Administrativo de "
            "la Presidencia de la República, el diseño y la ejecución de una "
            "estrategia anticipatoria para generar capacidades en prospectiva, "
            "vigilancia tecnológica e investigación sobre inteligencia artificial, "
            "así como mecanismos de articulación y seguimiento de la política "
            "(DNP, 2025).")
    parrafo(doc,
            "Conviene precisar el alcance de ese mandato, porque de ello depende "
            "cómo se presenta institucionalmente el trabajo. El documento CONPES "
            "no crea un observatorio ni lo menciona: encarga una estrategia "
            "anticipatoria y deja a la entidad responsable la definición de los "
            "instrumentos con los que la ejecuta. El Observatorio Nacional de "
            "Inteligencia Artificial es la figura mediante la cual la Dirección "
            "decide materializar esa estrategia. Es una decisión institucional "
            "legítima y trazable, pero su fundamento normativo debe enunciarse en "
            "esos términos y no como un mandato literal del documento CONPES.")
    parrafo(doc,
            "La segunda fuente de articulación es contractual. La modificación "
            "número 1 al contrato DNP-1025-2026 vinculó esta contratación al "
            "producto del Plan de Acción de la Dirección consistente en formular "
            "lineamientos con la propuesta de herramientas para incorporar "
            "enfoques diferenciales que promuevan la inclusión digital en el "
            "diseño, la implementación y la evaluación de iniciativas de política "
            "relacionadas con el desarrollo de habilidades digitales.")
    parrafo(doc,
            "Esa vinculación, lejos de ser un formalismo, ordena el trabajo: significa que la evaluación no es "
            "un ejercicio general sobre el sector, sino uno con foco definido: "
            "habilidades digitales, inclusión y enfoque diferencial. El capítulo 8 "
            "traduce ese foco en indicadores y el capítulo 9 lo prueba sobre datos "
            "reales.")
    parrafo(doc,
            "Como puede verse, los tres marcos convergen: el CONPES 4144 fija los ejes y las metas; "
            "el producto del Plan de Acción de la Dirección fija el foco "
            "diferencial; y el Observatorio Nacional de Inteligencia Artificial es "
            "el destinatario natural del inventario de fuentes y de la matriz de "
            "indicadores que este documento propone. El producto contractual "
            "alimenta esos tres niveles a la vez.")


def revision(doc):
    titulo(doc, "5. Revisión bibliográfica", 1)
    parrafo(doc,
            "La revisión se organiza en cuatro bloques que responden a las "
            "preguntas sucesivas del estudio. El primero recoge los fundamentos "
            "metodológicos de la evaluación, que definen cómo se mide. El segundo "
            "reúne la literatura sobre gobernanza de la inteligencia artificial y "
            "capacidad estatal, que define qué es razonable esperar de una "
            "política como la evaluada. El tercero inventaría los antecedentes "
            "nacionales de medición del sector, que aportan las líneas base. Y el "
            "cuarto sitúa los referentes internacionales que permiten comparar. En "
            "cada bloque se privilegian las fuentes que el propio CONPES 4144 "
            "utiliza como sustento, de modo que la evaluación dialogue con el "
            "mismo cuerpo de evidencia que la política y no con uno paralelo.")

    titulo(doc, "5.1. Fundamentos metodológicos de la evaluación", 2)
    parrafo(doc,
            "El texto de referencia obligada en evaluación de impacto sigue "
            "siendo el manual del Banco Mundial, que define el campo, sistematiza "
            "los diseños experimentales y cuasiexperimentales, y precisa las "
            "condiciones de validez de cada uno (Gertler et al., 2017). Su aporte "
            "central para este trabajo es la insistencia en que la calidad de una "
            "evaluación se decide en su diseño, antes de que existan datos: "
            "definir el indicador después de conocer el resultado es la vía más "
            "corta para producir un informe que confirme lo que ya se creía.")
    parrafo(doc,
            "En el plano de la construcción de indicadores, la referencia "
            "regional es el manual del Instituto Latinoamericano y del Caribe de "
            "Planificación Económica y Social, que además de la cadena de valor "
            "expuesta en el capítulo anterior aporta los criterios de calidad que "
            "debe cumplir un indicador para ser útil en la gestión pública "
            "—pertinencia, comparabilidad, oportunidad y economía en su "
            "producción— y advierte sobre la proliferación de indicadores que no "
            "se usan (Bonnefoy y Armijo, 2005). Esos criterios son los que el "
            "capítulo 8 aplica para depurar la matriz propuesta.")
    parrafo(doc,
            "Para el componente de análisis territorial, la literatura de "
            "economía regional ofrece el instrumental de medición de "
            "desigualdades espaciales, con los conceptos de concentración y "
            "convergencia que la Entrega 3 emplea para caracterizar la brecha "
            "digital entre departamentos. La aplicación de esos conceptos al "
            "sector digital, sin embargo, está todavía poco desarrollada en la "
            "literatura colombiana, lo que constituye en sí mismo un espacio de "
            "aporte para este producto.")

    titulo(doc, "5.2. Gobernanza de la inteligencia artificial y capacidad estatal", 2)
    parrafo(doc,
            "El campo de la gobernanza de la inteligencia artificial se consolidó "
            "en la última media década alrededor de un diagnóstico compartido: la "
            "regulación de esta tecnología enfrenta incertidumbre sobre efectos, "
            "asimetría de información frente a los desarrolladores y una "
            "velocidad de cambio que desactualiza los instrumentos durante su "
            "propia implementación (Taeihagh, 2021). Ese diagnóstico es "
            "especialmente pertinente para evaluar una política cuyo horizonte se "
            "extiende hasta 2030, porque obliga a preguntarse si sus acciones "
            "conservarán sentido a mitad de camino.")
    parrafo(doc,
            "En cuanto a la capacidad del Estado para gestionar esa "
            "transformación, la discusión sobre paradigmas administrativos aporta "
            "el contexto. La tesis del agotamiento de la Nueva Gestión Pública y "
            "su reemplazo por una gobernanza de la era digital, basada en la "
            "reintegración de funciones y la organización de servicios en torno "
            "al ciudadano, sigue orientando el análisis de las reformas "
            "contemporáneas del sector público (Dunleavy et al., 2006; Barzelay, "
            "2001). A ella se suma la lectura sobre el carácter sistémico y "
            "acelerado de la transformación tecnológica actual y sus exigencias "
            "institucionales (Schwab, 2016).")
    parrafo(doc,
            "En el ámbito iberoamericano, la traducción operativa de esas "
            "discusiones se ha dado a través de modelos de innovación pública. El "
            "Hexágono de Innovación Pública, construido sobre el análisis de "
            "ciento cinco metodologías y adoptado por la Secretaría General "
            "Iberoamericana, propone seis vectores de transformación "
            "organizacional (Oliván, 2020), y es el marco con el que la Dirección "
            "estructuró la propuesta del Observatorio Nacional de Inteligencia "
            "Artificial. La literatura sobre innovación en el sector público "
            "complementa ese modelo con la distinción entre innovación "
            "incremental y disruptiva, útil para calibrar las expectativas de "
            "resultado de una política de cuatro años (Bason, 2010), mientras que "
            "el ciclo de conversión del conocimiento organizacional aporta el "
            "esquema con el que un observatorio puede transformar información "
            "dispersa en práctica institucional (Nonaka y Takeuchi, 1995).")
    parrafo(doc,
            "Finalmente, el enfoque de anticipación aplicado al sector público "
            "sostiene que las instituciones deben desarrollar capacidad para "
            "considerar futuros alternativos como insumo de decisiones presentes, "
            "en lugar de reaccionar cuando los efectos se materializan (Miller, "
            "2018). Este es el fundamento conceptual de la acción del CONPES 4144 "
            "que da origen al observatorio.")

    titulo(doc, "5.3. Antecedentes nacionales de medición del sector", 2)
    parrafo(doc,
            "El antecedente más directo y más útil para este trabajo es el Índice "
            "de Pobreza Digital de la propia Dirección, construido sobre la "
            "Encuesta Nacional de Calidad de Vida y la Canasta Básica TIC, y "
            "planteado como una medición multidimensional de la exclusión digital "
            "de los hogares en tres dimensiones: conectividad, dispositivos y "
            "habilidades (Vanegas Barrero et al., 2024). Su valor es doble. En lo "
            "empírico, entrega la línea base de resultado que el sistema de "
            "seguimiento de la política no produce. En lo metodológico, demuestra "
            "que es posible construir mediciones nuevas sobre información "
            "estadística existente, sin depender de un levantamiento propio, que "
            "es exactamente la estrategia que este contrato adopta.")
    parrafo(doc,
            "Un segundo antecedente lo constituyen las mediciones sectoriales del "
            "Observatorio Nacional de TIC, que produce series de conectividad, "
            "apropiación y economía digital, incluido un índice de brecha digital "
            "para el período 2020-2024 (MinTIC, 2026). A diferencia de las "
            "encuestas, estas series provienen de registros administrativos de "
            "los operadores, lo que les da periodicidad trimestral y "
            "desagregación municipal a cambio de medir infraestructura contratada "
            "en lugar de uso efectivo; la combinación de ambos tipos de fuente es "
            "la que permite una lectura completa.")
    parrafo(doc,
            "El tercer antecedente es el conjunto de registros administrativos "
            "publicados en el portal de datos abiertos del Estado, cuya utilidad "
            "para el análisis con enfoque diferencial se demuestra en el capítulo "
            "9. Estos registros suelen carecer de representatividad estadística, "
            "pero incorporan variables de sexo, grupo poblacional y municipio que "
            "las operaciones agregadas no ofrecen, lo que los hace "
            "particularmente valiosos cuando la pregunta de política es "
            "distributiva y no de nivel.")
    parrafo(doc,
            "Debe señalarse, por último, un vacío relevante en la literatura "
            "nacional: no se identificaron evaluaciones de impacto publicadas "
            "sobre políticas de inteligencia artificial en Colombia, lo que se "
            "explica por la juventud del instrumento evaluado. Las evaluaciones "
            "disponibles se concentran en programas de conectividad y dotación "
            "tecnológica, de manera que este producto opera sobre un terreno con "
            "escasos precedentes metodológicos directos.")

    titulo(doc, "5.4. Referentes internacionales", 2)
    parrafo(doc,
            "El marco de medición de la transformación digital de la Organización "
            "para la Cooperación y el Desarrollo Económicos ofrece el sistema de "
            "indicadores multinivel más consolidado, con dimensiones de "
            "infraestructura, adopción, innovación e impacto económico; su "
            "formulación original data de 2019 y fue actualizada en la hoja de "
            "ruta de medición de 2026 (OCDE, 2019; OCDE, 2026). Su utilidad para "
            "este contrato es estructural más que numérica: aporta la arquitectura "
            "de dimensiones sobre la cual se organiza la matriz de indicadores del "
            "capítulo 8.")
    parrafo(doc,
            "La Unión Internacional de Telecomunicaciones, por su parte, publica "
            "anualmente los indicadores comparados de conectividad y uso de "
            "internet que permiten situar a Colombia frente a la región y frente "
            "a los países de referencia (UIT, 2024). En materia de posicionamiento "
            "comparado específicamente en inteligencia artificial, el propio "
            "CONPES 4144 se apoya en tres mediciones —el Global AI Index de "
            "Tortoise Media, el Government AI Readiness Index de Oxford Insights "
            "y el Índice Latinoamericano de Inteligencia Artificial del Centro "
            "Nacional de Inteligencia Artificial de Chile, las tres en su edición "
            "2023 (DNP, 2025)—, y conviene conservar esas mismas fuentes para que "
            "la comparación entre la línea base de la política y su evolución "
            "posterior resulte consistente.")
    parrafo(doc,
            "El Banco Interamericano de Desarrollo aporta el repositorio de "
            "evaluaciones y conjuntos de datos de acceso abierto sobre proyectos "
            "de desarrollo en la región, útil tanto como referencia metodológica "
            "para el diseño de indicadores como para la comparación con "
            "intervenciones equivalentes en otros países del continente (BID, "
            "2026). Su consulta sistemática, sugerida por la supervisión del "
            "contrato, quedó incorporada al inventario de fuentes del capítulo 7.")


def metodologia(doc):
    titulo(doc, "6. Metodología de esta entrega", 1)
    parrafo(doc,
            "El trabajo de esta entrega se organizó en cuatro operaciones. La "
            "primera fue la revisión documental de los instrumentos de política "
            "vigentes y de la literatura especializada en evaluación, con "
            "prioridad en las fuentes que el propio CONPES 4144 utiliza como "
            "referencia, de modo que la comparación posterior sea consistente.")
    parrafo(doc,
            "La segunda fue el inventario de fuentes de información. Cada fuente "
            "se caracterizó por contenido, periodicidad, nivel de desagregación y "
            "condición de acceso, y se verificó su disponibilidad efectiva "
            "consultándola en línea. El inventario distingue de forma explícita "
            "entre fuentes de acceso público inmediato y fuentes que requieren "
            "gestión institucional, porque de esa diferencia depende el "
            "cronograma de la Entrega 3.")
    parrafo(doc,
            "La tercera fue la propuesta de indicadores, construida a partir de "
            "los ejes priorizados del CONPES 4144 y sometida a cinco criterios de "
            "selección que se detallan en el capítulo 8.")
    parrafo(doc,
            "La cuarta fue la aplicación exploratoria de analítica de texto sobre "
            "el corpus de los documentos CONPES 3975 y 4144. El procedimiento, los "
            "parámetros y las limitaciones se documentan en el capítulo 9 y el "
            "código que lo reproduce se incluye como Anexo A.")
    parrafo(doc,
            "Dos criterios de calidad se aplicaron de forma transversal. Todas las "
            "cifras se verificaron contra la fuente primaria, y aquellas que no "
            "pudieron confirmarse se excluyeron del documento en lugar de "
            "matizarse. Todo uso de herramientas de inteligencia artificial "
            "generativa en la elaboración del informe se declara en el Anexo B, "
            "conforme a las pautas documentales de la Dirección.")
    parrafo(doc,
            "El documento tiene tres limitaciones que conviene enunciar de "
            "entrada. El análisis del avance físico y financiero de la política se "
            "limita a delimitar el universo evaluable, porque 40 de los 93 "
            "reportes de SisCONPES al corte están en revisión y sus valores pueden "
            "ajustarse. El ejercicio de analítica de texto es exploratorio y se "
            "limita a dos documentos. Y la prueba de viabilidad del capítulo 9 se "
            "realizó sobre un conjunto de datos de cobertura municipal, por lo que "
            "ilustra el método sin sustentar conclusiones nacionales.")


def fuentes(doc):
    titulo(doc, "7. Fuentes de información identificadas", 1)
    parrafo(doc,
            "El inventario se organiza en cuatro grupos según el origen y la "
            "condición de acceso de la información. Todas las fuentes relacionadas "
            "fueron consultadas durante la elaboración de este informe.")

    titulo(doc, "7.1. Repositorios sectoriales del MinTIC", 2)
    parrafo(doc,
            "El portal ColombiaTIC concentra la información estadística del sector "
            "y es la fuente sectorial de referencia. Su sección de indicadores "
            "organiza la información en cinco grupos —servicios de comunicaciones, "
            "infraestructura, apropiación de tecnologías, economía digital e "
            "indicadores de uso—, entre los que resulta de particular interés el "
            "índice de brecha digital para el período 2020-2024, por su "
            "correspondencia directa con el foco de inclusión de este contrato. La "
            "sección de publicaciones ofrece los informes y estudios sectoriales, "
            "incluidos los reportes de avance de las metas de gobierno, que "
            "permiten contrastar el dato estadístico con la lectura institucional "
            "del sector.")

    titulo(doc, "7.2. Datos abiertos del Estado colombiano", 2)
    parrafo(doc,
            "El portal de datos abiertos ofrece registros administrativos "
            "consultables mediante interfaz de programación, lo que permite "
            "automatizar la descarga y actualización de series sin intermediación. "
            "Su ventaja sobre las encuestas es la desagregación: varios conjuntos "
            "incorporan atributos de sexo, grupo poblacional y municipio que "
            "habilitan el análisis con enfoque diferencial que exige el producto "
            "del Plan de Acción de la Dirección. Su limitación es la cobertura, "
            "que suele corresponder a programas específicos y no al universo "
            "poblacional, como se documenta en el capítulo 9.")

    titulo(doc, "7.3. Fuentes del sistema estadístico y de seguimiento", 2)
    tabla(doc, "Fuentes institucionales y estadísticas identificadas",
          ["Fuente", "Contenido relevante", "Periodicidad", "Acceso"],
          [["Sinergia (DNP)", "Indicadores de seguimiento a metas del PND y a documentos CONPES, incluidas las acciones del CONPES 4144.", "Semestral", "Consulta pública; el detalle requiere gestión"],
           ["Plan de Acción y Seguimiento del CONPES 4144", "Las 106 acciones con objetivo, entidad responsable, dependencia, indicador, meta anual y costo programado por vigencia.", "Único, con ajustes", "Obtenido a través de la Dirección"],
           ["SisCONPES (DNP)", "Reportes de avance físico y financiero por acción, con estado de aprobación y periodo de reporte.", "Semestral", "Obtenido a través de la Dirección"],
           ["Índice de Pobreza Digital (DNP)", "Medición multidimensional de privaciones en conectividad, dispositivos y habilidades.", "Anual", "Pública"],
           ["ECV (DANE)", "Acceso a internet por hogar, con desagregación por cabecera y área rural dispersa y por departamento.", "Anual", "Pública"],
           ["Entic Hogares y Empresas (DANE)", "Tenencia y uso de TIC en hogares; adopción tecnológica en empresas, incluido el uso de IA.", "Anual", "Pública"],
           ["ColombiaTIC (MinTIC)", "Indicadores sectoriales de conectividad, apropiación y economía digital; índice de brecha digital.", "Trimestral y anual", "Pública"],
           ["Datos abiertos del Estado", "Registros administrativos de programas del sector, con desagregación territorial y poblacional.", "Variable", "Pública, con interfaz de programación"],
           ["Furag (Función Pública)", "Desempeño institucional en la dimensión de gestión del conocimiento e innovación.", "Anual", "Pública"]],
          "elaboración propia a partir de la consulta directa de cada fuente, agosto de 2026.",
          anchos=[3.4, 6.6, 2.6, 3.4])

    titulo(doc, "7.4. Fuentes internacionales comparadas", 2)
    parrafo(doc,
            "Tres fuentes internacionales completan el inventario. Los indicadores "
            "de la Unión Internacional de Telecomunicaciones permiten situar la "
            "conectividad nacional en el contexto regional. El marco de medición "
            "de la transformación digital de la OCDE aporta la estructura de "
            "dimensiones sobre la que se organizan los indicadores propuestos. Y "
            "el repositorio de datos abiertos del Banco Interamericano de "
            "Desarrollo ofrece conjuntos de datos y evaluaciones de proyectos "
            "comparables en la región, útiles como referencia metodológica.")

    titulo(doc, "7.5. Síntesis del inventario", 2)
    parrafo(doc,
            "En síntesis, el balance del inventario es favorable en lo público y pendiente en "
            "lo institucional. Las fuentes de acceso inmediato —encuestas del "
            "DANE, portal ColombiaTIC, datos abiertos, Índice de Pobreza Digital y "
            "referentes internacionales— bastan para construir las líneas base de "
            "los indicadores propuestos y para adelantar el análisis descriptivo y "
            "territorial de la Entrega 3.")
    parrafo(doc,
            "El seguimiento acción por acción de la política no puede resolverse "
            "con fuentes públicas. El Plan de Acción y Seguimiento, que contiene "
            "la entidad responsable, el indicador, la meta anual y el costo "
            "programado de cada una de las 106 acciones, se tramita como anexo en "
            "SisCONPES y no forma parte del documento CONPES de acceso público, "
            "cuyo Anexo A se publica sin contenido.")
    parrafo(doc,
            "Ese instrumento y los reportes de avance con corte al 21 de julio de "
            "2026 fueron facilitados por la Dirección durante la elaboración de "
            "este informe. Con ellos es posible analizar la distribución "
            "institucional de responsabilidades, la estructura de asignación de "
            "recursos y el avance diferenciado por objetivo, análisis que se "
            "presenta en el capítulo 10 y que constituye la base de la evaluación "
            "intermedia prevista para la Entrega 3.")


def indicadores(doc):
    titulo(doc, "8. Propuesta inicial de indicadores", 1)

    titulo(doc, "8.1. Criterios de selección", 2)
    parrafo(doc,
            "Los indicadores propuestos se seleccionaron aplicando cinco criterios. "
            "La pertinencia exige que el indicador mida el resultado esperado del "
            "eje y no una actividad intermedia. La disponibilidad exige que exista "
            "una fuente identificada en el capítulo 7 que lo alimente, sin depender "
            "de levantamientos nuevos. La comparabilidad exige que admita lectura "
            "territorial y temporal. La sensibilidad exige que pueda registrar "
            "cambios en un horizonte razonable de evaluación. Y la "
            "desagregabilidad exige que pueda descomponerse por sexo, zona o grupo "
            "poblacional, condición necesaria para el enfoque diferencial que "
            "ordena el producto del Plan de Acción de la Dirección.")
    parrafo(doc,
            "De los cinco, el último resultó el más restrictivo, pues llevó a descartar varios "
            "indicadores sectoriales de uso corriente que solo existen como "
            "agregado nacional.")

    titulo(doc, "8.2. Matriz de indicadores", 2)
    tabla(doc, "Indicadores de resultado e impacto propuestos",
          ["Eje", "Indicador", "Tipo", "Fuente"],
          [["1. Ética y gobernanza", "Avance de las acciones del CONPES 4144 sobre el total programado", "Resultado", "SisCONPES y Sinergia"],
           ["1. Ética y gobernanza", "Operación del Observatorio Nacional de IA y nivel de madurez alcanzado", "Resultado", "Furag y DENDD"],
           ["2. Datos e infraestructura", "Hogares con acceso a internet, por departamento y por zona", "Resultado", "DANE, ECV"],
           ["2. Datos e infraestructura", "Brecha de conectividad entre cabecera y área rural dispersa, en puntos porcentuales", "Impacto", "DANE, ECV"],
           ["2. Datos e infraestructura", "Incidencia de la pobreza digital y variación interanual", "Impacto", "DNP, IPD"],
           ["4. Capacidades y talento", "Población con brechas en habilidades digitales básicas", "Impacto", "DNP, IPD"],
           ["4. Capacidades y talento", "Personas formadas y certificadas en IA, por sexo y grupo poblacional", "Resultado", "Datos abiertos y MinTIC"],
           ["4. Capacidades y talento", "Tasa de certificación sobre inscripción en programas de formación en IA", "Resultado", "Datos abiertos"],
           ["6. Uso y adopción", "Empresas que usan IA en procesos productivos o de gestión, por tamaño y sector", "Impacto", "DANE, Entic Empresas"],
           ["6. Uso y adopción", "Entidades públicas con proyectos de IA implementados o en pilotaje", "Resultado", "Sinergia y MinTIC"],
           ["Transversal", "Posición de Colombia en las mediciones internacionales de referencia del CONPES 4144", "Impacto", "Tortoise, Oxford Insights y Cenia"]],
          "elaboración propia. La validación final de la matriz se realizará con la "
          "supervisión del contrato, según la disponibilidad efectiva de datos.",
          anchos=[3.4, 7.0, 2.0, 3.6])

    titulo(doc, "8.3. Sobre el indicador de tasa de certificación", 2)
    parrafo(doc,
            "El octavo indicador de la matriz merece una nota, porque no es de uso "
            "habitual en el sector. Los programas de formación suelen reportar "
            "personas inscritas o cupos ofrecidos, cifras que miden oferta y no "
            "aprendizaje. La relación entre quienes culminan y quienes se "
            "inscriben es más exigente y, sobre todo, admite lectura diferencial: "
            "permite distinguir si una brecha se produce en el acceso al programa "
            "o en la permanencia dentro de él, que son dos problemas de política "
            "distintos y requieren respuestas distintas. El capítulo 9 muestra por "
            "qué esa distinción no es teórica.")


def analitica(doc):
    titulo(doc, "9. Aplicación exploratoria de analítica de texto", 1)

    titulo(doc, "9.1. Procedimiento", 2)
    parrafo(doc,
            "Como primer uso de herramientas de analítica en el marco del "
            "contrato, se procesó el texto completo de los documentos CONPES 3975 "
            "de 2019 y 4144 de 2025, descargados de la biblioteca pública del DNP. "
            "La unidad de análisis es el párrafo. Se descartaron los bloques de "
            "menos de veinticinco palabras para excluir encabezados, pies de "
            "página y entradas de tabla de contenido, que no son texto "
            "argumentativo. El corpus resultante comprende 254 párrafos del "
            "documento de 2019 y 509 del de 2025.")
    parrafo(doc,
            "Sobre ese corpus se midieron dos cosas. La primera es la extensión "
            "temática: la proporción de párrafos que mencionan al menos un término "
            "de tres familias léxicas —territorial, de enfoque diferencial y de "
            "habilidades digitales—, comparadas sin tildes ni distinción de "
            "mayúsculas. La segunda es la frecuencia absoluta de un conjunto de "
            "términos de política, que permite observar el desplazamiento del "
            "vocabulario entre los dos documentos. El código que reproduce el "
            "ejercicio se incluye como Anexo A.")

    titulo(doc, "9.2. Resultados", 2)
    tabla(doc, "Extensión temática por familia léxica, en porcentaje de párrafos",
          ["Familia léxica", "CONPES 3975 (2019)", "CONPES 4144 (2025)"],
          [["Territorial", "5,1% (13 de 254)", "16,7% (85 de 509)"],
           ["Enfoque diferencial", "0,0% (0 de 254)", "7,1% (36 de 509)"],
           ["Habilidades y apropiación digital", "6,7% (17 de 254)", "12,6% (64 de 509)"]],
          "elaboración propia a partir del procesamiento de los documentos CONPES 3975 y 4144.",
          anchos=[6.0, 5.0, 5.0])

    tabla(doc, "Frecuencia absoluta de términos seleccionados",
          ["Término", "CONPES 3975 (2019)", "CONPES 4144 (2025)"],
          [["Inteligencia artificial", "17", "121"],
           ["Transformación digital", "80", "29"],
           ["Gobernanza", "1", "59"],
           ["Ética", "7", "126"],
           ["Riesgo", "9", "79"],
           ["Capacidades", "15", "96"],
           ["Datos", "42", "293"]],
          "elaboración propia a partir del procesamiento de los documentos CONPES 3975 y 4144.",
          anchos=[6.0, 5.0, 5.0])

    titulo(doc, "9.3. Lectura de los resultados", 2)
    parrafo(doc,
            "Los resultados muestran, en primer lugar, un desplazamiento nítido del eje discursivo entre los dos documentos. La transformación digital, que era el concepto articulador en "
            "2019 con ochenta menciones, cae a veintinueve en 2025, mientras la "
            "inteligencia artificial pasa de diecisiete a ciento veintiuna. El "
            "movimiento más marcado no está sin embargo en la tecnología sino en "
            "su encuadre: gobernanza pasa de una mención a cincuenta y nueve, y "
            "ética de siete a ciento veintiséis. La política de 2025 no solo habla "
            "de otra tecnología; habla de ella en términos de reglas y "
            "responsabilidad.")
    parrafo(doc,
            "El segundo hallazgo es, para efectos de este contrato, el más pertinente: la "
            "presencia de la dimensión territorial se triplica, al pasar del 5,1% "
            "al 16,7% de los párrafos, y el enfoque diferencial aparece por "
            "primera vez: está ausente por completo del documento de 2019 y "
            "alcanza el 7,1% de los párrafos en el de 2025, coherente con la "
            "creación del Grupo Interno de Trabajo de Enfoque Diferencial e "
            "Interseccional que el propio documento reseña. Las habilidades "
            "digitales casi duplican su presencia.")
    parrafo(doc,
            "La trayectoria, por tanto, es positiva y va en la dirección del "
            "producto al que está vinculado este contrato. El desafío ya no es "
            "que la política incorpore el enfoque diferencial y territorial en su "
            "formulación, porque lo hizo, sino que ese enfoque se traduzca en "
            "indicadores capaces de verificarlo en la implementación. Ahí es donde "
            "el aporte de este trabajo resulta pertinente.")
    parrafo(doc,
            "Estos resultados tienen, además de su valor descriptivo, una "
            "implicación directa para la matriz del capítulo 8: confirman que "
            "las familias léxicas del enfoque diferencial y de las habilidades "
            "digitales ya no son periféricas en el discurso de la política, de "
            "modo que exigirles indicadores desagregados no es imponerle a la "
            "política una vara ajena sino tomarle la palabra a su propia "
            "formulación.")

    titulo(doc, "9.4. Nota metodológica sobre el reproceso", 2)
    parrafo(doc,
            "La versión de junio de este informe reportó un resultado preliminar "
            "distinto, obtenido sobre una muestra parcial del corpus y sin el "
            "criterio de extensión mínima de párrafo. El procesamiento sobre el "
            "texto completo de los dos documentos, con el procedimiento descrito "
            "en el numeral 9.1, corrige esa cifra y la invierte: la presencia "
            "territorial no es baja, es tres veces mayor que en el instrumento "
            "anterior. Los resultados que aquí se reportan son los que reproduce "
            "el código del Anexo A.")
    parrafo(doc,
            "El informe de junio incluyó además una estimación sobre la "
            "concentración institucional de las acciones de la política, calculada "
            "sin acceso al instrumento que la sustenta. Esa estimación se "
            "reemplaza por el cálculo directo sobre el Plan de Acción y "
            "Seguimiento, que se presenta en el numeral siguiente.")

    titulo(doc, "9.5. Prueba de viabilidad sobre datos administrativos", 2)
    parrafo(doc,
            "Para verificar que el análisis con enfoque diferencial previsto para "
            "la Entrega 3 es factible con información pública, se procesó un "
            "registro administrativo de formación en inteligencia artificial "
            "disponible en el portal de datos abiertos, correspondiente a un "
            "programa ejecutado en 2025 con 1.549 participantes registrados "
            "(MinTIC, 2026b).")
    tabla(doc, "Resultados de la prueba de viabilidad",
          ["Variable", "Resultado"],
          [["Participantes registrados", "1.549"],
           ["Composición por sexo", "62,9% hombres y 37,1% mujeres"],
           ["Tasa de certificación general", "30,3%"],
           ["Tasa de certificación, mujeres", "30,1%"],
           ["Tasa de certificación, hombres", "30,4%"],
           ["Tasa de certificación por grupo poblacional",
            "80,0% en docentes, 34,1% en población general y 21,5% en estudiantes"],
           ["Cobertura territorial", "99,7% de los registros en un solo municipio"]],
          "elaboración propia a partir del conjunto de datos consultado en el portal "
          "de datos abiertos, agosto de 2026.", anchos=[6.5, 9.5])
    parrafo(doc,
            "El ejercicio arroja un resultado con implicación de política. La tasa "
            "de certificación de mujeres y hombres es prácticamente idéntica, con "
            "una diferencia de tres décimas de punto, mientras que las mujeres "
            "representan apenas el 37,1% de quienes se inscriben. Es decir que la "
            "brecha de género en esta intervención no se produce en la permanencia "
            "sino en el acceso: quienes ingresan culminan en la misma proporción, "
            "pero ingresan muchas menos. Una intervención dirigida a mejorar la "
            "retención no corregiría esa brecha, porque el problema está antes.")
    parrafo(doc,
            "La variación por grupo poblacional apunta en la misma dirección. Que "
            "los docentes culminen en el 80% de los casos y los estudiantes en el "
            "21,5% sugiere que el diseño del programa opera de manera muy distinta "
            "según el perfil del participante, algo que un indicador agregado de "
            "personas formadas ocultaría por completo.")
    parrafo(doc,
            "Ahora bien, el alcance de estas cifras es limitado y conviene decirlo "
            "con precisión: el 99,7% de los registros corresponde a un solo "
            "municipio, de modo que se trata de un programa de cobertura local y "
            "no de una medición nacional. El ejercicio demuestra que el método "
            "funciona y que los datos existen con la desagregación necesaria; no "
            "sustenta conclusiones sobre el país. El conjunto presenta además "
            "inconsistencias menores de normalización en los nombres de "
            "departamento, del tipo que la Entrega 3 deberá depurar de forma "
            "sistemática.")


def plan_accion(doc):
    titulo(doc, "10. Estructura del plan de acción y asignación de recursos", 1)
    parrafo(doc,
            "El Plan de Acción y Seguimiento facilitado por la Dirección permite "
            "examinar la política por dentro: cómo se reparten las 106 acciones "
            "entre objetivos y entidades, con qué tipo de indicador se mide cada "
            "una y cómo se distribuyen los recursos. El procesamiento del "
            "instrumento arroja un costo total de 479.042 millones de pesos, cifra "
            "que coincide con la reportada en el documento CONPES con una "
            "diferencia del 0,05%, atribuible a tres registros con referencias "
            "incompletas en el archivo.")

    titulo(doc, "10.1. Distribución institucional de las acciones", 2)
    parrafo(doc,
            "En la ejecución de la política participan 56 entidades, pero la "
            "distribución dista de ser uniforme. El Ministerio de Tecnologías de la "
            "Información y las Comunicaciones interviene en 64 de las 106 acciones "
            "y el Departamento Nacional de Planeación en 52. Sumando las acciones "
            "en que participa al menos una de las tres entidades con mayor "
            "presencia —MinTIC, DNP y el Ministerio de Ciencia, Tecnología e "
            "Innovación— se cubren 85 acciones, el 80,2% del plan. La mitad de "
            "las entidades restantes figura en una sola acción.")
    parrafo(doc,
            "La lectura para el seguimiento es directa: la capacidad de ejecución "
            "de tres entidades determina el desempeño de cuatro quintas partes de "
            "la política, lo que concentra el riesgo de implementación y hace del "
            "mecanismo de coordinación interinstitucional un factor crítico. La "
            "Dirección de Economía Naranja y Desarrollo Digital participa en 39 "
            "acciones, de las cuales 12 tienen vigencia hasta 2026.")

    titulo(doc, "10.2. Asignación de recursos por objetivo", 2)
    tabla(doc, "Acciones y recursos programados por objetivo, 2025-2030",
          ["Objetivo", "Acciones", "Recursos (millones de pesos)", "Participación"],
          [["2. Datos e infraestructura", "30", "284.117", "59,3%"],
           ["6. Uso y adopción de la IA", "29", "157.050", "32,8%"],
           ["4. Capacidades y talento digital", "14", "18.322", "3,8%"],
           ["3. Investigación, desarrollo e innovación", "6", "11.160", "2,3%"],
           ["5. Mitigación de riesgos", "21", "7.759", "1,6%"],
           ["1. Ética y gobernanza", "6", "634", "0,1%"],
           ["Total", "106", "479.042", "100%"]],
          "elaboración propia a partir del Plan de Acción y Seguimiento del CONPES 4144.",
          anchos=[6.0, 2.2, 4.4, 3.4])
    parrafo(doc,
            "La distribución revela una tensión que conviene poner sobre la mesa, "
            "porque toca directamente el foco de este contrato. La medición "
            "oficial de la Dirección establece que la privación digital dominante "
            "en el país no es la infraestructura sino la capacidad de uso: el 60% "
            "de la población presenta brechas en habilidades digitales básicas, "
            "frente al 39% con deficiencias de conectividad. La asignación de "
            "recursos de la política sigue el orden inverso, con el 59,3% "
            "destinado a datos e infraestructura y el 3,8% a capacidades y "
            "talento digital.")
    parrafo(doc,
            "La observación no implica que la asignación sea equivocada. La "
            "infraestructura tiene costos unitarios mucho más altos que la "
            "formación, de modo que una participación mayor en el presupuesto no "
            "equivale a una prioridad mayor en la política; y sin red no hay uso "
            "posible. Lo que sí exige la diferencia es una hipótesis explícita "
            "sobre cómo la inversión en infraestructura se traduce en cierre de la "
            "brecha de habilidades, hipótesis que hoy no está formulada en el "
            "instrumento y que la evaluación debería poner a prueba.")

    titulo(doc, "10.3. Naturaleza de los indicadores de seguimiento", 2)
    parrafo(doc,
            "El plan asigna a cada acción un indicador de cumplimiento clasificado "
            "por tipo. De los 106 indicadores, 76 son de gestión y 30 de producto. "
            "Ninguno es de resultado. Es decir que el instrumento de seguimiento "
            "está diseñado para verificar que las acciones se ejecuten y que los "
            "productos se entreguen, no para establecer si esos productos "
            "modifican la situación que la política busca corregir.")
    parrafo(doc,
            "Esa constatación confirma con evidencia del propio instrumento el "
            "argumento planteado en la introducción de este informe. La distancia "
            "entre seguimiento y evaluación no es una carencia de la "
            "implementación sino una característica del diseño, y explica por qué "
            "un análisis de resultados e impacto como el previsto en este contrato "
            "requiere construir indicadores adicionales a los que el plan "
            "contempla. La matriz del capítulo 8 responde a esa necesidad.")

    titulo(doc, "10.4. Horizonte de la vigencia 2026", 2)
    parrafo(doc,
            "El calendario del plan concentra el esfuerzo en el corto plazo. "
            "Dieciocho acciones tenían vigencia hasta 2025 y treinta la tienen "
            "hasta 2026, de modo que 48 acciones —el 45,3% del plan— vencen a más "
            "tardar en diciembre de este año. En recursos, los años 2025 y 2026 "
            "concentran 305.013 millones de pesos, el 63,7% del costo total.")
    parrafo(doc,
            "El reporte de avance de SisCONPES con corte al 21 de julio de 2026 "
            "muestra que 93 acciones tienen reporte registrado, con un avance "
            "físico promedio del 43,3% sobre las acciones reportadas, y que de las "
            "18 acciones con vigencia hasta 2025 solo tres cerraron con "
            "cumplimiento total. Estas cifras se toman del documento interno de "
            "seguimiento de la Dirección y su análisis detallado corresponde a la "
            "Entrega 3; se incorporan aquí porque delimitan el universo evaluable "
            "de la vigencia.")
    parrafo(doc,
            "Dos precauciones metodológicas se derivan de la revisión de los "
            "reportes. La primera es que conviven dos periodos de reporte, 2025-1 "
            "y 2025-2, con valores distintos para una misma acción, por lo que "
            "toda cifra debe citarse indicando el corte. La segunda es que 40 de "
            "los 93 reportes figuran en revisión y no aprobados, de modo que sus "
            "valores están sujetos a ajuste. El análisis de la Entrega 3 deberá "
            "distinguir ambos conjuntos en lugar de promediarlos sin advertencia.")


def brecha(doc):
    titulo(doc, "11. Situación deseada y punto de partida", 1)
    parrafo(doc,
            "Las pautas de elaboración documental de la Dirección piden que un "
            "informe técnico explicite qué situación se busca alcanzar y en qué "
            "punto se encuentra. Este capítulo lo hace para el objeto del "
            "contrato, que no es el desempeño del sector TIC en abstracto sino la "
            "capacidad de la Dirección para evaluar sus políticas con evidencia.")
    tabla(doc, "Distancia entre la situación deseada y el punto de partida",
          ["Dimensión", "Situación deseada", "Punto de partida (agosto de 2026)"],
          [["Marco de medición",
            "Un conjunto de indicadores de resultado e impacto acordado, con línea base y fuente definida para cada uno.",
            "Matriz de once indicadores propuesta en este informe, pendiente de validación con la supervisión."],
           ["Disponibilidad de datos",
            "Acceso corriente a las fuentes públicas e institucionales que alimentan cada indicador.",
            "Fuentes públicas verificadas y disponibles. Plan de Acción y Seguimiento y reportes de SisCONPES obtenidos a través de la Dirección."],
           ["Seguimiento de la política",
            "Lectura del avance de las 106 acciones por eje y por entidad responsable.",
            "Estructura y recursos ya analizados (capítulo 10). El análisis del avance físico y financiero por acción corresponde a la Entrega 3."],
           ["Enfoque diferencial",
            "Indicadores desagregados por sexo, zona y grupo poblacional en los ejes priorizados.",
            "Demostrado viable sobre registros administrativos; aún no incorporado a las series sectoriales."],
           ["Analítica aplicada",
            "Procesamiento sistemático del corpus de política y de las series sectoriales con resultados reproducibles.",
            "Ejercicio exploratorio realizado sobre dos documentos, con código reproducible."]],
          "elaboración propia.", anchos=[3.2, 6.4, 6.4])
    parrafo(doc,
            "La lectura de conjunto es que las condiciones para la Entrega 3 están "
            "dadas. El marco conceptual está definido, los indicadores están "
            "propuestos, las fuentes públicas alcanzan para construir las líneas "
            "base y el instrumento de seguimiento de la política ya está "
            "disponible y procesado. La restricción que subsiste es de "
            "desagregación: lograr que las series sectoriales admitan la lectura "
            "por sexo, zona y grupo poblacional que el enfoque diferencial exige, "
            "algo que la prueba de viabilidad demostró posible en los registros "
            "administrativos, pero que aún no es la norma en las fuentes del "
            "sector.")


def conclusiones(doc):
    titulo(doc, "12. Conclusiones", 1)
    parrafo(doc,
            "El marco conceptual y normativo del análisis quedó establecido. La "
            "cadena de valor de las intervenciones públicas ordena los indicadores "
            "según lo que efectivamente miden, los seis ejes del CONPES 4144 "
            "operan como unidades de análisis y los ejes primero, segundo, cuarto "
            "y sexto se priorizaron por disponibilidad de datos y correspondencia "
            "misional.")
    parrafo(doc,
            "El fundamento normativo del Observatorio Nacional de Inteligencia "
            "Artificial quedó precisado. El documento CONPES 4144 encarga al DNP "
            "una estrategia anticipatoria en prospectiva y vigilancia tecnológica, "
            "y el Observatorio es el instrumento con que la Dirección decide "
            "ejecutarla. Enunciarlo así protege la solidez del argumento "
            "institucional.")
    parrafo(doc,
            "La base de información quedó completa. Las fuentes públicas bastan "
            "para construir las líneas base de los indicadores propuestos, y el "
            "Plan de Acción y Seguimiento, junto con los reportes de SisCONPES al "
            "corte de julio de 2026, fue obtenido a través de la Dirección y "
            "procesado, lo que habilita la evaluación intermedia de la Entrega 3.")
    parrafo(doc,
            "El análisis del plan de acción arrojó tres constataciones. La "
            "ejecución está concentrada: tres entidades participan en el 80,2% de "
            "las acciones. La asignación de recursos privilegia la infraestructura "
            "sobre las capacidades en una proporción de quince a uno, mientras el "
            "diagnóstico oficial señala las habilidades digitales como la "
            "privación dominante, lo que exige explicitar la hipótesis que conecta "
            "una cosa con la otra. Y el sistema de seguimiento no contiene ningún "
            "indicador de resultado: 76 indicadores son de gestión y 30 de "
            "producto, lo que confirma que la medición de efectos debe "
            "construirse por fuera del instrumento, como propone la matriz del "
            "capítulo 8.")
    parrafo(doc,
            "El análisis del corpus de política muestra una trayectoria positiva y "
            "medible entre 2019 y 2025. La dimensión territorial triplicó su "
            "presencia y el enfoque diferencial, ausente del instrumento anterior, "
            "aparece en el 7,1% de los párrafos del vigente. El reto se desplazó "
            "de la formulación a la verificación: falta traducir ese enfoque en "
            "indicadores que permitan comprobarlo durante la implementación.")
    parrafo(doc,
            "La prueba de viabilidad confirmó que el análisis con enfoque "
            "diferencial es factible con datos públicos y arrojó un hallazgo "
            "sustantivo, aunque de alcance local: en el programa de formación "
            "analizado, la brecha de género se produce en el acceso y no en la "
            "permanencia, porque hombres y mujeres culminan en la misma "
            "proporción pero ellas se inscriben mucho menos.")
    parrafo(doc,
            "Finalmente, el punto de partida del sector está documentado con "
            "cifras oficiales. El 65,6% de los hogares tiene acceso a internet, "
            "con una brecha de más de treinta puntos entre cabeceras y área rural "
            "dispersa, y el 37,9% de la población vive en pobreza digital, siendo "
            "las habilidades digitales la dimensión más crítica, con el 60% de la "
            "población afectada.")


def recomendaciones(doc):
    titulo(doc, "13. Recomendaciones", 1)
    parrafo(doc,
            "Establecer con la supervisión un canal regular de actualización de "
            "los reportes de SisCONPES, de modo que la Entrega 3 y la entrega "
            "final trabajen sobre cortes comparables. La revisión de los reportes "
            "disponibles mostró que conviven periodos distintos con valores "
            "distintos para una misma acción, y que 40 de los 93 reportes están "
            "aún en revisión; fijar el corte y el estado de aprobación como "
            "convención de cita evita inconsistencias entre entregas.")
    parrafo(doc,
            "Validar la matriz de indicadores del capítulo 8 con la supervisión "
            "antes de iniciar el procesamiento de datos, de modo que la Entrega 3 "
            "trabaje sobre un conjunto acordado y no sobre una propuesta. Plazo "
            "sugerido: dos semanas.")
    parrafo(doc,
            "Adoptar la tasa de certificación sobre inscripción como indicador de "
            "resultado en los programas de formación en inteligencia artificial, y "
            "exigir su desagregación por sexo y grupo poblacional. El hallazgo del "
            "capítulo 9 muestra que sin esa desagregación el diagnóstico de una "
            "brecha puede apuntar al problema equivocado.")
    parrafo(doc,
            "Documentar, con las áreas responsables, el vínculo entre el "
            "Observatorio Nacional de Inteligencia Artificial y la acción de "
            "estrategia anticipatoria del CONPES 4144, en los términos precisados "
            "en el capítulo 4. Es un ajuste de redacción institucional de bajo "
            "costo que evita una objeción previsible en instancias de revisión.")
    parrafo(doc,
            "Formular, en el marco de la evaluación de la Entrega 3, la hipótesis "
            "de transmisión entre la inversión en infraestructura y el cierre de "
            "la brecha de habilidades digitales, de modo que la tensión "
            "identificada en el capítulo 10 entre la asignación de recursos y el "
            "diagnóstico oficial pueda examinarse con evidencia y no quede como "
            "una observación sin respuesta.")
    parrafo(doc,
            "Incorporar el inventario de fuentes del capítulo 7 y la base "
            "procesada del Plan de Acción y Seguimiento al acervo documental del "
            "Observatorio, de modo que su utilidad exceda la vigencia de este "
            "contrato.")


def referencias(doc):
    titulo(doc, "Referencias", 1)
    refs = [
        "Banco Interamericano de Desarrollo. (2026). Datos abiertos del BID [repositorio de conjuntos de datos]. https://data.iadb.org/dataset/",
        "Barzelay, M. (2001). The new public management: Improving research and policy dialogue. University of California Press.",
        "Bason, C. (2010). Leading public sector innovation: Co-creating for a better society. Policy Press.",
        "Bonnefoy, J. C., y Armijo, M. (2005). Indicadores de desempeño en el sector público (Serie Manuales n.º 45). Instituto Latinoamericano y del Caribe de Planificación Económica y Social, Cepal.",
        "Departamento Nacional de Planeación. (2000). Documento CONPES 3072. Agenda de conectividad. DNP.",
        "Departamento Administrativo Nacional de Estadística. (2025). Encuesta Nacional de Calidad de Vida 2024: boletín técnico. DANE.",
        "Departamento Nacional de Planeación. (2018). Documento CONPES 3920. Política nacional de explotación de datos (big data). DNP.",
        "Departamento Nacional de Planeación. (2019). Documento CONPES 3975. Política nacional para la transformación digital e inteligencia artificial. DNP.",
        "Departamento Nacional de Planeación. (2023). Estrategia Nacional Digital de Colombia 2023-2026. DNP.",
        "Departamento Nacional de Planeación. (2025). Documento CONPES 4144. Política nacional de inteligencia artificial. DNP. https://colaboracion.dnp.gov.co/CDT/Conpes/Económicos/4144.pdf",
        "Departamento Nacional de Planeación. (2025b). Índice de Pobreza Digital: resultados. DNP.",
        "Departamento Nacional de Planeación. (2025c). Anexo A. Plan de Acción y Seguimiento del documento CONPES 4144 [instrumento de seguimiento, SisCONPES]. DNP.",
        "Departamento Nacional de Planeación. (2026). Reporte y revisión del documento CONPES 4144 [módulo de SisCONPES, corte del 21 de julio de 2026]. DNP.",
        "Dunleavy, P., Margetts, H., Bastow, S., y Tinkler, J. (2006). New public management is dead—Long live digital-era governance. Journal of Public Administration Research and Theory, 16(3), 467-494.",
        "Gertler, P. J., Martínez, S., Premand, P., Rawlings, L. B., y Vermeersch, C. M. J. (2017). La evaluación de impacto en la práctica (2.ª ed.). Banco Mundial.",
        "Miller, R. (Ed.). (2018). Transforming the future: Anticipation in the 21st century. Unesco y Routledge.",
        "Ministerio de Tecnologías de la Información y las Comunicaciones. (2026). Indicadores y publicaciones del sector TIC [portal ColombiaTIC]. https://colombiatic.mintic.gov.co/",
        "Ministerio de Tecnologías de la Información y las Comunicaciones. (2026b). Participantes certificados en inteligencia artificial – SENATIC [conjunto de datos]. Portal de Datos Abiertos de Colombia. https://www.datos.gov.co/",
        "Kelsen, H. (2009). Teoría pura del derecho (2.ª ed.; R. Vernengo, Trad.). Universidad Nacional Autónoma de México. (Obra original publicada en 1934).",
        "Nonaka, I., y Takeuchi, H. (1995). The knowledge-creating company: How Japanese companies create the dynamics of innovation. Oxford University Press.",
        "Oliván, R. (2020). Modelo HIP: Hexágono de Innovación Pública. Secretaría General Iberoamericana y Laboratorio de Aragón Gobierno Abierto.",
        "Organización para la Cooperación y el Desarrollo Económicos. (2019). Measuring the digital transformation: A roadmap for the future. OECD Publishing.",
        "Organización para la Cooperación y el Desarrollo Económicos. (2026). The OECD going digital measurement roadmap 2026. OECD Publishing.",
        "Schwab, K. (2016). La cuarta revolución industrial. Foro Económico Mundial y Debate.",
        "Taeihagh, A. (2021). Governance of artificial intelligence. Policy and Society, 40(2), 137-157.",
        "Unión Internacional de Telecomunicaciones. (2024). Measuring digital development: Facts and figures 2024. UIT.",
        "Vanegas Barrero, V., Dávila Barragán, J., y Barreto Nieto, C. A. (2024, noviembre). Pobreza digital, una nueva perspectiva para fortalecer la inclusión social en Colombia. Planeación & Desarrollo. Departamento Nacional de Planeación.",
    ]
    for r in refs:
        parrafo(doc, r, sangria_francesa=True, space_after=6)


def anexos(doc):
    titulo(doc, "Anexos", 1)

    titulo(doc, "Anexo A. Código de reproducción del análisis de corpus", 2)
    parrafo(doc,
            "El ejercicio reportado en el capítulo 9 se reproduce con la rutina "
            "analisis_corpus.py, que descarga los dos documentos CONPES de la "
            "biblioteca pública del DNP, extrae su texto, segmenta por párrafos con "
            "el umbral de veinticinco palabras y calcula las dos medidas reportadas. "
            "La rutina no tiene dependencias externas más allá de la utilidad "
            "pdftotext y se ejecuta con la instrucción python3 analisis_corpus.py. "
            "Se entrega como archivo adjunto a este informe.")
    parrafo(doc,
            "Los parámetros que determinan el resultado son tres y quedan "
            "explícitos en el código: el umbral de extensión mínima de párrafo, la "
            "composición de cada familia léxica y el criterio de comparación sin "
            "tildes ni distinción de mayúsculas. Modificar cualquiera de ellos "
            "altera las cifras, razón por la cual se documentan.")

    titulo(doc, "Anexo B. Declaración de uso de herramientas de inteligencia artificial", 2)
    parrafo(doc,
            "Conforme al numeral sobre referencias de las Pautas para la "
            "elaboración, revisión y entrega de documentos técnicos de la "
            "Dirección, se declara el uso de herramientas de inteligencia "
            "artificial generativa como apoyo en la organización de la estructura "
            "del documento, en la redacción de borradores y en el procesamiento "
            "del corpus documental del capítulo 9.")
    parrafo(doc,
            "Todo contenido generado con ese apoyo fue verificado por el autor "
            "contra las fuentes primarias citadas en el capítulo de referencias. "
            "Las cifras que no pudieron confirmarse contra fuente primaria fueron "
            "retiradas del documento, según se explica en el numeral 9.4. Las "
            "herramientas empleadas no son autoras del documento ni sustituyen la "
            "consulta de las fuentes, y sus resultados se tratan como producto de "
            "un procedimiento y no como fuente de hechos.")


# --------------------------------------------------------------------------

def construir(salida):
    _CONTADOR_TABLAS[0] = 0
    _CONTADOR_FIGURAS[0] = 0
    doc = Document()
    configurar_estilos(doc)

    s0 = doc.sections[0]
    margenes(s0)
    s0.different_first_page_header_footer = True
    portada(doc,
            "Marco conceptual, normativo y de fuentes para el análisis de datos "
            "y la evaluación de impacto de las políticas públicas del sector TIC",
            "Informe técnico – Entrega 2 – Vigencia 2026", "Versión 2")

    s1 = doc.add_section(WD_SECTION.NEW_PAGE)
    margenes(s1)
    encabezado_paginado(s1)

    pagina_legal(doc)
    salto(doc)
    tabla_contenido(doc)
    salto(doc)
    siglas(doc)
    salto(doc)
    presentacion(doc)
    salto(doc)
    introduccion(doc)
    antecedentes(doc)
    marco_conceptual(doc)
    marco_normativo(doc)
    revision(doc)
    metodologia(doc)
    fuentes(doc)
    indicadores(doc)
    analitica(doc)
    plan_accion(doc)
    brecha(doc)
    conclusiones(doc)
    recomendaciones(doc)
    salto(doc)
    referencias(doc)
    salto(doc)
    anexos(doc)

    os.makedirs(os.path.dirname(salida), exist_ok=True)
    doc.save(salida)
    return salida


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--salida", default=SALIDA_DEFECTO)
    ap.add_argument("--figs", default=".", help="directorio con las figuras PNG")
    args = ap.parse_args()
    DIR_FIGURAS[0] = args.figs
    ruta = compilar(construir, args.salida)
    print(f"documento generado: {ruta}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
