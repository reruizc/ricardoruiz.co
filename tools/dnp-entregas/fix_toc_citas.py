#!/usr/bin/env python3
"""
Ajustes in situ sobre los .docx de las entregas 2 y 3 (sin regenerarlos, para
preservar las ediciones manuales del autor):

1. Las líneas de la tabla de contenido y de las listas de tablas y figuras
   pasan a los estilos de Word "TOC 1/2/3" (que se crean si no existen), con
   interlineado sencillo, sin espacio entre párrafos y sin negrita, como las
   tablas de contenido nativas de Word. Las líneas se identifican por su
   tabulación con relleno de puntos, que solo esas usan.

2. La referencia de la Entrega 2 deja de atribuirse al contratista y pasa a la
   Dirección de Economía Naranja y Desarrollo Digital, reubicándose en su
   posición alfabética.

Uso: python3 fix_toc_citas.py <archivo.docx> [archivo2.docx ...]
"""

import sys

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_TAB_LEADER
from docx.shared import Cm, Pt, RGBColor

REF_VIEJA = "Ruiz Castro, R. E. (2026)."
REF_NUEVA = ("Dirección de Economía Naranja y Desarrollo Digital. (2026). "
             "Marco conceptual, normativo y de fuentes para el análisis de "
             "datos y la evaluación de impacto de las políticas públicas del "
             "sector TIC (Informe técnico, Entrega 2, contrato DNP-1025-2026). "
             "Departamento Nacional de Planeación.")


def asegurar_estilos_toc(doc):
    nombres = {s.name for s in doc.styles}
    for nivel in (1, 2, 3):
        nombre = f"TOC {nivel}"
        if nombre in nombres:
            st = doc.styles[nombre]
        else:
            st = doc.styles.add_style(nombre, WD_STYLE_TYPE.PARAGRAPH,
                                      builtin=True)
            st.base_style = doc.styles["Normal"]
        st.font.name = "Arial"
        st.font.size = Pt(10)
        st.font.bold = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        pf = st.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(0.5 * (nivel - 1))
        pf.first_line_indent = Cm(0)


def es_linea_toc(p):
    return any(ts.leader == WD_TAB_LEADER.DOTS
               for ts in p.paragraph_format.tab_stops)


def nivel_de(p):
    ind = p.paragraph_format.left_indent
    cm = ind.cm if ind is not None else 0.0
    if cm >= 0.85:
        return 3
    if cm >= 0.35:
        return 2
    return 1


def arreglar_toc(doc):
    n = 0
    for p in doc.paragraphs:
        if not es_linea_toc(p):
            continue
        nivel = nivel_de(p)
        p.style = doc.styles[f"TOC {nivel}"]
        pf = p.paragraph_format
        # Se retira el formato directo para que gobierne el estilo; la
        # tabulación con puntos se conserva porque vive en el pPr directo.
        pf.left_indent = None
        pf.first_line_indent = None
        pf.space_before = None
        pf.space_after = None
        pf.line_spacing = None
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(10)
            r.font.bold = False
        n += 1
    return n


def arreglar_referencia(doc):
    origen = None
    destino = None
    for p in doc.paragraphs:
        if p.text.startswith(REF_VIEJA):
            origen = p
        if p.text.startswith("Dunleavy, P.") and destino is None:
            destino = p
    if origen is None:
        return False
    for r in list(origen.runs[1:]):
        r._r.getparent().remove(r._r)
    origen.runs[0].text = REF_NUEVA
    if destino is not None:
        destino._p.addprevious(origen._p)
    return True


def procesar(ruta):
    doc = docx.Document(ruta)
    asegurar_estilos_toc(doc)
    n = arreglar_toc(doc)
    ref = arreglar_referencia(doc)
    doc.save(ruta)
    print(f"{ruta}: {n} líneas de contenido/listas reformateadas"
          f"{' | referencia reatribuida a la DENDD' if ref else ''}")


if __name__ == "__main__":
    for ruta in sys.argv[1:]:
        procesar(ruta)
