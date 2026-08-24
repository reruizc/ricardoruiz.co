#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Documento Word (10-12 pág) para Tatiana Villarreal:
Cartagena a nivel de barrio — dónde concentrar el esfuerzo de 2ª vuelta (campaña Abelardo).
Incrusta Inter para tipografía estable en cualquier Office.
"""
import json, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "Bases de datos/output_abelardo_cartagena"
RES = json.load(open(f"{OUT}/cartagena_resumen.json"))
LOC = json.load(open(f"{OUT}/cartagena_localidades.json"))
BAR = json.load(open(f"{OUT}/cartagena_barrios.json"))
LAB = json.load(open(f"{OUT}/barrio_labels.json"))

NF = lambda n: f"{round(n):,}".replace(",", ".")   # miles con punto (Colombia)
NAVY = RGBColor(0x16, 0x24, 0x5e)
ABE = RGBColor(0x1f, 0x47, 0xcc)
INK = RGBColor(0x1a, 0x1a, 0x2e)
GR = RGBColor(0x5a, 0x54, 0x48)
FONT = "Inter"
HEXNAVY = "16245E"

d = Document()
# margenes
for s in d.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)
for s in d.styles:
    if s.name == "Normal": s.font.name = FONT; s.font.size = Pt(10.5)
_st = d.styles.element; _dd = _st.find(qn("w:docDefaults"))
if _dd is not None:
    _rprd = _dd.find(qn("w:rPrDefault"))
    if _rprd is None: _rprd = OxmlElement("w:rPrDefault"); _dd.insert(0, _rprd)
    _rpr = _rprd.find(qn("w:rPr"))
    if _rpr is None: _rpr = OxmlElement("w:rPr"); _rprd.append(_rpr)
    _rf = _rpr.find(qn("w:rFonts"))
    if _rf is None: _rf = OxmlElement("w:rFonts"); _rpr.insert(0, _rf)
    for _a in ("w:ascii", "w:hAnsi", "w:cs"): _rf.set(qn(_a), FONT)


def shade(c, h):
    e = OxmlElement("w:shd"); e.set(qn("w:fill"), h); c._tc.get_or_add_tcPr().append(e)


def sech(num, txt):
    p = d.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.page_break_before = True
    eyebrow = "ANEXO" if num == "Anexo" else f"SECCIÓN {num}"
    r = p.add_run(eyebrow); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GR
    p2 = d.add_paragraph(); r = p2.add_run(txt); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    p2.paragraph_format.space_after = Pt(6)


def h(txt, size=12.5, color=INK, before=8, after=4):
    p = d.add_paragraph(); r = p.add_run(txt); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)


def body(txt, size=10.5, after=6):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in re.split(r"(\*\*.*?\*\*)", txt):
        r = p.add_run(seg[2:-2] if seg.startswith("**") else seg)
        r.font.size = Pt(size); r.bold = seg.startswith("**"); r.font.color.rgb = INK


def note(txt): body(txt, size=8.8, after=5)


def bullet(txt):
    p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in re.split(r"(\*\*.*?\*\*)", txt):
        r = p.add_run(seg[2:-2] if seg.startswith("**") else seg)
        r.font.size = Pt(10.5); r.bold = seg.startswith("**"); r.font.color.rgb = INK


def tbl(headers, rows, widths=None, fs=8.4):
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; r = c.paragraphs[0].add_run(hh)
        r.bold = True; r.font.size = Pt(fs); r.font.color.rgb = RGBColor(255, 255, 255); shade(c, HEXNAVY)
    for ri, row in enumerate(rows):
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""; r = cs[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(fs)
            if ri % 2 == 1: shade(cs[i], "F1EEE6")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    return t


def img(name, w=6.5, cap=None):
    d.add_picture(f"{OUT}/{name}", width=Inches(w)); p = d.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(2)
    if cap:
        pc = d.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pc.add_run(cap); r.italic = True; r.font.size = Pt(8.4); r.font.color.rgb = GR
        pc.paragraph_format.space_after = Pt(8)


# ===== derivados
val = RES["validos"]; abe = RES["abelardo"]; cep = RES["cepeda"]; ama = RES["amarillo"]
blanco = RES["blanco"]; censo = RES["censo"]; votantes = RES["votantes"]
deficit = cep - abe
abst = round(100 * (1 - votantes / censo), 1)
amd = RES["amarillo_detalle"]
pro_abe = RES["pro_abe"]; disputa = RES["disputa"]; pro_cep = RES["pro_cep"]
verdes_hostil = sum(x["abelardo"] for x in BAR if x["tier"] in ("Bastión Cepeda", "Perdido cosechable"))
tier_agg = {}
for t in ["Disputado", "Perdido cosechable", "Bastión Cepeda", "Bastión Abelardo"]:
    sub = [x for x in BAR if x["tier"] == t]
    tier_agg[t] = dict(n=len(sub), verde=sum(x["abelardo"] for x in sub),
                       amarillo=sum(x["amarillo"] for x in sub),
                       cosecha=sum(x["cosecha_eff"] for x in sub),
                       validos=sum(x["validos"] for x in sub))

# ===== logo en el encabezado — esquina superior derecha, todas las páginas
_hdr = d.sections[0].header; _hdr.is_linked_to_previous = False
_hp = _hdr.paragraphs[0]; _hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
_hp.paragraph_format.space_after = Pt(0)
_hp.add_run().add_picture(f"{OUT}/logo_ricardoruiz.png", width=Inches(1.45))

# =========================================================== PORTADA
p = d.add_paragraph(); p.paragraph_format.space_before = Pt(60)
r = p.add_run("INTELIGENCIA ELECTORAL · SEGUNDA VUELTA 2026"); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ABE
p = d.add_paragraph(); r = p.add_run("Cartagena, barrio por barrio:"); r.bold = True; r.font.size = Pt(27); r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(0)
p = d.add_paragraph(); r = p.add_run("dónde concentrar el esfuerzo de campaña"); r.bold = True; r.font.size = Pt(27); r.font.color.rgb = NAVY
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(10)
p = d.add_paragraph(); r = p.add_run(
    "El semáforo político-electoral (rojos, amarillos y verdes) llevado a los 213 barrios de la ciudad, "
    "con el preconteo oficial de la primera vuelta. Dónde se cosechan amarillos, dónde se cuida a los verdes "
    "y dónde no vale la pena gastar la última semana.")
r.font.size = Pt(12.5); r.font.color.rgb = INK
p = d.add_paragraph(); p.paragraph_format.space_before = Pt(26)
r = p.add_run("Preparado para  "); r.font.size = Pt(11); r.font.color.rgb = GR
r = p.add_run("Tatiana Villarreal"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
p = d.add_paragraph(); r = p.add_run("Documento de trabajo · uso interno · confidencial · 8 de junio de 2026")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GR
p = d.add_paragraph(); r = p.add_run(
    "Base: preconteo oficial Registraduría de la 1ª vuelta (31-may-2026) a nivel de mesa · "
    "137 puestos de votación de Cartagena georreferenciados y cruzados con la cartografía barrial oficial. "
    "Cifras de preconteo, sujetas a escrutinio definitivo.")
r.italic = True; r.font.size = Pt(8.8); r.font.color.rgb = GR

# índice
h("Contenido", size=12, color=NAVY, before=22)
for t in [
    "Resumen ejecutivo (una página)",
    "1 · Diagnóstico y recomendaciones (en detalle)",
    "2 · Qué hicimos y con qué datos (metodología y lectura honesta de límites)",
    "3 · El panorama: cómo votó Cartagena en primera vuelta",
    "4 · El semáforo barrio por barrio: rojos, amarillos y verdes",
    "5 · Dónde concentrar el esfuerzo: los barrios prioritarios",
    "6 · Qué hacer en la última semana",
    "Anexo · Tabla completa de los 82 barrios con dato directo",
]:
    pp = d.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
    r = pp.add_run("· " + t); r.font.size = Pt(10.5); r.font.color.rgb = INK

# =========================================================== RESUMEN EJECUTIVO (1 página)
def bullet_sm(txt):
    p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in re.split(r"(\*\*.*?\*\*)", txt):
        r = p.add_run(seg[2:-2] if seg.startswith("**") else seg)
        r.font.size = Pt(9.8); r.bold = seg.startswith("**"); r.font.color.rgb = INK


p = d.add_paragraph(); p.paragraph_format.page_break_before = True; p.paragraph_format.space_after = Pt(2)
r = p.add_run("PARA DECIDIR EN UNA PÁGINA"); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GR
p2 = d.add_paragraph(); r = p2.add_run("Resumen ejecutivo"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
p2.paragraph_format.space_after = Pt(6)
body(f"**Cepeda ganó Cartagena en primera vuelta con {RES['cep_pct']}% frente al {RES['abe_pct']}% de "
     f"Abelardo** — unos {NF(deficit)} votos y {abs(RES['margen']):.0f} puntos de diferencia. La ciudad no se "
     f"voltea a una semana. La jugada no es ganar la plaza, sino **gastar bien los últimos siete días**: "
     f"cuidar a los verdes que ya votaron por Abelardo y cosechar el amarillo afín, sin perder tiempo con el "
     f"votante de Cepeda. Tu estrategia —enfocar donde Abelardo fue débil, buscando amarillos y verdes y no al "
     f"petrista— es la correcta; aquí está, barrio por barrio, dónde está cada color.", after=7)
tbl(["Semáforo", "Qué es", "Votos (1ª vuelta)", "Qué hacer"],
    [["Verde", "Ya votaron por Abelardo", f"{NF(abe)} ({RES['abe_pct']}%)", "Movilizar y cuidar el voto"],
     ["Rojo", "Base de Cepeda (petristas)", f"{NF(cep)} ({RES['cep_pct']}%)", "No persuadir a una semana"],
     ["Amarillo", "Votaron por otros candidatos", f"{NF(ama)} ({RES['ama_pct']}%)", "Persuadir (la mayoría es afín)"]],
    widths=[1.0, 2.0, 1.5, 2.0], fs=8.6)
h("Tres hallazgos", size=11.5, color=NAVY, before=10, after=3)
bullet_sm(f"**La bolsa amarilla es chica pero rema a favor:** de {NF(ama)} votos disponibles, ~{NF(pro_abe)} "
          f"son de derecha (Paloma, Lizcano, Botero) afines a Abelardo y ~{NF(disputa)} de Fajardo en disputa; "
          f"apenas {NF(pro_cep)} se inclinan a Cepeda.")
bullet_sm(f"**{NF(verdes_hostil)} verdes —el {round(100*verdes_hostil/abe)}% de la base de Abelardo— viven en "
          f"barrios que ganó Cepeda**, sin estructura que los cuide. Movilizarlos es la tarea de mayor retorno.")
bullet_sm("**El voto se ordena por estrato:** Abelardo arrasa en Bocagrande, Manga, Centro y Crespo, y pierde "
          "fuerte en los populares (Nelson Mandela, El Pozón, Olaya). El campo de batalla son los barrios "
          "medios de la Localidad Industrial.")
h("Qué hacer esta semana", size=11.5, color=NAVY, before=8, after=3)
bullet_sm("**Cuidar a los verdes primero:** turnout y testigos para los ~89.000 verdes regados en territorio "
          "de Cepeda. Si no vuelven a votar, la pérdida es mayor que cualquier amarillo que se gane.")
bullet_sm("**Cosechar amarillos donde son afines:** barrios medios de la Localidad Industrial (Ternera, La "
          "Carolina, El Socorro, El Country, Blas de Lezo). Mensaje de cierre y contraste, no pedagogía.")
bullet_sm("**No persuadir petristas:** en Nelson Mandela, El Pozón u Olaya solo entra el operativo de cuidar "
          "a los verdes que ya existen, no el convencimiento al votante de Cepeda.")
bullet_sm("**La unidad de despliegue es el barrio, no la localidad:** dentro de cada localidad conviven "
          "bastiones de Cepeda y barrios disputados (ver mapa operativo, sección 5).")
bullet_sm("**Cartagena es defensa, no ofensiva:** el objetivo aquí es minimizar el déficit neto para la "
          "cuenta nacional. Eso libera recursos para donde sí se decide la elección.")

# =========================================================== 1 DIAGNÓSTICO
sech("1", "Diagnóstico y recomendaciones")
body(f"**Cartagena fue de Cepeda en primera vuelta, y por mucho.** Sobre {NF(val)} votos válidos en la "
     f"ciudad, Cepeda obtuvo **{RES['cep_pct']}% ({NF(cep)})** y Abelardo **{RES['abe_pct']}% ({NF(abe)})**: "
     f"una diferencia de **{abs(RES['margen']):.0f} puntos** y cerca de **{NF(deficit)} votos**. Esa es la "
     f"realidad de partida y conviene decirla sin rodeos: Cartagena no es un terreno que se voltee a una "
     f"semana hablándole al votante de Cepeda. La pregunta correcta no es «cómo ganar la ciudad», sino "
     f"**dónde rinde cada peso y cada caminante en los próximos siete días**.")
body(f"**La estrategia que planteas es la correcta, y los datos la respaldan.** Enfocarse en los lugares "
     f"donde Abelardo no fue fuerte —pero buscando ahí a los **amarillos** (votaron por otros y están en "
     f"disputa) y a los **verdes** (ya votaron por Abelardo y hay que asegurarlos), no al **petrista**— es "
     f"exactamente lo que la geografía del voto aconseja. Este documento ubica, barrio por barrio, dónde "
     f"está cada color.")
body(f"**El semáforo de Cartagena, en cifras:** de los {NF(val)} válidos, **{NF(abe)} son verdes** "
     f"(base de Abelardo, hay que cuidarlos), **{NF(cep)} son rojos** (base de Cepeda, no se persuaden a "
     f"esta altura) y **{NF(ama)} son amarillos** ({RES['ama_pct']}%: votaron por Fajardo, Paloma, Lizcano y "
     f"los demás, y siguen en juego para la segunda vuelta). Más {NF(blanco)} votos en blanco.")
body(f"**Hallazgo 1 — la bolsa amarilla es chica pero mayoritariamente afín.** Solo el {RES['ama_pct']}% de "
     f"la ciudad quedó disponible, pero de esos {NF(ama)} votos, **~{NF(pro_abe)} son de derecha o "
     f"centro-derecha** (Paloma, Lizcano, Botero) que transfieren con naturalidad a Abelardo frente a "
     f"Cepeda; **~{NF(disputa)} son de Fajardo** (el centro en disputa real); y apenas **{NF(pro_cep)} "
     f"se inclinan a Cepeda**. La cosecha eficiente existe: está en consolidar el bloque de derecha y "
     f"pelear el voto de Fajardo.")
body(f"**Hallazgo 2 — el problema no es solo persuadir, es cuidar lo propio.** De los {NF(abe)} verdes de "
     f"Abelardo, **{NF(verdes_hostil)} ({round(100*verdes_hostil/abe)}%) viven en barrios que ganó "
     f"Cepeda**. Son justamente los «verdes sin quién los cuide» que mencionas: gente que ya votó por "
     f"Abelardo pero rodeada de territorio adverso, sin testigos ni estructura que les garantice volver a "
     f"las urnas. Movilizarlos y protegerles el voto es la tarea de mayor retorno en la ciudad.")
body(f"**Hallazgo 3 — el mapa se ordena por estrato.** Abelardo arrasa en los barrios de estrato alto "
     f"(Bocagrande 79%, Castillogrande, Manga 72%, Centro 58%, Crespo 53%) y pierde fuerte en los populares "
     f"(Nelson Mandela 22%, El Pozón 26%, Olaya, San Fernando). El campo de batalla real son los **barrios "
     f"medios de la Localidad Industrial** (Ternera, La Carolina, El Socorro, El Country, Blas de Lezo, "
     f"Las Gaviotas): ahí Abelardo perdió, pero por poco y con base grande.")
h("Las cinco recomendaciones", size=12.5, color=NAVY, before=10)
bullet("**Cuida a los verdes primero.** La prioridad número uno es turnout y testigos para los ~89.000 "
       "verdes que viven en territorio de Cepeda. Si no vuelven a votar, la pérdida es mayor que cualquier "
       "amarillo que se gane.")
bullet("**Cosecha amarillos donde son afines y están concentrados:** los barrios medios de la Localidad "
       "Industrial. Mensaje de cierre anti-continuidad / cambio, dirigido a quien votó Paloma, Lizcano o "
       "Fajardo — no pedagogía de cero.")
bullet("**No gastes la última semana persuadiendo petristas.** En Nelson Mandela, El Pozón u Olaya el "
       "esfuerzo de persuasión no rinde; ahí solo entra el operativo de cuidar a los verdes que existen.")
bullet("**Prioriza por barrio, no por localidad.** Dentro de una misma localidad conviven bastiones de "
       "Cepeda y barrios disputados. La unidad de despliegue es el barrio (sección 5).")
bullet("**Cartagena es defensa, no ofensiva.** Aun ganando todo el amarillo, la ciudad seguiría siendo de "
       "Cepeda; el objetivo aquí es **minimizar el déficit neto** para la cuenta nacional, no dar vuelta a "
       "la plaza. Eso libera recursos para donde sí se decide.")

# =========================================================== 2 METODOLOGIA
sech("2", "Qué hicimos y con qué datos")
body("**La fuente es el dato oficial, no una encuesta.** Trabajamos sobre el preconteo de la Registraduría "
     "de la primera vuelta (snapshot final del 31 de mayo, el mismo que cuadra con el resultado publicado), "
     "a nivel de **mesa de votación**. Para Cartagena son **137 puestos** (departamento 05, municipio 001), "
     "con el voto de cada candidato en cada mesa.")
body("**Cómo bajamos el dato al barrio.** Cada puesto de votación está georreferenciado (coordenadas del "
     "archivo oficial de puestos de la Registraduría). Tomamos esas coordenadas y las cruzamos, por "
     "ubicación, contra los **213 polígonos barriales oficiales** de Cartagena (cartografía catastral / "
     "IDECA). Así, los votos de cada puesto se suman al barrio que físicamente lo contiene. "
     f"**{RES['n_barrios_con_dato']} de los 213 barrios tienen puesto de votación propio** y por tanto dato "
     "directo. Para que los mapas se lean completos, los barrios sin puesto propio —en su mayoría rural, "
     "insular o sin mesa— se pintan con la tendencia de su **barrio vecino más cercano**, en tono atenuado "
     "para que se distingan a simple vista del dato directo.")
body("**El semáforo, definido con precisión.** Sobre el voto de primera vuelta de cada barrio:")
bullet("**Verde (base de Abelardo)** = votos que sacó Abelardo. Es la base a proteger y movilizar.")
bullet("**Rojo (base de Cepeda)** = votos que sacó Cepeda. A una semana, no es público persuadible.")
bullet("**Amarillo (disponible)** = todo lo que votó por los demás candidatos (Fajardo, Paloma, Lizcano, "
       "Botero, etc.). Es el universo que sigue en juego para la segunda vuelta. Dentro del amarillo "
       "distinguimos lo **afín a Abelardo** (derecha: Paloma, Lizcano, Botero), lo de **disputa** (centro: "
       "Fajardo) y lo **afín a Cepeda** (izquierda menor: Roy, Caicedo, Murillo).")
body("**Y los límites, con honestidad.** Esto es un análisis del comportamiento **agregado** de cada barrio, "
     "no una lista de personas: el voto es secreto y nadie viene marcado de un color. \"Rojo\", \"amarillo\" "
     "y \"verde\" describen cómo votó el conjunto del barrio en primera vuelta, no a individuos. Que el "
     "amarillo afín «transfiera» a Abelardo en segunda vuelta es un **supuesto razonable**, no un dato: la "
     "gente puede quedarse en casa o cambiar de opinión. Por eso las cifras de amarillo se leen como "
     "**tamaño del universo a trabajar**, no como votos garantizados. Donde un barrio no tiene puesto propio "
     "no inventamos su resultado: queda sin dato. Y todo son cifras de preconteo, sujetas al escrutinio "
     "definitivo.")

# =========================================================== 3 PANORAMA
sech("3", "El panorama: cómo votó Cartagena")
body(f"En la ciudad votaron {NF(votantes)} personas de un potencial de {NF(censo)} ({abst}% de abstención). "
     f"El resultado de primera vuelta fue contundente a favor de Cepeda:")
tbl(["", "Votos", "% válidos"],
    [["Cepeda (rojo)", NF(cep), f"{RES['cep_pct']}%"],
     ["Abelardo (verde)", NF(abe), f"{RES['abe_pct']}%"],
     ["Otros candidatos (amarillo)", NF(ama), f"{RES['ama_pct']}%"],
     ["Voto en blanco", NF(blanco), f"{round(100*blanco/val,1)}%"],
     ["Total válidos", NF(val), "100%"]],
    widths=[2.6, 1.5, 1.2])
body("**Cepeda ganó la mayor parte del territorio de la ciudad.** El mapa de ganadores muestra el morado "
     "(Cepeda) cubriendo casi toda Cartagena, con islas azules (Abelardo) en los barrios de estrato alto: "
     "la península de Bocagrande–Castillogrande–El Laguito, Manga, Centro y Crespo.")
img("m1_ganador.png", 6.4, "Mapa 1 · Quién ganó cada barrio en primera vuelta. Azul = Abelardo · morado = Cepeda.")
h("Las tres localidades", size=12.5, color=NAVY)
body("Cartagena se organiza en tres localidades. Su comportamiento fue muy distinto, y eso ya marca dónde "
     "hay algo que pelear:")
tbl(["Localidad", "Barrios", "Válidos", "Abe %", "Cep %", "Margen", "Verde (Abe)", "Amarillo"],
    [[l["loc_nombre"], l["n_barrios"], NF(l["validos"]), f"{l['abe_pct']}%", f"{l['cep_pct']}%",
      f"{l['margen']:+.0f}", NF(l["abelardo"]), NF(l["amarillo"])] for l in LOC],
    widths=[1.9, 0.6, 0.85, 0.65, 0.65, 0.7, 0.95, 0.85], fs=8.2)
body("**Histórica y del Caribe Norte** es la más competida (Abelardo a solo 7 puntos): concentra sus "
     "bastiones (Bocagrande, Manga, Centro, Crespo) junto a barrios disputados como Pie del Cerro. "
     "**Industrial y de la Bahía** es donde más se juega la cosecha: Abelardo pierde por ~29 puntos en el "
     "agregado, pero contiene los grandes barrios medios donde la elección está abierta (Ternera, La "
     "Carolina, El Socorro). **De la Virgen y Turística** es el terreno más adverso (−36): los grandes "
     "barrios populares —El Pozón, Nelson Mandela, Olaya— donde Cepeda fue dominante.")
img("m3_margen.png", 6.0, "Mapa 2 · Margen Abelardo − Cepeda por barrio. Azul intenso = bastión Abelardo · "
    "morado intenso = bastión Cepeda · tonos claros = barrios disputados (cerca del empate).")

# =========================================================== 4 SEMAFORO
sech("4", "El semáforo barrio por barrio")
body("El mapa de fuerza de Abelardo confirma el patrón por estrato: el azul oscuro (donde más votó) se "
     "concentra en la franja de estrato alto, y se aclara hacia los barrios populares del sur y la ciénaga.")
img("m2_abelardo_pct.png", 6.0, "Mapa 3 · Fuerza de Abelardo (% sobre válidos). Más azul = más votó Abelardo.")
h("De qué está hecha la bolsa amarilla", size=12.5, color=NAVY)
body("La bolsa amarilla de la ciudad —los " + NF(ama) + " votos en disputa— se reparte así. Importa porque "
     "dice **a quién hay que hablarle**:")
tbl(["Candidato (1ª vuelta)", "Votos", "% de la bolsa", "Afinidad en 2ª vuelta"],
    [["Sergio Fajardo", NF(amd["fajardo"]), f"{round(100*amd['fajardo']/ama)}%", "Centro — en disputa"],
     ["Paloma Valencia", NF(amd["paloma"]), f"{round(100*amd['paloma']/ama)}%", "Derecha — afín a Abelardo"],
     ["Mauricio Lizcano", NF(amd["lizcano"]), f"{round(100*amd['lizcano']/ama)}%", "Centro-derecha — afín"],
     ["Santiago Botero", NF(amd["botero"]), f"{round(100*amd['botero']/ama)}%", "Derecha — afín a Abelardo"],
     ["Resto (8 candidatos)", NF(ama - amd["fajardo"] - amd["paloma"] - amd["lizcano"] - amd["botero"]),
      f"{round(100*(ama-amd['fajardo']-amd['paloma']-amd['lizcano']-amd['botero'])/ama)}%", "Mezcla"]],
    widths=[1.9, 1.0, 1.1, 2.3], fs=8.4)
body(f"**Leído por bloques: ~{NF(pro_abe)} votos amarillos son de derecha/centro-derecha** (afines a "
     f"Abelardo), **~{NF(disputa)} son de Fajardo** (la pelea real del centro) y solo **{NF(pro_cep)} se "
     f"inclinan a Cepeda**. Es decir: la mayoría del amarillo rema a favor. La cosecha eficiente es "
     f"**blindar el voto de derecha que no es de Abelardo y disputarle a Cepeda el voto de Fajardo**.")
body("El mapa siguiente muestra dónde está concentrada esa bolsa: los barrios medios (no los bastiones de "
     "ningún lado) son los que más amarillo tienen en proporción.")
img("m4_amarillo.png", 6.0, "Mapa 4 · Bolsa amarilla por barrio (% que votó por otros candidatos).")

# =========================================================== 5 PRIORIDAD
sech("5", "Dónde concentrar el esfuerzo")
body("Clasificamos cada barrio en cuatro categorías operativas, combinando qué tan cerca quedó la elección "
     "y cuánto hay para cosechar y para cuidar. Este es el mapa para repartir caminantes, testigos y "
     "recursos en la última semana.")
img("m5_prioridad.png", 6.5, "Mapa 5 · Mapa operativo de segunda vuelta. Los números corresponden al "
    "ranking de la tabla de abajo (barrios prioritarios urbanos).")
tbl(["Categoría", "Barrios", "Verde (a cuidar)", "Amarillo", "Qué hacer"],
    [["Disputado", tier_agg["Disputado"]["n"], NF(tier_agg["Disputado"]["verde"]),
      NF(tier_agg["Disputado"]["amarillo"]), "Máxima prioridad: cosechar y cuidar"],
     ["Perdido cosechable", tier_agg["Perdido cosechable"]["n"], NF(tier_agg["Perdido cosechable"]["verde"]),
      NF(tier_agg["Perdido cosechable"]["amarillo"]), "Cosecha selectiva + turnout de verdes"],
     ["Bastión Cepeda", tier_agg["Bastión Cepeda"]["n"], NF(tier_agg["Bastión Cepeda"]["verde"]),
      NF(tier_agg["Bastión Cepeda"]["amarillo"]), "Recursos mínimos: solo testigos"],
     ["Bastión Abelardo", tier_agg["Bastión Abelardo"]["n"], NF(tier_agg["Bastión Abelardo"]["verde"]),
      NF(tier_agg["Bastión Abelardo"]["amarillo"]), "Cuidar el voto; no persuadir"]],
    widths=[1.6, 0.6, 1.2, 0.85, 2.3], fs=8.3)
body("**Cómo leerlo.** Los **disputados** (oro en el mapa) son donde la elección quedó cerca: cada verde "
     "movilizado y cada amarillo convencido cuenta doble. Los **perdidos cosechables** (naranja) son los "
     "grandes barrios donde Abelardo perdió pero tiene miles de verdes y una bolsa amarilla apreciable: ahí "
     "el oro está en el turnout de los propios. Los **bastiones de Cepeda** (morado) concentran 37 barrios "
     "con poco que cosechar: solo se entra a proteger a los pocos verdes con testigos. Los **bastiones de "
     "Abelardo** (verde) ya están ganados: se cuidan, no se persuaden.")
h("Los barrios prioritarios, uno por uno", size=12.5, color=NAVY)
body("Tabla de los barrios urbanos a trabajar, en orden de prioridad (el número coincide con el del Mapa 5). "
     "**Amarillo** = votos por cosechar; **Verde** = votos de Abelardo por cuidar.")
LOC_SHORT = {"Histórica y del Caribe Norte": "Histórica", "De la Virgen y Turística": "La Virgen",
             "Industrial y de la Bahía": "Industrial"}
_LOWER = {"de", "del", "la", "las", "los", "y", "el", "san"}  # 'san' queda; ajustamos abajo


def titlecase(s):
    out = []
    for i, w in enumerate(s.lower().split()):
        out.append(w if (w in _LOWER and i > 0 and w != "san") else w.capitalize())
    return " ".join(out)


def fmt_mrg(x):
    r = round(x)
    return "0" if r == 0 else f"{r:+d}"


rows = []
for k in sorted(LAB, key=lambda x: int(x)):
    v = LAB[k]
    rows.append([k, titlecase(v["barrio"]), LOC_SHORT.get(v["loc"], v["loc"][:11]),
                 f"{v['abe']}%", f"{v['cep']}%", fmt_mrg(v["margen"]),
                 NF(v["amarillo"]), NF(v["verde"]), v["tier"].split(" ")[0]])
tbl(["#", "Barrio", "Localidad", "Abe", "Cep", "Mrg", "Amar.", "Verde", "Categoría"],
    rows, widths=[0.3, 1.75, 1.0, 0.55, 0.55, 0.5, 0.7, 0.7, 1.05], fs=7.9)
note("Nota: además de estos barrios urbanos, el operativo de cuidado de verdes debe cubrir los grandes "
     "bastiones de Cepeda donde Abelardo igual tiene base (El Pozón ~3.800 verdes, San José de los Campanos, "
     "San Fernando, Olaya): no se persuade ahí, pero esos votos verdes hay que protegerlos. La tabla "
     "completa de los 82 barrios está en el Anexo (CSV).")

# =========================================================== 6 ULTIMA SEMANA
sech("6", "Qué hacer en la última semana")
h("Verdes — cuidar el voto (la prioridad #1)", size=12, color=ABE)
body(f"Abelardo tiene {NF(abe)} votos ya emitidos en Cartagena, y **{NF(verdes_hostil)} de ellos están en "
     f"barrios que ganó Cepeda**, sin estructura propia que los acompañe. A una semana, el mayor retorno no "
     f"es convencer a nadie nuevo: es **garantizar que los que ya votaron por Abelardo vuelvan a las urnas** "
     f"y que su voto se cuente. Eso significa testigos electorales en los puestos de los bastiones de Cepeda "
     f"(donde hay verdes sueltos), recordatorio puerta a puerta a la base identificada, y logística de "
     f"transporte el día de la elección en los barrios perdidos-cosechables (El Socorro, Blas de Lezo, "
     f"El Country, Almirante Colón).")
h("Amarillos — persuadir donde rinde", size=12, color=ABE)
body("La cosecha está en la **Localidad Industrial** y en los disputados de la Histórica: Ternera, La "
     "Carolina, Las Gaviotas, Pie del Cerro, El Bosque, San Isidro. Ahí está la mayor concentración de "
     "amarillo afín, junto a una base verde grande. El mensaje es de **cierre y contraste** (continuidad vs. "
     "cambio, orden, lo que está en juego), dirigido a quien votó por Paloma, Lizcano o Fajardo. No es "
     "momento de pedagogía ni de construir identidad desde cero: a siete días, eso no ocurre. Es activar a "
     "quien ya está cerca.")
h("Rojos — no malgastar", size=12, color=ABE)
body("En Nelson Mandela, El Pozón, Olaya, San Fernando, Torices y los demás bastiones de Cepeda, la "
     "persuasión no rinde y consume recursos que hacen falta en otro lado. La única operación que justifica "
     "entrar ahí es **proteger y movilizar a los verdes que existen** — nada de puerta a puerta de "
     "convencimiento al votante de Cepeda.")
h("La síntesis", size=12, color=ABE)
body("**Cuidar lo verde, cosechar lo amarillo afín, no tocar lo rojo.** Concentrar el operativo en los "
     "barrios medios de la Localidad Industrial y los disputados de la Histórica para persuasión; y un "
     "operativo de testigos y turnout que cubra a los ~89.000 verdes regados en territorio adverso. "
     "Cartagena no se gana, pero se puede perder por menos — y ese margen, sumado al resto del país, importa.")

# =========================================================== ANEXO
sech("Anexo", "Tabla completa de los 82 barrios con dato directo")
body("El archivo **Anexo_Cartagena_barrios.csv** (adjunto) trae los 82 barrios con puesto de votación propio, "
     "ordenados por prioridad, con todas las columnas: ranking, barrio, localidad, categoría, ganador, % de "
     "Abelardo y de Cepeda, margen, % amarillo, votos de Abelardo (verde), de Cepeda (rojo), amarillo total, "
     "amarillo afín a la derecha, voto en blanco, válidos, censo, número de puestos y la acción recomendada. "
     "Se abre directamente en Excel.")
note("Metodología y fuentes — Preconteo oficial de la Registraduría Nacional, primera vuelta presidencial "
     "del 31 de mayo de 2026, a nivel de mesa (snapshot final, el que cuadra con el resultado publicado). "
     "Georreferenciación de los 137 puestos de Cartagena a partir del archivo oficial de puestos de votación; "
     "asignación a barrio por cruce de ubicación (punto en polígono / vecino más cercano) contra los 213 "
     "polígonos barriales oficiales de la ciudad. Porcentajes sobre votos válidos. El \"semáforo\" "
     "(rojo/amarillo/verde) describe el comportamiento agregado de primera vuelta de cada barrio, no a "
     "personas; la transferencia del amarillo a la segunda vuelta es un supuesto de trabajo, no un dato. "
     "Cifras de preconteo, sujetas al escrutinio definitivo. Análisis: Ricardo Ruiz · inteligencia electoral.")

DOCX = f"{OUT}/Cartagena_Barrios_2V_Tatiana_Villarreal.docx"
d.save(DOCX); print("✓ Word guardado:", DOCX)


def embed_inter(docx_path):
    import zipfile, uuid, re, shutil
    FDIR = "tools/pacto-1v-2026/fonts"
    FONTS = [("embedRegular", f"{FDIR}/Inter-Regular.ttf"),
             ("embedBold", f"{FDIR}/Inter-Bold.ttf"),
             ("embedItalic", f"{FDIR}/Inter-Italic.ttf")]
    def obf(data, guid):
        key = bytes.fromhex(guid.strip("{}").replace("-", ""))[::-1]
        b = bytearray(data)
        for i in range(32): b[i] ^= key[i % 16]
        return bytes(b)
    zin = zipfile.ZipFile(docx_path, "r"); items = {n: zin.read(n) for n in zin.namelist()}; zin.close()
    embeds = []
    for idx, (kind, path) in enumerate(FONTS, 1):
        guid = "{" + str(uuid.uuid4()).upper() + "}"
        items[f"word/fonts/font{idx}.odttf"] = obf(open(path, "rb").read(), guid)
        embeds.append((kind, f"rIdFont{idx}", f"fonts/font{idx}.odttf", guid))
    childs = "".join(f'<w:{k} r:id="{rid}" w:fontKey="{g}"/>' for k, rid, t, g in embeds)
    node = ('<w:font w:name="Inter"><w:charset w:val="00"/><w:family w:val="swiss"/>'
            '<w:pitch w:val="variable"/>' + childs + "</w:font>")
    items["word/fontTable.xml"] = items["word/fontTable.xml"].decode("utf-8").replace("</w:fonts>", node + "</w:fonts>").encode("utf-8")
    rels = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            + "".join(f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="{t}"/>' for k, rid, t, g in embeds)
            + "</Relationships>")
    items["word/_rels/fontTable.xml.rels"] = rels.encode("utf-8")
    ct = items["[Content_Types].xml"].decode("utf-8")
    if 'Extension="odttf"' not in ct:
        ct = ct.replace("</Types>", '<Default Extension="odttf" ContentType="application/vnd.openxmlformats-officedocument.obfuscatedFont"/></Types>')
    items["[Content_Types].xml"] = ct.encode("utf-8")
    st = items["word/settings.xml"].decode("utf-8")
    if "<w:embedTrueTypeFonts/>" not in st:
        st = st.replace("<w:proofState", "<w:embedTrueTypeFonts/><w:proofState", 1) if "<w:proofState" in st \
            else re.sub(r"(<w:settings[^>]*>)", r"\1<w:embedTrueTypeFonts/>", st, count=1)
    items["word/settings.xml"] = st.encode("utf-8")
    tmp = docx_path + ".tmp"; zo = zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED)
    for n, data in items.items(): zo.writestr(n, data)
    zo.close(); shutil.move(tmp, docx_path)
    return embeds


_emb = embed_inter(DOCX); print(f"✓ Inter incrustada ({len(_emb)} variantes)")
