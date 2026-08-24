#!/usr/bin/env python3
"""
RENADIA · Propuesta de estrategias de relacionamiento (mesas temáticas y
diálogos bilaterales), versión 2, en formato .docx conforme a las Pautas para
la elaboración, revisión y entrega de documentos técnicos de la DENDD.

Diferencias frente a la versión 1 (el .docx que venía de Word):
  · idioma del documento configurado en español (es-CO) en estilos, runs,
    settings y propiedades del archivo; la versión anterior traía en-US y por
    eso Word marcaba todo el texto como error ortográfico;
  · formato DNP completo (portada institucional, control documental, siglas,
    Arial 11, interlineado 1,5, márgenes de 2,54 cm, numeración decimal hasta
    el tercer nivel, encabezado con logo y "Página X de Y", tablas numeradas
    con título arriba y fuente abajo). SIN tabla de contenido, por decisión
    del solicitante;
  · la primera mesa temática pasa a ser la de creación, uso y administración
    de tableros en entidades públicas nacionales y territoriales, desarrollada
    en un capítulo propio;
  · se integran al cuerpo las dos notas sueltas que traía el borrador (el rol
    de mentor de la comunidad y la invitación a compartir contactos) y se
    agrega un capítulo de decisiones pendientes para la coordinación.

Reusa la infraestructura de formato de tools/dnp-entregas/build_entrega2.py.

Uso:
    python3 build_mesas.py [--salida RUTA.docx] [--pdf]
"""

import argparse
import os
import subprocess
import sys

RAIZ = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, os.path.join(RAIZ, "tools", "dnp-entregas"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import build_entrega2 as base
from build_entrega2 import (FUENTE, LOGO, SOFFICE, configurar_estilos,
                            encabezado_paginado, margenes, parrafo, salto,
                            titulo, vineta)

def tabla(doc, *args, **kwargs):
    """base.tabla fija el ancho en cada celda, que es lo que lee Word, pero no
    toca la cuadrícula (w:tblGrid), que es lo que lee LibreOffice y lo que se
    usa para la vista previa en PDF. Sin esto las columnas salen todas del
    mismo ancho al convertir."""
    t = base.tabla(doc, *args, **kwargs)
    anchos = kwargs.get("anchos") or (args[4] if len(args) > 4 else None)
    if anchos:
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for col, ancho in zip(grid.findall(qn("w:gridCol")), anchos):
                col.set(qn("w:w"), str(int(Cm(ancho).twips)))
    return t


SALIDA_DEFECTO = os.path.join(
    RAIZ, "Bases de datos", "DNP", "mesas-tematicas",
    "RENADIA-Mesas-Tematicas-y-Dialogos-Bilaterales-v2.docx",
)

IDIOMA = "es-CO"


# --------------------------------------------------------------------------
# Idioma
# --------------------------------------------------------------------------

def _marcar_idioma(rPr, cod=IDIOMA):
    for previo in rPr.findall(qn("w:lang")):
        rPr.remove(previo)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), cod)
    lang.set(qn("w:eastAsia"), cod)
    lang.set(qn("w:bidi"), "ar-SA")
    rPr.append(lang)


def idioma_por_defecto(doc, cod=IDIOMA):
    """Fija el español en los valores por defecto y en todos los estilos.

    Word toma el idioma de tres lugares distintos —docDefaults, el estilo y el
    run— y basta con que uno quede en inglés para que el corrector marque el
    documento completo. Se cubren los tres."""
    estilos = doc.styles.element
    defecto = estilos.find(qn("w:docDefaults"))
    if defecto is None:
        defecto = OxmlElement("w:docDefaults")
        estilos.insert(0, defecto)
    rpr_def = defecto.find(qn("w:rPrDefault"))
    if rpr_def is None:
        rpr_def = OxmlElement("w:rPrDefault")
        defecto.insert(0, rpr_def)
    rpr = rpr_def.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_def.append(rpr)
    _marcar_idioma(rpr, cod)
    for estilo in doc.styles:
        el = estilo.element
        if el.tag == qn("w:style"):
            _marcar_idioma(el.get_or_add_rPr(), cod)
    # themeFontLang, que es de donde Word deduce el idioma de la interfaz
    settings = doc.settings.element
    for previo in settings.findall(qn("w:themeFontLang")):
        settings.remove(previo)
    tfl = OxmlElement("w:themeFontLang")
    tfl.set(qn("w:val"), cod)
    settings.append(tfl)


def idioma_en_runs(doc, cod=IDIOMA):
    """Marca cada run del cuerpo, encabezados y pies. Se llama al final."""
    partes = [doc.element]
    for sec in doc.sections:
        for contenedor in (sec.header, sec.footer, sec.first_page_header,
                           sec.first_page_footer, sec.even_page_header,
                           sec.even_page_footer):
            try:
                partes.append(contenedor._element)
            except Exception:
                pass
    for parte in partes:
        for run in parte.iter(qn("w:r")):
            rPr = run.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                run.insert(0, rPr)
            _marcar_idioma(rPr, cod)


# --------------------------------------------------------------------------
# Portada propia (documento interno, no entrega contractual)
# --------------------------------------------------------------------------

def portada(doc):
    if os.path.exists(LOGO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(18)
        p.add_run().add_picture(LOGO, width=Cm(4.6))
    C = WD_ALIGN_PARAGRAPH.CENTER
    parrafo(doc, "Dirección de Economía Naranja y Desarrollo Digital", tam=12,
            align=C, space_after=6)
    parrafo(doc, "Red Nacional de Analítica de Datos e Inteligencia Artificial "
                 "(RENADIA)", tam=12, align=C, space_after=30)
    parrafo(doc, "Estrategias para un relacionamiento cercano con los miembros "
                 "de la red", tam=18, negrita=True, align=C, space_after=10)
    parrafo(doc, "Mesas temáticas y diálogos bilaterales, con mesa piloto de "
                 "tableros para la decisión pública", tam=13, align=C,
            space_after=30)
    parrafo(doc, "Versión 2 · Documento interno de trabajo", tam=11,
            negrita=True, align=C, space_after=30)
    parrafo(doc, "Elaborado por:", tam=11, negrita=True, align=C, space_after=2)
    parrafo(doc, "Ricardo Esteban Ruiz Castro", tam=11, align=C, space_after=2)
    parrafo(doc, "Centro de Pensamiento de Desarrollo Digital", tam=11,
            align=C, space_after=18)
    parrafo(doc, "Dirigido a:", tam=11, negrita=True, align=C, space_after=2)
    parrafo(doc, "Vanessa Monroy", tam=11, align=C, space_after=2)
    parrafo(doc, "Coordinadora de RENADIA", tam=11, align=C, space_after=30)
    parrafo(doc, "Departamento Nacional de Planeación", tam=11, negrita=True,
            align=C, space_after=2)
    parrafo(doc, "Bogotá, D. C.", tam=11, align=C, space_after=2)
    parrafo(doc, "Agosto de 2026", tam=11, align=C, space_after=0)


# --------------------------------------------------------------------------
# Elementos preliminares
# --------------------------------------------------------------------------

def control_documental(doc):
    titulo(doc, "Página legal y control documental", 1)
    parrafo(doc,
            "Este documento es una propuesta interna de trabajo elaborada para "
            "la coordinación de la Red Nacional de Analítica de Datos e "
            "Inteligencia Artificial (RENADIA), en el marco de la operación del "
            "Carril 1 de la red. No constituye un acto administrativo ni un "
            "compromiso institucional: describe alternativas operativas para "
            "que la coordinación decida cuáles adopta, con qué alcance y en qué "
            "momento. Las cifras de la convocatoria «Mundial 2026» que se citan "
            "provienen del informe de resultados de esa actividad y "
            "corresponden a las 33 sesiones registradas entre el 6 y el 9 de "
            "julio de 2026.")
    tabla(doc, "Control documental",
          ["Campo", "Contenido"],
          [["Título", "Estrategias para un relacionamiento cercano con los "
                      "miembros de la red"],
           ["Tipo documental", "Propuesta interna de trabajo"],
           ["Versión", "2"],
           ["Dependencia responsable", "Dirección de Economía Naranja y "
                                       "Desarrollo Digital"],
           ["Elabora", "Ricardo Esteban Ruiz Castro, Centro de Pensamiento de "
                       "Desarrollo Digital"],
           ["Dirigido a", "Vanessa Monroy, Coordinadora de RENADIA"],
           ["Fecha de la versión", "Agosto de 2026"],
           ["Cambios frente a la versión 1",
            "Se adopta el formato de documento técnico de la Dirección; se "
            "define la mesa de tableros como mesa piloto y se desarrolla en un "
            "capítulo propio; se incorpora el rol de mentor de la comunidad y "
            "el paso de referidos en el diálogo bilateral; se agrega el "
            "capítulo de decisiones pendientes."]],
          "elaboración propia.", anchos=[4.5, 11.5])


def siglas(doc):
    titulo(doc, "Siglas y abreviaturas", 1)
    tabla(doc, "Siglas y abreviaturas empleadas en el documento",
          ["Sigla", "Significado"],
          [["CONPES", "Consejo Nacional de Política Económica y Social"],
           ["DANE", "Departamento Administrativo Nacional de Estadística"],
           ["DDD", "Dirección de Desarrollo Digital"],
           ["DENDD", "Dirección de Economía Naranja y Desarrollo Digital"],
           ["DNP", "Departamento Nacional de Planeación"],
           ["IA", "Inteligencia artificial"],
           ["MIPG", "Modelo Integrado de Planeación y Gestión"],
           ["MinTIC", "Ministerio de Tecnologías de la Información y las "
                      "Comunicaciones"],
           ["RENADIA", "Red Nacional de Analítica de Datos e Inteligencia "
                       "Artificial"],
           ["TIC", "Tecnologías de la información y las comunicaciones"]],
          "elaboración propia.", anchos=[3.0, 13.0])


def presentacion(doc):
    titulo(doc, "Presentación", 1)
    parrafo(doc,
            "La coordinación de RENADIA solicitó el diseño de espacios de "
            "relacionamiento que complementen los webinars y los boletines, y "
            "que produzcan contacto recurrente y conversación de doble vía con "
            "los miembros de la red. Este documento responde a esa solicitud "
            "con dos estrategias operativas —mesas temáticas y diálogos "
            "bilaterales— formuladas para ejecutarse en Microsoft Teams con la "
            "capacidad instalada actual del equipo.")
    parrafo(doc,
            "Frente a la versión anterior, esta incorpora una definición que la "
            "coordinación pidió precisar: cuál es la primera mesa que abre el "
            "ciclo. La propuesta es que sea la mesa de creación, uso y "
            "administración de tableros en entidades públicas nacionales y "
            "territoriales, desarrollada en el capítulo 4 con su alcance "
            "temático, su agenda de sesiones, sus criterios de selección de "
            "casos, sus productos, sus riesgos y sus indicadores de "
            "seguimiento.")
    parrafo(doc,
            "El documento cierra con un capítulo de decisiones pendientes. Son "
            "seis definiciones que no dependen del diseño sino de la "
            "coordinación, y que conviene resolver antes de convocar la primera "
            "sesión. Para cada una se presentan las opciones y una "
            "recomendación argumentada.")


# --------------------------------------------------------------------------
# 1. Contexto
# --------------------------------------------------------------------------

def contexto(doc):
    titulo(doc, "1. Contexto y motivación", 1)
    parrafo(doc,
            "RENADIA opera hoy con dos vehículos principales de comunicación "
            "con sus miembros: los webinars sectoriales y temáticos, que ya se "
            "ejecutan de manera sostenida desde la Dirección de Desarrollo "
            "Digital, y los boletines, próximos a entrar en producción. Estos "
            "canales tienen una virtud clara —llegan a una audiencia amplia con "
            "contenidos estructurados de calidad— y han permitido que la red "
            "tenga presencia y voz en la agenda nacional de analítica de datos "
            "e inteligencia artificial.")
    parrafo(doc,
            "Sin embargo, ambos comparten una característica que limita su "
            "capacidad para construir comunidad: son canales unidireccionales y "
            "de baja frecuencia individual. Un miembro de una entidad nacional "
            "puede pasar dos o tres meses sin tener una interacción real con la "
            "red, lo cual erosiona progresivamente el sentido de pertenencia "
            "que sostiene una comunidad como esta. En redes orientadas a la "
            "cocreación —y RENADIA lo es por diseño, en línea con el CONPES "
            "4144 y el espíritu de gobernanza participativa— la frecuencia y la "
            "calidad del contacto son tan importantes como el alcance.")
    parrafo(doc,
            "La lectura no es solo conceptual. En la convocatoria «Mundial "
            "2026», que reunió 33 sesiones de participantes de entidades "
            "nacionales, territoriales, del sector privado y de la academia, la "
            "forma preferida de vincularse a la red fueron precisamente las "
            "mesas temáticas de trabajo: 22 de las respuestas, el 61 %, las "
            "escogieron por encima de los webinars (31 %) y de los diálogos uno "
            "a uno (8 %). Y los dos frenos más citados para avanzar en datos e "
            "inteligencia artificial fueron la falta de capacidades o talento "
            "(33 %) y la falta de espacios con quién intercambiar experiencias "
            "(28 %). Existe entonces una demanda expresada por los miembros, y "
            "no solamente una hipótesis de diseño.")
    parrafo(doc,
            "La coordinación de la red identificó esta brecha y solicitó el "
            "diseño de espacios complementarios que generen contacto "
            "recurrente, conversación de doble vía y vínculos horizontales "
            "entre los miembros, no solamente verticales hacia el DNP. El "
            "presente documento responde a esa solicitud con dos estrategias "
            "operativas, complementarias entre sí.")


# --------------------------------------------------------------------------
# 2. Marco estratégico
# --------------------------------------------------------------------------

def marco(doc):
    titulo(doc, "2. Marco estratégico de las dos propuestas", 1)
    parrafo(doc,
            "Las dos estrategias propuestas operan en niveles distintos pero "
            "complementarios del relacionamiento. La primera —mesas temáticas— "
            "construye comunidad horizontal entre los miembros a través de "
            "espacios grupales pequeños y recurrentes. La segunda —diálogos "
            "bilaterales— construye una relación vertical personalizada entre "
            "el DNP y cada entidad miembro. Juntas configuran un sistema de "
            "relacionamiento que cubre tanto la dimensión colectiva como la "
            "dimensión individual de la red.")
    tabla(doc, "Comparación de las dos estrategias de relacionamiento",
          ["Dimensión", "Mesas temáticas", "Diálogos bilaterales"],
          [["Tipo de relación", "Grupal y horizontal (entre pares)",
            "Bilateral y vertical (DNP–entidad)"],
           ["Frecuencia sugerida", "Mensual por mesa", "Trimestral por entidad"],
           ["Tamaño del grupo", "10 a 15 entidades por mesa",
            "Uno a uno (o equipos pequeños)"],
           ["Foco principal", "Aprendizaje entre pares en torno a un tema",
            "Acompañamiento a la trayectoria de cada entidad"],
           ["Plataforma", "Microsoft Teams", "Microsoft Teams"],
           ["Producto principal",
            "Casos, lecciones aprendidas e insumos para estándares",
            "Mapeo de necesidades y plan de articulación por entidad"],
           ["Riesgo si no se hace",
            "La red no genera vínculos entre miembros y depende del DNP para "
            "cada interacción",
            "La red no distingue trayectorias y trata igual a entidades con "
            "madurez muy distinta"]],
          "elaboración propia.", anchos=[3.6, 6.2, 6.2])
    parrafo(doc,
            "Ambas estrategias se articulan con la operación del Carril 1 de la "
            "red: alimentan los webinars con casos y ponentes identificados, "
            "nutren los boletines con contenidos surgidos de las "
            "conversaciones, y aportan insumos a la construcción de estándares "
            "y lineamientos técnicos en los temas prioritarios de RENADIA. "
            "Ninguna de las dos compite con los canales existentes ni los "
            "reemplaza: los abastece.")


# --------------------------------------------------------------------------
# 3. Mesas temáticas
# --------------------------------------------------------------------------

def mesas(doc):
    titulo(doc, "3. Estrategia 1: mesas temáticas RENADIA", 1)

    titulo(doc, "3.1. Qué son las mesas temáticas", 2)
    parrafo(doc,
            "Las mesas temáticas son espacios de trabajo grupal pequeños y "
            "recurrentes, organizados alrededor de un sector específico (salud, "
            "comercio, hacienda, sector territorial, entre otros) o de una "
            "temática transversal (tableros para la decisión, ética e "
            "inteligencia artificial, analítica territorial, gobernanza "
            "participativa de datos, calidad del dato). A diferencia del "
            "webinar —donde la mayoría escucha y pocos hablan— en la mesa todos "
            "los participantes hablan. El facilitador, de la Dirección de "
            "Desarrollo Digital, abre la sesión, modera y captura aprendizajes, "
            "pero el centro de la conversación está en los pares.")
    parrafo(doc,
            "La lógica de las mesas se inspira en las comunidades de práctica: "
            "una entidad presenta un caso concreto que tiene entre manos —un "
            "piloto en marcha, una dificultad metodológica, un dilema ético, "
            "una decisión de gobernanza pendiente— y las demás aportan desde su "
            "propia experiencia y contexto. La sesión termina con un cierre "
            "breve que sintetiza lo aprendido y compromete a las siguientes.")
    parrafo(doc,
            "Tres reglas de conversación sostienen ese formato y conviene "
            "enunciarlas desde la primera sesión: se comparte el método y no el "
            "dato reservado; lo que se dice en la mesa no se atribuye por fuera "
            "de ella sin autorización de quien lo dijo; y se trae el problema "
            "abierto, no la vitrina terminada. Sin la tercera regla, la mesa se "
            "convierte en una sucesión de presentaciones institucionales y "
            "pierde su razón de ser.")

    titulo(doc, "3.2. Formato operativo", 2)
    tabla(doc, "Formato operativo de las mesas temáticas",
          ["Atributo", "Descripción"],
          [["Plataforma",
            "Microsoft Teams, con grabación opcional para los miembros que no "
            "puedan asistir en vivo."],
           ["Duración",
            "90 minutos en la mesa piloto. Si la asistencia se resiente, se "
            "ajusta a 60 minutos reduciendo el bloque de intercambio (véase el "
            "capítulo 8)."],
           ["Frecuencia",
            "Mensual por mesa. Cada participante asiste solo a la mesa o mesas "
            "en las que tenga interés temático o sectorial."],
           ["Tamaño",
            "Entre 10 y 15 entidades por mesa. Si una mesa supera ese tamaño de "
            "forma sostenida, se divide en dos."],
           ["Convocatoria",
            "Agenda compartida con 48 horas de anticipación, con el caso de la "
            "sesión y la pregunta orientadora ya enunciados."],
           ["Facilitación",
            "Un facilitador titular del equipo de la Dirección de Desarrollo "
            "Digital, con apoyo rotativo de un cofacilitador para construir "
            "capacidad interna."],
           ["Estructura típica",
            "Apertura (5 min) · presentación del caso por una entidad (20 min) "
            "· intercambio entre pares (40 min) · aportes técnicos de la "
            "Dirección (15 min) · cierre y compromisos (10 min)."],
           ["Memoria",
            "Nota de una página publicada dentro de las 48 horas siguientes a "
            "la sesión, con aprendizajes, compromisos y responsables."]],
          "elaboración propia.", anchos=[3.6, 12.4])

    titulo(doc, "3.3. Roles y responsabilidades", 2)
    parrafo(doc,
            "El éxito operativo de las mesas depende de roles asignados desde "
            "el inicio. Los tres primeros son indispensables para que la mesa "
            "funcione; el cuarto es el que permite que la red devuelva valor a "
            "quien más aporta.")
    tabla(doc, "Roles dentro de cada mesa temática",
          ["Rol", "Quién lo asume", "Responsabilidad"],
          [["Facilitador titular", "Dirección de Desarrollo Digital",
            "Convoca, modera, captura aprendizajes y asegura la continuidad "
            "entre sesiones. Es el rostro estable de la mesa."],
           ["Anfitrión de sesión", "Entidad miembro, rotativo",
            "Cada mes una entidad presenta un caso y propone la pregunta "
            "orientadora. La rotación distribuye protagonismo y compromiso."],
           ["Relator", "Dirección de Desarrollo Digital, rotativo",
            "Toma nota de aprendizajes y compromisos y produce la nota de una "
            "página que alimenta boletines y eventos de la red."],
           ["Mentor o padrino temático", "Miembro de la comunidad, voluntario",
            "Persona de la comunidad con interés en capacitar o acompañar. "
            "Aporta experiencia en el tema de la mesa, orienta a las entidades "
            "que van empezando y puede liderar una sesión práctica. Es la "
            "figura que reconoce y retiene al miembro más activo."]],
          "elaboración propia.", anchos=[3.4, 4.0, 8.6])
    parrafo(doc,
            "El rol de mentor no exige dedicación fija ni vínculo contractual "
            "con el DNP: se ofrece como reconocimiento y se activa por "
            "invitación, sesión a sesión. Su identificación es uno de los "
            "resultados esperados de los diálogos bilaterales del capítulo 5.")

    titulo(doc, "3.4. Ruta de mesas para el primer año", 2)
    parrafo(doc,
            "La selección de mesas debe responder a los temas prioritarios de "
            "RENADIA y a la demanda real de los miembros. La siguiente ruta es "
            "un punto de partida para discusión y ajuste con la coordinación, "
            "no una lista cerrada. El orden importa: la primera mesa fija el "
            "tono y demuestra que el formato produce algo.")
    tabla(doc, "Ruta propuesta de mesas temáticas para el primer año",
          ["Orden", "Mesa", "Foco principal"],
          [["Piloto",
            "Tableros para la decisión pública: creación, uso y administración",
            "Ciclo de vida de los tableros en entidades nacionales y "
            "territoriales: de la decisión que soportan a su sostenibilidad. Se "
            "desarrolla en el capítulo 4."],
           ["Segunda", "Ética e inteligencia artificial responsable",
            "Casos de uso responsable de IA en entidades públicas: sesgos, "
            "transparencia, evaluación de impacto algorítmico, marcos éticos "
            "aplicados."],
           ["Tercera", "Analítica territorial",
            "Casos y retos de analítica en entidades territoriales o con "
            "presencia regional: calidad de datos, conectividad, talento, "
            "articulación nación–territorio."],
           ["Cuarta", "Gobernanza de datos",
            "Modelos de gobernanza interna, roles, comités de datos, políticas "
            "de acceso e interoperabilidad sectorial."],
           ["Sectoriales",
            "Mesas por sector, empezando por salud y por comercio y "
            "productividad",
            "Analítica e IA aplicada a un sector: predicción epidemiológica y "
            "gestión clínica en salud; facilitación al comercio, formalización "
            "e indicadores de cadena en el sector productivo."]],
          "elaboración propia.", anchos=[2.2, 4.8, 9.0])
    parrafo(doc,
            "Las mesas sectoriales se abren cuando exista masa crítica de "
            "entidades del sector interesadas, criterio que los diálogos "
            "bilaterales permiten verificar antes de convocar. Abrir una mesa "
            "sin demanda comprobada y tener que cerrarla después es más costoso "
            "para la credibilidad de la red que abrirla tarde.")

    titulo(doc, "3.5. Productos esperados de las mesas", 2)
    parrafo(doc, "Cada mesa, en un año de operación, produce:")
    for t in [
        "Entre 10 y 12 notas cortas de aprendizaje, una por sesión, que pueden "
        "alimentar boletines, micrositios y eventos de RENADIA.",
        "Un repositorio de casos de uso reales del sector público colombiano, "
        "categorizado por sector y por desafío.",
        "La identificación temprana de ponentes para webinars y de validadores "
        "potenciales para los espacios públicos de la red.",
        "Insumos para la construcción de estándares y lineamientos técnicos en "
        "los temas prioritarios de RENADIA.",
        "Un producto colectivo anual por mesa —guía corta, listado de buenas "
        "prácticas o contribución a un estándar— que da visibilidad a las "
        "entidades participantes.",
    ]:
        vineta(doc, t)


# --------------------------------------------------------------------------
# 4. Mesa piloto: tableros
# --------------------------------------------------------------------------

def mesa_piloto(doc):
    salto(doc)
    titulo(doc, "4. Mesa piloto: creación, uso y administración de tableros en "
                "entidades públicas nacionales y territoriales", 1)
    parrafo(doc,
            "Este capítulo desarrolla la primera mesa temática. Se propone "
            "abrir el ciclo con el tema de los tableros —los cuadros de mando o "
            "dashboards con los que las entidades siguen su gestión y publican "
            "sus cifras— por razones de demanda, de transversalidad y de "
            "capacidad de mostrar resultados en poco tiempo.")

    titulo(doc, "4.1. Por qué esta mesa abre el ciclo", 2)
    parrafo(doc,
            "El tablero es el punto donde converge todo lo demás. Construir uno "
            "obliga a resolver, en un solo producto y en un plazo corto, la "
            "calidad del dato, su gobernanza, su visualización y la decisión "
            "que debería soportar. Es, por eso, una puerta de entrada concreta "
            "a discusiones que en abstracto no convocan: pocas entidades "
            "inscriben a su equipo en una mesa sobre «gobernanza del dato», "
            "pero casi todas tienen un tablero que no se actualiza, que nadie "
            "usa o que muestra una cifra distinta a la de otra dependencia.")
    parrafo(doc,
            "Es además el tema más transversal disponible. No excluye a ningún "
            "sector ni a ningún nivel de gobierno: una secretaría de planeación "
            "municipal y un ministerio enfrentan el mismo problema a distinta "
            "escala, lo que hace posible una mesa mixta nación–territorio desde "
            "la primera sesión. Esa mezcla es difícil de lograr en mesas "
            "sectoriales y es justamente la que RENADIA necesita demostrar.")
    parrafo(doc,
            "Los insumos de la convocatoria «Mundial 2026» apuntan en la misma "
            "dirección. Entre los desafíos que los participantes describieron "
            "en sus propias palabras aparecen con nitidez los tres dolores que "
            "esta mesa aborda: la dispersión del dato —«las entidades tienen "
            "muchos datos al interior de forma dispersa en distintas bases de "
            "datos y poco personal para explotarlos»—, la centralización y la "
            "interoperabilidad, y la pérdida de memoria técnica: «el trabajo se "
            "acaba cuando cambian de contratistas y se pierde el conocimiento, "
            "falta de empalme y guardar información». Este último testimonio, "
            "de una entidad territorial, describe con precisión el modo de "
            "falla más común de los tableros públicos.")
    parrafo(doc,
            "Finalmente, es una mesa que produce evidencia visible rápido. Al "
            "cierre del primer ciclo de seis sesiones existirán un inventario "
            "de tableros del sector público y una guía corta de buenas "
            "prácticas: dos productos mostrables que sirven para justificar la "
            "continuidad del formato ante la coordinación y ante los propios "
            "miembros.")

    titulo(doc, "4.2. Alcance temático: el ciclo de vida de un tablero", 2)
    parrafo(doc,
            "Para que la mesa no se disperse, su alcance se organiza como el "
            "ciclo de vida de un tablero, de la decisión que lo origina a su "
            "retiro. Cada etapa corresponde a una conversación distinta y a un "
            "tipo de caso distinto, y ninguna es exclusivamente técnica.")
    tabla(doc, "Ciclo de vida de un tablero y contenido de la discusión",
          ["Etapa", "Qué se discute en la mesa"],
          [["1. Decisión y pregunta",
            "Qué decisión concreta soporta el tablero, quién la toma y con qué "
            "periodicidad. Un tablero sin decisión asociada es un informe con "
            "colores. Se discute cómo pasar del «quiero ver mis datos» a una "
            "pregunta accionable."],
           ["2. Fuentes y calidad del dato",
            "De dónde sale cada indicador, con qué frecuencia se actualiza, qué "
            "ocurre cuando la fuente cambia de estructura o de responsable. "
            "Mínimos de calidad, trazabilidad hasta el registro original y "
            "manejo honesto del dato faltante o preliminar."],
           ["3. Modelo e infraestructura",
            "Herramientas disponibles y sus costos reales de licenciamiento, "
            "alternativas libres, dónde vive el dato, cómo se automatiza la "
            "actualización y cómo se evita que el tablero dependa del computador "
            "de una sola persona."],
           ["4. Diseño y lectura",
            "Qué gráfico responde a qué pregunta, jerarquía visual, lenguaje "
            "claro, accesibilidad —contraste, tamaños, lectura en móvil— y "
            "errores frecuentes: exceso de indicadores en una pantalla, "
            "semáforos sin umbral definido, cifras sin denominador."],
           ["5. Publicación y apertura",
            "Cuándo el tablero es interno y cuándo debe ser público; deberes de "
            "transparencia activa de la Ley 1712 de 2014; tratamiento de datos "
            "personales y anonimización bajo la Ley 1581 de 2012; publicación "
            "del dato fuente en el portal de datos abiertos."],
           ["6. Administración y sostenibilidad",
            "Roles de dueño del dato, responsable técnico y aprobador; "
            "documentación mínima; ficha técnica del tablero; empalme y "
            "continuidad ante la rotación de contratistas y los cambios de "
            "administración."],
           ["7. Uso, evaluación y retiro",
            "Cómo se mide si el tablero se usa y si cambió alguna decisión; qué "
            "hacer con los tableros que nadie consulta; criterios para "
            "archivarlos sin perder la serie histórica."]],
          "elaboración propia.", anchos=[4.0, 12.0])
    parrafo(doc,
            "El alcance excluye deliberadamente dos cosas. No es una mesa de "
            "capacitación en una herramienta: no se enseña a usar un software "
            "determinado, porque para eso existen los webinars y la oferta "
            "formativa del mercado. Y no es un espacio de proveedores: las "
            "sesiones discuten decisiones y datos, no productos comerciales.")

    titulo(doc, "4.3. Preguntas que la mesa se propone responder", 2)
    parrafo(doc,
            "Enunciar las preguntas por adelantado ayuda a que cada entidad "
            "sepa si el espacio le sirve, y le da al facilitador un criterio "
            "para reconducir la conversación cuando se dispersa. Las preguntas "
            "de arranque son:")
    for t in [
        "¿Cuántos tableros tiene realmente la entidad, quién los mantiene y "
        "cuántos siguen vivos?",
        "¿Qué ocurre con un tablero cuando se va el contratista que lo "
        "construyó, y qué documentación mínima evita empezar de cero?",
        "¿Cómo se evita que dos dependencias publiquen cifras distintas sobre "
        "el mismo fenómeno, y quién arbitra cuando ocurre?",
        "¿Cuándo un tablero debe ser público, qué se publica y qué no, y cómo "
        "se maneja el dato preliminar o en revisión?",
        "¿Qué mínimos debería cumplir un tablero público del Estado colombiano "
        "para que la ciudadanía pueda confiar en él?",
        "¿Cómo se pasa del tablero descriptivo a la analítica avanzada y a la "
        "inteligencia artificial sin saltarse los pasos previos?",
    ]:
        vineta(doc, t)

    titulo(doc, "4.4. Agenda propuesta para el primer ciclo", 2)
    parrafo(doc,
            "El primer ciclo son seis sesiones mensuales. La secuencia recorre "
            "el ciclo de vida del tablero, pero cada sesión es autocontenida: "
            "quien llegue en la cuarta no queda descolgado. Cada sesión tiene un "
            "caso anfitrión y un producto parcial, de modo que la guía final se "
            "va escribiendo sesión a sesión y no al final.")
    tabla(doc, "Agenda del primer ciclo de la mesa de tableros",
          ["Sesión", "Tema", "Caso que se busca", "Producto parcial"],
          [["1", "Qué tenemos: inventario y dolores",
            "Cada entidad presenta en tres minutos su tablero más usado y su "
            "mayor dificultad. No hay caso único.",
            "Mapa inicial de tableros y de dolores compartidos"],
           ["2", "Del indicador a la decisión",
            "Una entidad cuyo tablero cambió —o no logró cambiar— una decisión "
            "concreta de gestión.",
            "Lista de preguntas previas a la construcción"],
           ["3", "La cocina: fuentes, calidad y actualización",
            "Un tablero que depende de una fuente frágil o de un cargue manual.",
            "Mínimos de calidad y trazabilidad"],
           ["4", "Herramientas, licencias y sostenibilidad técnica",
            "Una entidad territorial con restricción presupuestal que resolvió "
            "con recursos limitados.",
            "Comparación de opciones y sus costos"],
           ["5", "Diseño, lenguaje claro y accesibilidad",
            "Clínica de tableros: dos entidades someten un tablero real a "
            "revisión entre pares.",
            "Lista de verificación de diseño"],
           ["6", "Publicar o no publicar",
            "Un caso con datos sensibles o con cifras en revisión, y la decisión "
            "que se tomó.",
            "Lista de verificación de publicación y cierre de la guía"]],
          "elaboración propia.", anchos=[1.8, 3.8, 6.0, 4.4])
    parrafo(doc,
            "Las sesiones 7 a 12 se definen con los participantes al cierre del "
            "primer ciclo, sobre la base de lo que la propia mesa haya "
            "identificado. Los temas candidatos, en orden de interés previsible, "
            "son la articulación nación–territorio en tableros compartidos, la "
            "cifra única y el arbitraje entre dependencias, la medición del uso "
            "real de los tableros, el paso del tablero descriptivo al modelo "
            "predictivo y el uso de asistentes conversacionales sobre tableros "
            "existentes. Dejar la segunda mitad abierta es deliberado: es la "
            "prueba de que la mesa la construyen los miembros.")

    titulo(doc, "4.5. Selección de casos y composición de la mesa", 2)
    parrafo(doc,
            "La calidad de la mesa depende menos del facilitador que de los "
            "casos que se traen. Se proponen cuatro criterios de selección: que "
            "el caso esté en curso y no terminado; que tenga un problema abierto "
            "que la entidad esté dispuesta a exponer; que pueda describirse sin "
            "comprometer información reservada; y que aporte una escala distinta "
            "a la de la sesión anterior, alternando entidades nacionales y "
            "territoriales.")
    parrafo(doc,
            "La composición se cura, no se abre. Para el piloto se propone un "
            "grupo de 12 a 15 entidades con tres condiciones: al menos cuatro "
            "entidades territoriales, para que la conversación no quede anclada "
            "en la escala nacional; al menos dos entidades con tableros "
            "públicos consolidados, que puedan actuar como mentoras; y perfiles "
            "mixtos dentro de cada entidad, invitando tanto a quien construye "
            "el tablero como a quien lo usa para decidir. Una mesa compuesta "
            "solo por perfiles técnicos reproduce el problema que la mesa quiere "
            "resolver.")
    parrafo(doc,
            "A modo ilustrativo, y sujeto a validación con la coordinación, el "
            "universo de invitación puede construirse a partir de entidades "
            "nacionales con tableros públicos en operación —el propio DNP con "
            "TerriData y el seguimiento a documentos CONPES, el DANE, el "
            "Ministerio de Salud y Protección Social, el MinTIC, la Función "
            "Pública—, de secretarías de planeación y de TIC de ciudades "
            "capitales y de gobernaciones, y de las entidades que participaron "
            "en la convocatoria «Mundial 2026» y manifestaron interés en las "
            "mesas temáticas. Este último grupo tiene la ventaja de que ya "
            "expresó su disposición y su desafío por escrito.")

    titulo(doc, "4.6. Productos de la mesa piloto", 2)
    parrafo(doc,
            "La mesa entrega cinco productos, todos derivados de la propia "
            "conversación y sin trabajo adicional significativo para las "
            "entidades participantes:")
    for t in [
        "Inventario colaborativo de tableros del sector público, con una ficha "
        "por tablero: entidad, propósito, decisión que apoya, fuentes, "
        "frecuencia de actualización, herramienta, carácter público o interno y "
        "responsable. Es el primer mapa de este tipo en la red.",
        "Guía corta «Diez preguntas antes de construir un tablero», escrita en "
        "lenguaje claro y dirigida tanto al equipo técnico como al directivo "
        "que lo solicita.",
        "Lista de verificación de publicación, que integra en una página los "
        "deberes de transparencia, el tratamiento de datos personales, la "
        "accesibilidad y la ficha técnica mínima.",
        "Plantilla de ficha técnica del tablero, que es el antídoto concreto "
        "contra la pérdida de memoria por rotación de contratistas.",
        "Seis notas de aprendizaje, una por sesión, y al menos dos casos "
        "identificados y listos para llevar a un webinar de la red.",
    ]:
        vineta(doc, t)
    parrafo(doc,
            "Los cinco productos, en conjunto, constituyen el insumo natural "
            "para un lineamiento técnico de RENADIA sobre tableros en el sector "
            "público, si la coordinación decide llevarlo a ese nivel.")

    titulo(doc, "4.7. Riesgos de la mesa y forma de mitigarlos", 2)
    tabla(doc, "Riesgos identificados para la mesa piloto",
          ["Riesgo", "Señal temprana", "Mitigación"],
          [["La mesa se convierte en una feria de herramientas",
            "Las sesiones giran alrededor de marcas y funcionalidades",
            "Regla explícita desde la sesión 1: se discuten decisiones y datos, "
            "no productos. Sin proveedores en la mesa piloto."],
           ["Hablan solo las entidades grandes",
            "Dos o tres entidades concentran la palabra en sesiones seguidas",
            "Rotación obligatoria del anfitrión y compromiso de que al menos una "
            "de cada tres sesiones la abra una entidad territorial."],
           ["Solo llegan casos de vitrina",
            "Nadie presenta un problema sin resolver",
            "El facilitador pide explícitamente el tablero que no funcionó, y la "
            "sesión 1 se abre con un dolor del propio DNP."],
           ["Se pide compartir información reservada",
            "Incomodidad o silencio al describir las fuentes",
            "Regla de conversación: se comparte el método, no el dato. Lo dicho "
            "en la mesa no se atribuye por fuera sin autorización."],
           ["La mesa produce conversación pero no memoria",
            "Pasan dos sesiones sin nota publicada",
            "Relator designado por sesión y nota de una página dentro de las 48 "
            "horas siguientes, como condición para convocar la siguiente."],
           ["Deserción por agenda",
            "La asistencia cae por debajo de la mitad en la tercera sesión",
            "Fecha fija mensual reservada con un ciclo de anticipación, "
            "grabación del bloque de caso y nota que permite seguir el hilo a "
            "quien faltó."]],
          "elaboración propia.", anchos=[4.2, 4.4, 7.4])

    titulo(doc, "4.8. Seguimiento de la mesa", 2)
    parrafo(doc,
            "Al cierre del primer ciclo, la coordinación debe poder decidir con "
            "evidencia si el formato continúa, se ajusta o se cierra. Se "
            "proponen seis metas de referencia, ajustables antes de la primera "
            "convocatoria:")
    tabla(doc, "Metas de referencia para el primer ciclo de seis sesiones",
          ["Aspecto", "Meta de referencia"],
          [["Asistencia sostenida",
            "Al menos el 60 % de las entidades convocadas asiste a cuatro o más "
            "de las seis sesiones"],
           ["Rotación efectiva",
            "Seis entidades anfitrionas distintas, de las cuales al menos dos "
            "territoriales"],
           ["Memoria",
            "Seis notas de aprendizaje publicadas dentro del plazo de 48 horas"],
           ["Inventario",
            "Al menos veinte fichas de tablero cargadas por las entidades "
            "participantes"],
           ["Aporte a otros canales",
            "Al menos dos ponentes y dos contenidos derivados hacia webinars y "
            "boletines"],
           ["Producto colectivo",
            "Guía corta y listas de verificación publicadas al cierre del ciclo"]],
          "elaboración propia.", anchos=[4.5, 11.5])


# --------------------------------------------------------------------------
# 5. Diálogos bilaterales
# --------------------------------------------------------------------------

def dialogos(doc):
    salto(doc)
    titulo(doc, "5. Estrategia 2: diálogos bilaterales RENADIA", 1)

    titulo(doc, "5.1. Qué son los diálogos bilaterales", 2)
    parrafo(doc,
            "Los diálogos bilaterales son reuniones uno a uno entre el equipo "
            "de la Dirección de Desarrollo Digital y cada entidad miembro de "
            "RENADIA. A diferencia de las mesas temáticas no son grupales: son "
            "conversaciones personalizadas en las que se revisa la situación "
            "específica de la entidad —dónde está, qué necesita, qué "
            "oportunidades tiene— y se ajusta la relación con la red en función "
            "de esa realidad concreta.")
    parrafo(doc,
            "Esta estrategia parte de un reconocimiento que las mesas no pueden "
            "atender por sí solas: cada entidad de RENADIA tiene una "
            "trayectoria propia y no avanza al ritmo del colectivo. Una entidad "
            "con alta madurez analítica puede estar lista para aportar como "
            "mentora; una entidad en niveles básicos necesita acompañamiento "
            "puntual y orientación sobre por dónde empezar. Tratar a todas por "
            "igual es ineficiente y desmotivador. Los diálogos bilaterales son "
            "el mecanismo para calibrar la relación con cada miembro.")

    titulo(doc, "5.2. Formato operativo", 2)
    tabla(doc, "Formato operativo de los diálogos bilaterales",
          ["Atributo", "Descripción"],
          [["Plataforma",
            "Microsoft Teams. La sesión se agenda directamente con el punto "
            "focal de la entidad."],
           ["Duración",
            "30 a 45 minutos. Una agenda más larga desincentiva la frecuencia y "
            "reduce la calidad de la conversación."],
           ["Frecuencia",
            "Trimestral para entidades prioritarias y semestral para el resto "
            "de los miembros. Esta cadencia se ajusta según el mapeo inicial."],
           ["Participantes por el DNP",
            "Un líder técnico de la Dirección de Desarrollo Digital, "
            "acompañado idealmente por el profesional de relacionamiento de la "
            "red."],
           ["Participantes por la entidad",
            "Punto focal técnico y, cuando aplique, líder estratégico —director "
            "o jefe de área—. No requiere convocatoria masiva interna."],
           ["Preparación previa",
            "Revisión del historial de interacción con la entidad: mesas en que "
            "participa, asistencia a webinars, contenidos que ha consumido y "
            "compromisos previos."],
           ["Registro",
            "Ficha de relacionamiento actualizada el mismo día, con "
            "compromisos, responsables y fecha de revisión."]],
          "elaboración propia.", anchos=[4.2, 11.8])

    titulo(doc, "5.3. Criterios para priorizar entidades", 2)
    parrafo(doc,
            "Dado que el equipo de la Dirección de Desarrollo Digital tiene "
            "capacidad operativa limitada, no es realista mantener una cadencia "
            "trimestral con todas las entidades miembro desde el primer día. Se "
            "propone una clasificación en tres niveles, revisable anualmente:")
    tabla(doc, "Niveles de priorización de entidades para el diálogo bilateral",
          ["Nivel", "Características de la entidad", "Cadencia sugerida"],
          [["Nivel A — Estratégico",
            "Entidades cabeza de sector, con alta capacidad analítica, alta "
            "visibilidad pública o con proyectos insignia en datos e "
            "inteligencia artificial.",
            "Cada 3 meses"],
           ["Nivel B — Activo",
            "Entidades con participación regular en mesas o webinars, con "
            "avances reportados o con iniciativas en marcha.",
            "Cada 4 a 6 meses"],
           ["Nivel C — En vinculación",
            "Entidades miembro con baja interacción reciente o en proceso de "
            "incorporación a la dinámica de la red.",
            "Cada 6 a 12 meses"]],
          "elaboración propia.", anchos=[3.6, 8.8, 3.6])
    parrafo(doc,
            "La clasificación describe el momento de la relación, no la calidad "
            "de la entidad, y por eso es móvil: una entidad de nivel C que "
            "empieza a participar en la mesa piloto pasa a nivel B en la "
            "siguiente revisión. Conviene que el criterio de movilidad sea "
            "explícito para que el nivel no se lea como una jerarquía.")

    titulo(doc, "5.4. Agenda sugerida para una sesión bilateral", 2)
    parrafo(doc,
            "Una agenda estándar permite que la sesión sea ágil y comparable "
            "entre entidades, lo que facilita el mapeo posterior y la "
            "trazabilidad. La estructura sugerida, para una sesión de 30 a 45 "
            "minutos, es la siguiente:")
    tabla(doc, "Estructura de una sesión bilateral",
          ["Momento", "Duración", "Contenido"],
          [["Apertura y contexto", "5 min",
            "Saludo, recordatorio del propósito del espacio y revisión de los "
            "compromisos del diálogo anterior."],
           ["¿En qué están?", "10 min",
            "La entidad cuenta en qué proyectos de datos e inteligencia "
            "artificial está trabajando, qué ha cambiado desde la última "
            "conversación y qué retos enfrenta."],
           ["¿Qué necesitan?", "10 min",
            "Identificación de necesidades concretas —datos, talento, "
            "articulación, lineamientos, conexión con pares— y discusión de "
            "cómo RENADIA puede aportar."],
           ["¿Qué pueden aportar?", "5 min",
            "Capacidades de la entidad útiles para la red: un caso de éxito, un "
            "ponente, una metodología, un conjunto de datos abierto o la "
            "disposición a actuar como mentor en una mesa."],
           ["¿A quién más deberíamos hablarle?", "5 min",
            "Invitación explícita a compartir contactos: áreas de la misma "
            "entidad que no están vinculadas, entidades pares del sector o del "
            "territorio, y personas con interés en capacitar o acompañar. Es la "
            "vía de crecimiento más barata que tiene la red."],
           ["Cierre y compromisos", "5 min",
            "Acuerdos concretos para los siguientes 90 días, con responsable y "
            "fecha. Sin compromisos no hay seguimiento."]],
          "elaboración propia.", anchos=[4.0, 2.0, 10.0])
    parrafo(doc,
            "El paso de referidos merece una precisión operativa: el contacto "
            "no se pide en abstracto sino con un motivo concreto —una mesa "
            "abierta, un webinar en preparación, una convocatoria—, y quien lo "
            "entrega recibe copia de la invitación que se envíe. Pedir "
            "contactos sin decir para qué desgasta la confianza que el espacio "
            "busca construir.")

    titulo(doc, "5.5. Productos esperados de los diálogos", 2)
    for t in [
        "Una ficha de relacionamiento por entidad miembro, actualizada después "
        "de cada diálogo, que mapea madurez, necesidades, capacidades y "
        "compromisos.",
        "Un mapa consolidado del ecosistema RENADIA: dónde están las "
        "fortalezas, dónde las brechas y dónde las oportunidades de "
        "articulación nación–territorio o intersectorial.",
        "La detección temprana de entidades en riesgo de desvinculación, antes "
        "de que la pérdida sea irreversible.",
        "La identificación oportuna de casos, ponentes, mentores y aliados para "
        "mesas, webinars, boletines y otros productos de la red.",
        "Una percepción más precisa de cómo está siendo recibida RENADIA por "
        "sus miembros, en términos cualitativos que no aparecen en los reportes "
        "operativos.",
    ]:
        vineta(doc, t)


# --------------------------------------------------------------------------
# 6, 7, 8, 9
# --------------------------------------------------------------------------

def implementacion(doc):
    titulo(doc, "6. Implementación por fases", 1)
    parrafo(doc,
            "Las dos estrategias pueden implementarse de manera escalonada, sin "
            "sobrecargar al equipo. Se propone un arranque en tres fases que "
            "permite estabilizar los procesos antes de ampliar el alcance. La "
            "fase 1 arranca con una sola mesa —la piloto de tableros— porque "
            "abrir dos formatos nuevos al tiempo duplica el riesgo operativo sin "
            "duplicar el aprendizaje.")
    tabla(doc, "Fases de implementación propuestas",
          ["Fase", "Alcance", "Hitos principales"],
          [["Fase 1: pilotaje (meses 1 y 2)",
            "Mesa piloto de tableros y diálogos bilaterales con cinco entidades "
            "de nivel A.",
            "Decisiones del capítulo 8 resueltas, curaduría de entidades "
            "invitadas, primera y segunda sesión de la mesa, primeros cinco "
            "diálogos, ajuste del formato con la retroalimentación recibida."],
           ["Fase 2: consolidación (meses 3 a 6)",
            "Mesa piloto en régimen mensual, apertura de la segunda mesa y "
            "diálogos con todas las entidades de nivel A y arranque del nivel B.",
            "Cierre del primer ciclo de seis sesiones con guía e inventario "
            "publicados, primera ronda completa de diálogos de nivel A, fichas "
            "de relacionamiento iniciales y primera nota consolidada para la "
            "coordinación."],
           ["Fase 3: régimen completo (mes 7 en adelante)",
            "Tres a cinco mesas en operación y diálogos bilaterales para los "
            "tres niveles según cadencia.",
            "Operación sostenida, balance anual, ajuste de mesas y de niveles, "
            "e integración plena con webinars, boletines y estándares."]],
          "elaboración propia.", anchos=[4.0, 5.5, 6.5])


def condiciones(doc):
    titulo(doc, "7. Condiciones para el éxito", 1)
    parrafo(doc,
            "Las dos estrategias son operativamente factibles con la capacidad "
            "actual del equipo, pero requieren tres condiciones para sostenerse "
            "en el tiempo:")
    for t in [
        "Constancia antes que perfección. Las mesas y los diálogos generan "
        "valor por su recurrencia. Es preferible una mesa mensual sencilla y "
        "bien preparada que una mesa trimestral sobreproducida: la regularidad "
        "es lo que construye comunidad.",
        "Capacidad de captura y memoria. Una nota de una página después de cada "
        "mesa y una ficha actualizada después de cada diálogo. Sin captura "
        "sistemática, los espacios producen conversación pero no producen "
        "aprendizaje organizacional para RENADIA ni para el DNP.",
        "Reciprocidad institucional. Cuando una entidad plantea una necesidad "
        "concreta en una mesa o en un diálogo debe haber respuesta visible "
        "dentro de un plazo razonable. La credibilidad de la red se construye, "
        "sesión a sesión, por la utilidad real que produce.",
    ]:
        vineta(doc, t)


def decisiones(doc):
    titulo(doc, "8. Decisiones que requiere la coordinación", 1)
    parrafo(doc,
            "Las definiciones que siguen no dependen del diseño sino de la "
            "coordinación de la red, y conviene resolverlas antes de convocar "
            "la primera sesión. Para cada una se presentan las opciones y una "
            "recomendación argumentada.")
    tabla(doc, "Decisiones pendientes, opciones y recomendación",
          ["Decisión", "Opciones", "Recomendación"],
          [["Duración de la sesión", "60 o 90 minutos",
            "90 minutos en el piloto. Con 60 minutos el bloque de intercambio "
            "entre pares queda en 20 y la mesa se parece a un webinar. Si la "
            "asistencia se resiente, se baja a 60 desde la sesión 4."],
           ["Segunda mesa del ciclo",
            "Ética e IA responsable, analítica territorial o una sectorial",
            "Ética e inteligencia artificial responsable, porque cubre la otra "
            "mitad de la agenda de RENADIA y convoca a un público distinto al "
            "de tableros, en vez de competir por la misma audiencia."],
           ["Grabación de las sesiones",
            "Grabar todo, grabar solo el caso o no grabar",
            "Grabar únicamente la presentación del caso. Grabar el intercambio "
            "entre pares inhibe la franqueza, que es el activo del formato."],
           ["Convocatoria", "Abierta a todos los miembros o curada",
            "Curada en el piloto, con los criterios de composición del apartado "
            "4.5, y abierta a partir del segundo ciclo."],
           ["Repositorio de la memoria",
            "Canal de Teams, SharePoint de la red o micrositio de RENADIA",
            "Definirlo antes de la sesión 1 y anunciarlo en la convocatoria. La "
            "nota de la primera sesión debe tener dónde publicarse el mismo día "
            "en que se escribe."],
           ["Nombre público de la mesa",
            "«Mesa de tableros» o «Mesa de tableros para la decisión pública»",
            "La segunda. Nombra el propósito y no la herramienta, que es "
            "justamente la distinción que la mesa quiere instalar."]],
          "elaboración propia.", anchos=[3.4, 5.0, 7.6])


def cierre(doc):
    titulo(doc, "9. Cierre y próximos pasos", 1)
    parrafo(doc,
            "Las mesas temáticas y los diálogos bilaterales no buscan "
            "reemplazar los webinars ni los boletines: los potencian. Los "
            "webinars seguirán siendo el espacio de mayor alcance y visibilidad "
            "pública y los boletines seguirán siendo el canal de difusión de "
            "información estructurada. Lo que las dos estrategias aportan es lo "
            "que hoy falta en la red: contacto frecuente, conversación de doble "
            "vía y una relación viva con cada miembro.")
    parrafo(doc,
            "Si la coordinación lo considera viable, los próximos pasos "
            "sugeridos son cuatro: resolver las seis decisiones del capítulo 8; "
            "curar la lista de entidades invitadas a la mesa piloto de tableros "
            "y definir la entidad anfitriona de la primera sesión; seleccionar "
            "las cinco entidades de nivel A con las que arrancan los diálogos "
            "bilaterales; y fijar la fecha mensual de la mesa para todo el "
            "primer ciclo, de modo que las entidades puedan reservarla con "
            "anticipación. El equipo de la Dirección de Desarrollo Digital queda "
            "atento a observaciones, ajustes y a las decisiones de la "
            "coordinación para iniciar la implementación.")


# --------------------------------------------------------------------------

def construir(salida):
    doc = Document()
    configurar_estilos(doc)
    idioma_por_defecto(doc)

    # La portada va en su propia sección, sin encabezado: el logo ya está en
    # el cuerpo de la página y repetirlo arriba lo duplica.
    s0 = doc.sections[0]
    margenes(s0)
    s0.different_first_page_header_footer = True
    portada(doc)

    s1 = doc.add_section(WD_SECTION.NEW_PAGE)
    margenes(s1)
    encabezado_paginado(s1)

    control_documental(doc)
    siglas(doc)
    presentacion(doc)
    salto(doc)
    contexto(doc)
    marco(doc)
    mesas(doc)
    mesa_piloto(doc)
    dialogos(doc)
    implementacion(doc)
    condiciones(doc)
    decisiones(doc)
    cierre(doc)

    idioma_en_runs(doc)
    props = doc.core_properties
    props.title = ("Estrategias para un relacionamiento cercano con los "
                   "miembros de RENADIA")
    props.subject = "Mesas temáticas y diálogos bilaterales · versión 2"
    props.author = "Ricardo Esteban Ruiz Castro"
    props.category = "Documento interno de trabajo"
    props.language = IDIOMA
    doc.save(salida)
    return salida


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--salida", default=SALIDA_DEFECTO)
    ap.add_argument("--pdf", action="store_true",
                    help="convierte a PDF para revisión visual")
    args = ap.parse_args()
    os.makedirs(os.path.dirname(args.salida), exist_ok=True)
    construir(args.salida)
    print(f"Documento generado: {args.salida}")
    if args.pdf:
        subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf",
                        "--outdir", os.path.dirname(args.salida), args.salida],
                       capture_output=True, check=True)
        print("PDF de revisión generado.")


if __name__ == "__main__":
    main()
