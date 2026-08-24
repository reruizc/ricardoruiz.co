#!/usr/bin/env python3
# Genera Excel de soporte + sección Word del bloque Oviedo.
import json
from openpyxl import Workbook
from openpyxl.styles import Font,PatternFill,Alignment,Border,Side
from openpyxl.utils import get_column_letter
OUT='Bases de datos/output_pacto_1v_2026'
loc=json.load(open(f'{OUT}/oviedo_bogota_localidad.json'))
nac=json.load(open(f'{OUT}/oviedo_nacional.json'))
cor=json.load(open(f'{OUT}/oviedo_correlaciones.json'))
OX='8A1E16'; HDR='2A241B'

# ================= EXCEL =================
wb=Workbook(); thin=Side(style='thin',color='DDDDDD'); bd=Border(thin,thin,thin,thin)
def style_header(ws,ncol,row=1):
    for c in range(1,ncol+1):
        cell=ws.cell(row=row,column=c); cell.font=Font(bold=True,color='FFFFFF',name='Calibri',size=11)
        cell.fill=PatternFill('solid',fgColor=OX); cell.alignment=Alignment(horizontal='center',vertical='center',wrap_text=True); cell.border=bd
def autosize(ws):
    for col in ws.columns:
        w=max((len(str(c.value)) for c in col if c.value is not None),default=8)
        ws.column_dimensions[get_column_letter(col[0].column)].width=min(max(w+3,11),40)

# hoja 0 resumen
ws=wb.active; ws.title='Resumen'
res=[['Bloque Centro — El voto de Oviedo','',''],
     ['','',''],
     ['Indicador','Valor','Nota'],
     ['Oviedo · Gran Consulta (nacional)','1.259.004','2º lugar, detrás de Paloma (3.248.589)'],
     ['Oviedo · Bogotá','504.807','42,4% de todo su voto nacional'],
     ['Paloma · Gran Consulta','3.248.589','Ganó la consulta'],
     ['Paloma · 1ª vuelta','1.639.683','Perdió ~1,6M de su base de consulta'],
     ['Hueco de Cepeda en Bogotá (vs Petro 2022)','~258.000','6,3 pp — la mayor caída del país'],
     ['Oviedo Bogotá DISPUTABLE (localidades que ganó Cepeda)','179.208','36%'],
     ['Oviedo Bogotá YA en la derecha (localidades que ganó Abelardo)','317.763','64%'],
     ['','',''],
     ['VEREDICTO','',''],
     ['Adhesión de Oviedo a nivel NACIONAL','Poco relevante','Su voto es chico y muy concentrado en Bogotá'],
     ['Adhesión de Oviedo en BOGOTÁ','Relevante pero parcial','Solo ~36% disputable; el resto ya votó Abelardo'],
     ['Hipótesis del cliente (Paloma 1V = voto de Oviedo)','REFUTADA','Paloma 1V se parece a su propia base, no a la de Oviedo']]
for r in res: ws.append(r)
ws['A1'].font=Font(bold=True,size=15,color=OX); ws['A12'].font=Font(bold=True,size=12,color=HDR)
style_header(ws,3,row=3)
for rr in (12,):
    pass
autosize(ws); ws.column_dimensions['A'].width=46; ws.column_dimensions['C'].width=46

# hoja nacional
ws2=wb.create_sheet('Oviedo nacional (depto)')
ws2.append(['Departamento','Votos Oviedo (consulta)','% del Oviedo nacional'])
for r in nac['top_deptos']: ws2.append([r['depto'],r['oviedo'],r['pct_nac']/100])
style_header(ws2,3)
for row in ws2.iter_rows(min_row=2):
    row[1].number_format='#,##0'; row[2].number_format='0.0%'
    for c in row: c.border=bd
autosize(ws2)

# hoja Bogotá localidad
ws3=wb.create_sheet('Oviedo Bogotá (localidad)')
ws3.append(['Localidad','Votos Oviedo (consulta)','Ganó 1ª vuelta','Cepeda 1V %','Abelardo 1V %','Fajardo 1V %','Abstención %','Lectura'])
for r in sorted(loc,key=lambda x:-x['oviedo']):
    lect='DISPUTABLE' if r['win1v']=='Cepeda' else 'ya en la derecha'
    ws3.append([r['localidad'],r['oviedo'],r['win1v'],r['cep']/100,r['abe']/100,r['faj']/100,(r['abst'] or 0)/100,lect])
style_header(ws3,8)
for row in ws3.iter_rows(min_row=2):
    row[1].number_format='#,##0'
    for i in (3,4,5,6): row[i].number_format='0.0%'
    fill='EDE9F5' if row[2].value=='Cepeda' else 'E7E2D4'
    for c in row: c.border=bd; c.fill=PatternFill('solid',fgColor=fill)
autosize(ws3)

# hoja correlaciones
ws4=wb.create_sheet('Correlaciones')
ws4.append(['Correlación del voto de Oviedo (consulta) con el voto de 1ª vuelta, por puesto'])
ws4.append([])
ws4.append(['Ámbito','vs Cepeda','vs Fajardo','vs Abelardo','vs Paloma','n puestos'])
for k in ('nacional','bogota'):
    c=cor[k]; ws4.append([k.capitalize(),c['cepeda'],c['fajardo'],c['abelardo'],c['paloma'],c['n_puestos']])
ws4.append([])
ws4.append(['Interpretación: + = mismas zonas; − = zonas opuestas. Oviedo es voto de centro, anti-Abelardo.'])
ws4['A1'].font=Font(bold=True,size=12,color=HDR); style_header(ws4,6,row=3)
for row in ws4.iter_rows(min_row=4,max_row=5,min_col=2,max_col=5):
    for c in row: c.number_format='0.000'
autosize(ws4)
wb.save(f'{OUT}/Soporte_Oviedo.xlsx'); print('✓ Soporte_Oviedo.xlsx')

# ================= WORD =================
from docx import Document
from docx.shared import Pt,RGBColor,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
OXR=RGBColor(0x8a,0x1e,0x16); INKR=RGBColor(0x1a,0x15,0x10); GR=RGBColor(0x55,0x54,0x48)
d=Document()
for s in d.styles:
    if s.name=='Normal': s.font.name='Helvetica Neue'; s.font.size=Pt(11)
def shade(cell,hexc):
    el=OxmlElement('w:shd'); el.set(qn('w:fill'),hexc); cell._tc.get_or_add_tcPr().append(el)
def h(txt,size=15,color=OXR,after=6,before=10):
    p=d.add_paragraph(); r=p.add_run(txt); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); return p
def body(txt,size=11):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(8)
    # negritas con **
    import re
    for i,seg in enumerate(re.split(r'(\*\*.*?\*\*)',txt)):
        if seg.startswith('**'): r=p.add_run(seg[2:-2]); r.bold=True
        else: r=p.add_run(seg)
        r.font.size=Pt(size); r.font.color.rgb=INKR
    return p
def tbl(headers,rows,shade_col=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'
    for i,hh in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(255,255,255)
        shade(c,'8A1E16')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(9.5)
            if shade_col is not None and i==0:
                pass
    return t

# Título
p=d.add_paragraph(); r=p.add_run('Análisis Nacional Electoral · Primera vuelta 2026'); r.font.size=Pt(10); r.font.color.rgb=GR
p=d.add_paragraph(); r=p.add_run('Bloque Centro — El voto de Oviedo: ¿conviene una adhesión a Iván Cepeda?'); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=OXR
p.paragraph_format.space_after=Pt(4)
p=d.add_paragraph(); r=p.add_run('Preconteo Registraduría (escrutinio ~99,9%) + Gran Consulta del 8 de marzo de 2026. Cifras sujetas a escrutinio definitivo.'); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GR

h('La pregunta del cliente')
body('El cliente plantea una hipótesis: que la mayoría del voto que Paloma Valencia obtuvo en la **Gran Consulta** (3.248.589) se fue a Abelardo de la Espriella en primera vuelta, y que los 1.639.683 votos que Paloma sacó en 1ª vuelta serían, en realidad, el **voto de Juan Daniel Oviedo** —2º en esa consulta con 1.259.004— que se negó a irse con Abelardo. Sobre esa base, exploran una eventual **adhesión de Oviedo** que ayude a recuperar Bogotá, donde Cepeda quedó 6,3 puntos por debajo de Petro-2022.')

h('Hallazgo 1 — La hipótesis no se sostiene')
body('Cruzando a nivel de puesto el voto de Paloma en 1ª vuelta contra el de la consulta (proporciones, controlando tamaño), Paloma-1V se parece **más a su propia base de consulta (+0,21) que a la de Oviedo (−0,11)**, que incluso va levemente en contra. En otras palabras: el voto que Paloma **conservó** en 1ª vuelta es un residuo de su propia base dura, no el voto de Oviedo. **El hecho que lo explica:** Oviedo es un perfil técnico-urbano (ex-director del DANE, excandidato a la Alcaldía de Bogotá en 2023); su votante de consulta es un antiuribista de centro, distinto del núcleo uribista que sostuvo a Paloma.')

h('Hallazgo 2 — El voto de Oviedo es, ante todo, bogotano')
body('De su ~1,26 millones de votos de consulta, **Bogotá concentra ~505.000 (42,4%)** — cuatro veces más que el siguiente departamento. Fuera de la capital y su área metropolitana (Cundinamarca), su voto se diluye. Esto define el problema: **cualquier jugada con Oviedo es una jugada de Bogotá, no de país.**')
tbl(['Departamento','Votos Oviedo','% del Oviedo nacional'],[[r['depto'],f"{r['oviedo']:,}".replace(',','.'),f"{r['pct_nac']:.1f}%"] for r in nac['top_deptos'][:6]])
d.add_paragraph()

h('Hallazgo 3 — Es un voto de centro, anti-Abelardo (y por eso es disputable)')
body('La composición ideológica del voto de Oviedo es favorable para el Pacto: a nivel nacional va de la mano del de **Fajardo (+0,60)** y de **Cepeda (+0,32)**, y es lo **opuesto al de Abelardo (−0,41)**. En Bogotá el patrón es aún más nítido: **+0,55 con Cepeda y −0,61 con Abelardo**. Es decir, no es un voto amarrado a la ultraderecha; es el profesional urbano incómodo con Abelardo, justamente el perfil que una segunda vuelta a dos miedos pone en disputa.')
tbl(['Ámbito','vs Cepeda','vs Fajardo','vs Abelardo','vs Paloma'],
    [['Nacional','+0,32','+0,60','−0,41','−0,12'],['Bogotá','+0,55','+0,19','−0,61','−0,57']])
d.add_paragraph()

h('Hallazgo 4 — Pero dos tercios de ese voto ya están en la derecha')
body('Al bajarlo por localidad aparece el matiz decisivo: el voto de Oviedo se parte en la **línea de clase** de la ciudad. En el **norte rico** (Usaquén, Suba, Chapinero, Barrios Unidos) el votante de Oviedo **ya se tapó la nariz y votó Abelardo** en 1ª vuelta; en el **occidente y sur popular** (Kennedy, Bosa, Ciudad Bolívar, San Cristóbal, Usme) es anti-Abelardo y sigue disputable. El resultado:')
p=d.add_paragraph()
for txt,col in [('• DISPUTABLE (localidades que ganó Cepeda): 179.208 votos — 36%   ',RGBColor(0x53,0x4a,0x8f)),
                ('\n• YA EN LA DERECHA (localidades que ganó Abelardo): 317.763 — 64%',RGBColor(0x16,0x16,0x6b))]:
    r=p.add_run(txt); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=col
d.add_picture(f'{OUT}/g_oviedo_bogota_localidad.png',width=Inches(6.3))
d.add_paragraph()

h('Veredicto operativo',color=INKR)
body('**Nacional:** poco relevante. El voto de Oviedo es chico frente a la brecha nacional (la derecha ronda 51% y la izquierda 41%, ~2,4 millones de distancia) y está demasiado concentrado en Bogotá. Como jugada de país, es necesaria pero muy insuficiente.')
body('**Bogotá:** aquí sí tiene sentido, pero con asterisco. Traer a Oviedo **no le devuelve a Cepeda ~505 mil votos**; le pone en juego, realista, la franja anti-Abelardo del centro bogotano —del orden de **180 mil votos** en Kennedy, Bosa, Engativá y el sur—, que es donde se decide la ciudad. Eso **por sí solo no cierra los 6,3 puntos**, pero **combinado con bajar la abstención** en los bastiones del sur (Ciudad Bolívar 34%, Usme 33%, San Cristóbal 35%) sí dibuja un camino para volver al 48% de Petro. En una frase: **Oviedo es una jugada de Bogotá, no de país; y dentro de Bogotá vale por la mitad de su titular, pero esa mitad cae donde más la necesitan.**')

h('Nota metodológica',size=11,color=GR)
body('Voto de Oviedo: Gran Consulta del 8-mar-2026 a nivel de puesto. Voto de 1ª vuelta: preconteo Registraduría a nivel de mesa (~99,9% escrutado), agregado a puesto y unido al georreferenciado de puestos (barrio/localidad). Correlaciones de Pearson sobre proporciones por puesto (excluye zonas de censo y cárceles). "Disputable" = localidad donde Cepeda superó a Abelardo en 1ª vuelta; no implica transferencia automática de votos.')
d.save(f'{OUT}/Bloque_Oviedo.docx'); print('✓ Bloque_Oviedo.docx')
