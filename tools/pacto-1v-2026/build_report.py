#!/usr/bin/env python3
import json
from docx import Document
from docx.shared import Pt,RGBColor,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
OUT='Bases de datos/output_pacto_1v_2026'
BA=json.load(open(f'{OUT}/blocks_all.json')); BF=json.load(open(f'{OUT}/blocks_full.json')); DIF=json.load(open(f'{OUT}/dif_2022.json'))
EST=json.load(open(f'{OUT}/estrato_bogota.json'))
M2V=json.load(open(f'{OUT}/twov_model.json')); TER=json.load(open(f'{OUT}/twov_territorial.json'))
STRAT=json.load(open(f'{OUT}/twov_estrategias.json'))
NF=lambda n: f"{round(n):,}".replace(',','.')
OX=RGBColor(0x8a,0x1e,0x16); INK=RGBColor(0x1a,0x15,0x10); GR=RGBColor(0x5a,0x54,0x48)
FONT='Inter'
d=Document()
for s in d.styles:
    if s.name=='Normal': s.font.name=FONT; s.font.size=Pt(10.5)
# fuente por defecto del documento (docDefaults) → asegura que TODO (incl. celdas de tabla) use Inter
_st=d.styles.element; _dd=_st.find(qn('w:docDefaults'))
if _dd is not None:
    _rprd=_dd.find(qn('w:rPrDefault'))
    if _rprd is None: _rprd=OxmlElement('w:rPrDefault'); _dd.insert(0,_rprd)
    _rpr=_rprd.find(qn('w:rPr'))
    if _rpr is None: _rpr=OxmlElement('w:rPr'); _rprd.append(_rpr)
    _rf=_rpr.find(qn('w:rFonts'))
    if _rf is None: _rf=OxmlElement('w:rFonts'); _rpr.insert(0,_rf)
    for _a in ('w:ascii','w:hAnsi','w:cs'): _rf.set(qn(_a),FONT)
def shade(c,h): e=OxmlElement('w:shd'); e.set(qn('w:fill'),h); c._tc.get_or_add_tcPr().append(e)
def cap(num,txt):
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.page_break_before=True   # cada capítulo arranca en página nueva
    r=p.add_run(f'Capítulo {num}'); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=GR
    p2=d.add_paragraph(); r=p2.add_run(txt); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=OX; p2.paragraph_format.space_after=Pt(5)
def h(txt,size=12.5,color=INK,before=8,after=4):
    p=d.add_paragraph(); r=p.add_run(txt); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
def body(txt,size=10.5,after=6):
    import re; p=d.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in re.split(r'(\*\*.*?\*\*)',txt):
        r=p.add_run(seg[2:-2] if seg.startswith('**') else seg); r.font.size=Pt(size); r.bold=seg.startswith('**'); r.font.color.rgb=INK
def note(txt): body(txt,size=9,after=6)   # un solo tamaño para notas al pie / metodológicas
def bullet(txt):
    import re; p=d.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in re.split(r'(\*\*.*?\*\*)',txt):
        r=p.add_run(seg[2:-2] if seg.startswith('**') else seg); r.font.size=Pt(10.5); r.bold=seg.startswith('**'); r.font.color.rgb=INK
def tbl(headers,rows):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Table Grid'
    for i,hh in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(255,255,255); shade(c,'8A1E16')
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row): cs[i].text=''; r=cs[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(8.5)
def img(name,w=6.3):
    d.add_picture(f'{OUT}/{name}',width=Inches(w)); p=d.paragraphs[-1]
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(5)
BIG=lambda v,th=20000: v['base']>=th

# PORTADA
p=d.add_paragraph(); r=p.add_run('ANÁLISIS NACIONAL ELECTORAL'); r.font.size=Pt(11); r.bold=True; r.font.color.rgb=GR
p=d.add_paragraph(); r=p.add_run('Primera vuelta presidencial 2026'); r.bold=True; r.font.size=Pt(26); r.font.color.rgb=OX; p.paragraph_format.space_after=Pt(2)
p=d.add_paragraph(); r=p.add_run('Dónde están los votos y cómo se recupera la elección — diagnóstico por bloque, hasta nivel de barrio y puesto.'); r.font.size=Pt(12); r.font.color.rgb=INK
p=d.add_paragraph(); r=p.add_run('Documento de trabajo · uso interno · preparado para el Pacto Histórico. Preconteo Registraduría (~99,9%) + escrutinios 2022 a nivel de mesa + Gran Consulta de marzo 2026. Cifras sujetas a escrutinio definitivo.'); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GR
# índice
h('Contenido',size=12,color=OX,before=14)
for t in ['Resumen ejecutivo','Capítulo 1 — Petro 2022: ¿dónde creció de primera a segunda vuelta?','Capítulo 2 — Cepeda 2026 frente a Petro 2022 (con mapas nacional y municipal)','Capítulo 3 — El voto de centro (Fajardo, Claudia, Oviedo)','Capítulo 4 — La derecha hoy (Abelardo + Paloma)','Capítulo 5 — Bogotá por estrato socioeconómico','Capítulo 6 — El mapa de la segunda vuelta: dónde se gana','Capítulo 7 — La reserva dormida: geografía de la abstención','Capítulo 8 — Tres campañas para ganar la 2ª vuelta (centro · recuperación · abstención)','Nota metodológica e índice de soportes (Excel)']:
    pp=d.add_paragraph(); pp.paragraph_format.space_after=Pt(2); r=pp.add_run('· '+t); r.font.size=Pt(10.5); r.font.color.rgb=INK

# RESUMEN EJECUTIVO
h('Resumen ejecutivo',size=17,color=OX,before=16)
body('**1. No fue una caída de la izquierda, fue una consolidación de la derecha.** Abelardo (43,7%) y Cepeda (40,9%) van a segunda vuelta el 21 de junio. La derecha (Abelardo + Paloma) suma **51,1%**, casi idéntico al 51,9% de Fico + Rodolfo en la 1ª vuelta de 2022: **el bloque no creció, dejó de estar dividido**. El primer lugar de Petro en 2022 fue el espejismo de una derecha partida.')
body('**2. Cepeda igualó el techo de Petro a nivel nacional, pero la geografía se movió.** Creció frente a Petro-1V en la periferia y la frontera (Arauca +10, Santander +8, Amazonía) y **cayó en las grandes ciudades** (Bogotá, Barranquilla, Cali). La izquierda ganó territorio y perdió donde están los votos.')
body('**3. Bogotá es el campo de batalla — y la oportunidad.** La izquierda se erosionó ~4–5 puntos frente a Petro-1V y la derecha avanzó +6,5 ahí (el espejo exacto). Pero el **techo de la izquierda en Bogotá es 58,6%** (Petro en 2ª vuelta 2022): Cepeda está hoy en ~43%, con **~16 puntos de recuperación**. Igual en Barranquilla (+16) y Cali (+12); en Medellín el techo es estructuralmente bajo (~34%).')
body('**4. La izquierda crece pero a veces se queda sola.** En 187 municipios —sobre todo en Santander, Boyacá y Chocó— Cepeda subió su porcentaje mientras caía la participación: consolida una porción mayor de un electorado que se encoge. Su mayor reserva no es convencer, es **movilizar**.')
body('**5. El centro es urbano y bogotano, y votó para el otro lado.** Fajardo, Claudia y el voto de Oviedo se concentran en Bogotá, Antioquia y el Eje Cafetero. Cruzado **mesa a mesa**, apenas **~8% del voto de Oviedo** —el 2º de la consulta de la derecha— se fue a Paloma; el 87% se volcó a la izquierda y el centro (~67% Cepeda, ~20% Fajardo). Oviedo no fue el salvavidas de Paloma; su electorado de centro sigue en disputa.')
body('**6. La ruta de Cepeda pasa por las ciudades, no por el mapa.** Tres palancas: recuperar las grandes ciudades (techo alto), bajar la abstención en sus bastiones (la Costa y el sur de Bogotá), y disputar el centro anti-Abelardo. La derecha parte con la tarea más simple: que el antipetrismo no se vuelva a dividir.')
body('**7. Y dentro de Bogotá, el blanco es la clase media.** Cruzamos los 1.038 puestos de la ciudad con la estratificación socioeconómica oficial por manzana (SDP/IDECA): el voto se ordena casi perfectamente por estrato. Cepeda gana en estratos 1 y 2 (62% y 55%), **empata técnicamente en estrato 3** (40,9% contra 40,2% de Abelardo) y se desploma del 4 al 6. La caída frente a Petro-1V se concentra en los **estratos 3 y 4 (−4,7 y −6,3 puntos) —la clase media—**, no en los pobres (1-2, casi intactos) ni en los ricos. Ahí, en el estrato medio del occidente y el centro, está la bolsa de recuperación más grande y más alcanzable.')
body(f'**8. La cuenta de la segunda vuelta.** Si la derecha se une detrás de Abelardo, su piso —solo con votos ya emitidos— es de **~{NF(M2V["abe_floor"])}**; Cepeda, con su base más **todo** el centro, llega a **~{NF(M2V["cep_ceiling"])}**. Faltan **~1,9 millones** — y no son hipotéticos: son casi exactamente los votos que la izquierda ya tuvo con Petro en la 2ª vuelta de 2022 y todavía no recupera, que dejamos mapeados ciudad por ciudad y puesto por puesto. La elección no se gana solo persuadiendo al centro: se gana **recuperando ese techo en las grandes ciudades y movilizando los bastiones** que hoy se abstienen.')

# CAP 1
cap(1,'Petro 2022: ¿dónde creció de primera a segunda vuelta?')
body('Entre la primera y la segunda vuelta de 2022, Petro pasó de **40,4% a 50,4% (+10 puntos)**, concentrando el crecimiento en el Caribe, Bogotá, el Pacífico y el Eje Cafetero, y mínimo en los Santanderes.')
b1=sorted([(k,v) for k,v in BA['depto'].items() if v['dif_petro2v1v'] is not None and k!='Exterior'],key=lambda x:-x[1]['dif_petro2v1v'])
tbl(['Departamento','Petro 1V %','Petro 2V %','Creció (pp)'],[[k,f"{v['petro1v']:.1f}",f"{v['petro2v']:.1f}",f"+{v['dif_petro2v1v']:.1f}"] for k,v in b1[:7]])
h('En las cuatro grandes ciudades')
tbl(['Ciudad','Petro 1V %','Petro 2V %','Creció (pp)'],[['Barranquilla','51,4','63,4','+12,0'],['Bogotá','47,1','58,6','+11,5'],['Cali','53,4','63,9','+10,5'],['Medellín','24,4','33,7','+9,3']])
body('**El hecho — la "alianza" de 2022:** en segunda vuelta Petro absorbió al **centro** (los ~4 millones de Fajardo) y al voto **anti-Rodolfo**. Por eso disparó donde había una bolsa urbana y costeña (Atlántico +12, Bogotá +11,5, Valle +10,5) y casi no se movió en los Santanderes (+5), tierra dura del antipetrismo. La lección para 2026: **el crecimiento de la izquierda en segunda vuelta vive en las ciudades y en el voto de centro, no en su mapa rural.** (Detalle por municipio en la hoja "1·Petro 2V-1V (municipio)" del Excel.)',after=10)

# CAP 2
cap(2,'Cepeda 2026 frente a Petro 2022')
body('A nivel nacional Cepeda igualó el techo de Petro (~41%). El diferencial por territorio revela el problema: **creció en la periferia y cayó en las grandes ciudades.**')
h('El mapa del resultado')
body('El resultado dibuja dos países. Cepeda ganó la **periferia y la frontera** —toda la Costa Caribe, el Pacífico, el sur (Nariño, Cauca) y la Orinoquía-Amazonía—; Abelardo ganó el **corazón andino**: Antioquia, el Eje Cafetero, los Santanderes, Boyacá, Cundinamarca, Tolima, Huila. La izquierda pinta más mapa; la derecha pinta donde vive más gente.')
img('m_ganador_dep.png',4.7)
b2=sorted([(k,v) for k,v in BA['depto'].items() if v['dif_cep'] is not None and k!='Exterior'],key=lambda x:x[1]['dif_cep'])
tbl(['Dónde MÁS cayó la izquierda','Cep 26 %','Petro 1V-22 %','Dif (pp)'],[[k,f"{v['cep26']:.1f}",f"{v['petro1v']:.1f}",f"{v['dif_cep']:+.1f}"] for k,v in b2[:5]])
h('El giro, municipio a municipio')
body('El mismo movimiento, a escala municipal, deja ver el patrón con nitidez: **verde donde la izquierda creció frente a Petro-1V, rojo donde cayó.** El verde domina la Orinoquía, la Amazonía y buena parte de la frontera; el rojo se concentra en focos del sur (Nariño, Cauca) y, sobre todo, en las cabeceras urbanas.')
img('m_swing_mun.png',4.8)
body('Y el nivel absoluto de Cepeda —no el cambio, sino cuánto sacó— confirma el bastión: oscuro (alto) en la Costa, el Pacífico y el sur; claro (bajo) en el centro andino y el Eje Cafetero.')
img('m_cepeda_mun.png',4.8)
mbig=[v for v in BF['muni'].values() if v['dif_cep'] is not None and BIG(v,15000) and v['dep']!='Exterior']
h('Municipios grandes donde más se erosionó (y dónde resistió)')
tbl(['Cayó · Municipio','Dif (pp)','— Creció · Municipio','Dif (pp)'],
    [[a['muni'],f"{a['dif_cep']:+.1f}",b['muni'],f"{b['dif_cep']:+.1f}"] for a,b in zip(sorted(mbig,key=lambda x:x['dif_cep'])[:6],sorted(mbig,key=lambda x:-x['dif_cep'])[:6])])
h('El número estrella no es la caída, es el techo')
body('La pregunta política no es cuánto bajó la izquierda sino cuánto puede subir. La referencia es Petro en la **segunda vuelta** de 2022 — su techo real en cada ciudad:')
img('g_ciudades_techo.png',6.4)
body('Cepeda está muy por debajo de ese techo en las grandes ciudades: **+16 puntos en Barranquilla, +16 en Bogotá, +12 en Cali**. No necesita "conquistar" terreno nuevo; necesita **volver a donde ya estuvo en 2022.**')
body('Dentro de Bogotá, la erosión frente a Petro-1V fue **transversal —cayó en todas las localidades— y más fuerte en las localidades medias del occidente y el centro** (Antonio Nariño, Puente Aranda, Fontibón, Engativá, Kennedy):')
img('g_bogota_recuperacion.png',5.5)
h('Cepeda crece y, a la vez, sube la abstención')
body('En **187 municipios** la izquierda subió su porcentaje mientras caía la participación. No es un patrón nacional —donde votó la gente, Cepeda mobilizó— sino **un fenómeno rural concentrado en Santander, Boyacá y Chocó**: la izquierda consolida una porción mayor de un electorado que se encoge.')
ma=BF['muni_abst'][:6]
tbl(['Municipio','Depto','Cepeda crece (pp)','Participación 2022 a 2026'],[[r['muni'],r['dep'],f"+{r['dif_cep']:.1f}",f"{r['part22']:.0f}% a {r['part26']:.0f}%"] for r in ma])
body('Su mayor reserva, entonces, no es convencer sino **movilizar** a su propia base. (Lista completa en la hoja "2·Crece+abstención (muni)".)')
h('¿Cuánto tiene que crecer, entonces? — la cuenta, en un vistazo')
body(f'Antes de seguir, conviene fijar el tamaño del reto, porque es la brújula de todo lo que viene. Si la derecha se une detrás de Abelardo, su piso de 2ª vuelta —solo con votos ya emitidos— es de **{NF(STRAT["abe_floor"])}**. Cepeda parte de sus **{NF(M2V["votos"]["cepeda"])}** y, para empatar ese piso, necesita **sumar ~{NF(STRAT["need_over_1v"])}** a su primera vuelta:')
img('g_brecha_2v.png',6.6)
body(f'Esos ~{NF(STRAT["need_over_1v"])} salen de **dos bolsas distintas**: **~{NF(STRAT["centro_total"])} del centro** (persuadir a quien ya votó por Fajardo o Claudia) y **~{NF(STRAT["gap"])} de movilización** (recuperar o sacar a votar a quien hoy no está con Cepeda). **El Capítulo 8 desarrolla ese reto en tres estrategias de campaña** —Centro, Recuperación y Abstención—, cada una mapeada hasta el barrio. Los capítulos que siguen (centro, derecha, estratos, abstención) son el diagnóstico que alimenta ese plan.',after=10)

# CAP 3
cap(3,'El voto de centro (Fajardo, Claudia, Oviedo)')
body('El voto de centro es **urbano, paisa y bogotano**. Fajardo (4,3% nacional) se concentra en Antioquia (218 mil), Bogotá (252 mil) y el Eje Cafetero; Claudia López (~225 mil, ~1%) refuerza el patrón en Bogotá. Ese es el universo que una segunda vuelta a dos miedos pone en disputa.')
h('El caso Oviedo: ¿conviene una adhesión a Cepeda?')
body('Se plantea la posibilidad de que el voto de Paloma en 1ª vuelta (1,64M) sea en realidad el de **Juan Daniel Oviedo** —2º en la Gran Consulta con 1,26M— que no quiso irse con Abelardo, y se evalúa una eventual adhesión a Cepeda para recuperar Bogotá. **Los datos no respaldan la hipótesis, pero abren una más útil:**')
bullet('**La hipótesis no se sostiene:** el voto de Paloma-1V se parece más a su propia base de consulta (+0,21) que a la de Oviedo (−0,11). Lo que conservó es su núcleo uribista.')
bullet('**Oviedo es un fenómeno bogotano:** ~505 mil votos en Bogotá, 42% de su total nacional.')
bullet('**Es voto de centro, anti-Abelardo:** correlaciona con Fajardo (+0,60) y Cepeda (+0,32) y en contra de Abelardo (−0,41); en Bogotá, +0,55 con Cepeda y −0,61 con Abelardo.')
bullet('**Pero dos tercios ya se perdieron:** el 64% (≈318 mil) está en localidades del norte rico que ya ganó Abelardo; solo 36% (≈180 mil) sigue disputable, en el sur.')
h('La pregunta directa: ¿cuánto del voto de Oviedo votó por Paloma?')
body('Conviene darle la vuelta a la pregunta: si Paloma ganó la Gran Consulta pero se desplomó en 1ª vuelta, ¿no será que el voto que conservó es el de Oviedo, su escolta en esa consulta? Corrimos el mismo cruce **mesa a mesa** sobre el bloque de Oviedo (1,26 millones, 2º en la Gran Consulta) para estimar a dónde se fue su base. La respuesta es lapidaria para esa hipótesis:')
img('g_oviedo_destino.png',6.5)
body('Apenas **~8% del voto de Oviedo se fue a Paloma.** El grueso —**~67% a Cepeda y ~20% a Fajardo**— se volcó a la izquierda y al centro; a la derecha dura (Paloma + Abelardo) se fue solo ~13%. Tiene toda la lógica: Oviedo es un perfil técnico y urbano (exdirector del DANE) y su electorado de centro no migró al uribismo de Paloma sino al centro-izquierda. **El segundo de la consulta de la derecha tenía votantes que se fueron para el otro lado.**')
body('Es coherente con las correlaciones por puesto: donde Oviedo fue fuerte, Fajardo (+0,60) y Cepeda (+0,32) crecieron, mientras Abelardo (−0,41) y la propia Paloma (−0,12) cayeron; en Bogotá el rechazo es aún más nítido (correlación Oviedo–Paloma −0,57). En una frase: **Oviedo no fue el salvavidas de Paloma.**')
body('Salvedad metodológica honesta: la inferencia ecológica no fija la transferencia a un punto exacto —la cota dura admite hasta un techo teórico de ~64% hacia Paloma—, pero tanto la estimación como el perfil ideológico del votante apuntan en la misma dirección.',size=9,after=9)
img('g_oviedo_bogota_localidad.png',6.3)
body('**Veredicto:** la adhesión de Oviedo **no mueve la aguja nacional** y dentro de Bogotá **vale por la mitad de su titular** (~180 mil disputables). No cierra sola la brecha, pero esos votos caen donde se decide la ciudad y, con menos abstención en el sur, dibujan el camino hacia el techo del 58,6%.',after=10)

# CAP 4
cap(4,'La derecha hoy (Abelardo + Paloma): ¿dónde avanzó?')
body('La derecha presidencial vale hoy **51,1%** (Abelardo 43,7 + Paloma 6,9), casi lo mismo que el **51,9%** de Fico + Rodolfo en la 1ª vuelta de 2022. **No creció: se consolidó.**')
body('El mapa de su fuerza es el negativo del de Cepeda: intensa en el bloque andino —Eje Cafetero, Santanderes, Boyacá, Cundinamarca, Antioquia, Tolima, Huila— y pálida en la Costa y el Pacífico.')
img('m_derecha_dep.png',4.7)
b4=sorted([(k,v) for k,v in BA['depto'].items() if v['dif_der'] is not None and k!='Exterior'],key=lambda x:-x[1]['dif_der'])
tbl(['Dónde AVANZÓ la derecha','Der 26 %','Fico+Rodolfo 22 %','Avance (pp)'],[[k,f"{v['der26']:.1f}",f"{v['der22']:.1f}",f"+{v['dif_der']:.1f}"] for k,v in b4[:6]])
mder=[v for v in BF['muni'].values() if v['dif_der'] is not None and BIG(v,15000) and v['dep']!='Exterior']
h('Municipios grandes donde más avanzó la derecha')
tbl(['Municipio','Departamento','Der 26 %','Avance (pp)'],[[v['muni'],v['dep'],f"{v['der26']:.1f}",f"+{v['dif_der']:.1f}"] for v in sorted(mder,key=lambda x:-x['dif_der'])[:6]])
body('**El avance de la derecha es el espejo exacto de la caída de la izquierda:** creció donde Cepeda perdió —**Bogotá +6,5**, Atlántico +3,1, las ciudades del Pacífico— y retrocedió donde la izquierda subió (Santander −12, Arauca −13). En una frase: **la segunda vuelta se juega en las grandes ciudades, y hoy ahí la derecha llega con ventaja pero la izquierda con un techo histórico muy por encima de su piso actual.** (Detalle por municipio y comuna en el Excel.)',after=10)

# CAP 5 — BOGOTÁ POR ESTRATO
cap(5,'Bogotá por estrato socioeconómico: dónde está el voto recuperable')
body('Para precisar **dónde** dentro de Bogotá está el voto que la izquierda perdió y puede recuperar, cruzamos los **1.038 puestos** de la ciudad (georreferenciados) con la **estratificación socioeconómica oficial por manzana** de la Secretaría Distrital de Planeación (44.260 manzanas, vigente a octubre de 2025), mediante un cruce geográfico punto-en-polígono. Cada puesto queda etiquetado con el estrato de su entorno. El resultado es el hallazgo más limpio del análisis: **el voto bogotano se ordena casi perfectamente por estrato.**')
body('Así se ve la ciudad por estrato (orientada con el norte a la izquierda, como se acostumbra ver Bogotá) — el estrato 6 (azul oscuro) en el nororiente (Chapinero, Usaquén), los estratos 1 y 2 (rojos) en el sur y el borde occidental, la clase media (estratos 3-4) llenando el centro:')
img('m_bogota_estrato.png',6.4)
ES={e:EST['estratos'][str(e)] for e in range(1,7)}
tot=sum(ES[e]['censo'] for e in range(1,7)) or 1
tbl(['Estrato','Censo (miles)','% ciudad','Cepeda 26 %','Abelardo 26 %','Dif vs Petro 1V','Techo (Petro 2V)'],
    [[ES[e]['nombre'],f"{ES[e]['censo']/1000:.0f}",f"{ES[e]['censo']/tot*100:.0f}%",f"{ES[e]['cepeda']:.1f}",f"{ES[e]['abelardo']:.1f}",f"{ES[e]['dif_cep_pet']:+.1f}",f"+{ES[e]['techo_2v']:.1f}"] for e in range(1,7)])
img('g_bogota_estrato.png',6.6)
body('El mapa del voto refleja el de los estratos casi exactamente. Visto por UPL (las 33 Unidades de Planeamiento Local, más fino que las 20 localidades), Cepeda gana el **sur** (Usme, Lucero, Arborizadora, la cuenca del Tunjuelo) y pierde el **norte rico** (Usaquén, Chapinero, Toberín). El centro y el occidente —la clase media— son la tierra de nadie que decide la ciudad:')
img('m_bogota_upl_cepeda.png',6.4)
h('Tres lecturas políticas')
bullet('**El empate se define en el estrato 3 (medio-bajo).** Es el estrato más grande de la ciudad (2,4 millones de censo, 40% de Bogotá) y ahí Cepeda y Abelardo están en un empate técnico (40,9 vs 40,2). Quien gane el estrato 3 gana Bogotá. Es el terreno decisivo.')
bullet('**La izquierda no perdió en los pobres: perdió en la clase media.** En estratos 1 y 2 Cepeda casi igualó a Petro (−2,2 y −3,4). El boquete está en los **estratos 3 y 4 (−4,7 y −6,3)**: la clase media del occidente y el centro fue la que se corrió. No es un problema de base popular, es un problema de clase media urbana.')
bullet('**El estrato 4 es la mejor oportunidad, pero con una traba: la abstención.** Es donde más cayó la izquierda (−6,3) y donde el techo recuperable es alto (+17,5 hasta el 44,8% que sacó Petro en 2ª vuelta), pero también es el estrato que **menos participó** (59% frente al 68-69% del resto). Recuperar el estrato 4 es a la vez convencer y movilizar.')
body('**En síntesis:** la pelea de Bogotá no es por los barrios populares —que la izquierda mayormente conserva— sino por el cinturón de clase media (estratos 3-4) del occidente y el centro de la ciudad. Es la misma conclusión del Capítulo 2 (las localidades medias del occidente), ahora con el dato socioeconómico que la explica. (Detalle por estrato en la hoja "Bogotá·estrato" del Excel.)',after=8)
body('Nota de lectura: el estrato asignado a cada puesto es el de la manzana donde se ubica, usado como aproximación del entorno socioeconómico de su electorado; los votantes de un puesto provienen de manzanas vecinas que pueden mezclar estratos. La asignación es muy precisa en lo espacial (distancia mediana puesto-manzana de 2,3 m) y el gradiente es tan marcado que la conclusión es robusta a esa aproximación.',size=9,after=6)

# CAP 6 — MAPA DE LA SEGUNDA VUELTA
cap(6,'El mapa de la segunda vuelta: dónde se gana el 21 de junio')
body('Hasta aquí el diagnóstico de lo que pasó. Este capítulo lo gira hacia adelante: **dónde están, en el mapa, los votos que Cepeda necesita para ganar la segunda vuelta.** La vara no es lo que sacó el 31 de mayo, sino lo que la izquierda ya logró en su mejor noche reciente —Petro en la 2ª vuelta de 2022—. La diferencia entre ese techo y el piso de hoy es el espacio por recuperar.')
body('El mapa pinta ese margen disponible por departamento. Cuanto más verde, más espacio entre el techo de 2022 y Cepeda hoy: la Costa Caribe, el Valle, Bogotá y el Pacífico concentran la mayor reserva.')
img('m_recuperacion_dep.png',4.7)
h('Las prioridades, en votos y no en puntos')
body('Un punto porcentual no vale lo mismo en Bogotá que en Vaupés. Si se pondera el espacio por recuperar por el tamaño del electorado, **el orden de prioridad es inequívocamente urbano y costeño:**')
rec=[]
for k,v in BA['depto'].items():
    if k=='Exterior': continue
    if v.get('petro2v') and v.get('cep26') and v.get('base_votos'):
        esp=v['petro2v']-v['cep26']; votos=esp/100*v['base_votos']
        if esp>0: rec.append((k,v['cep26'],v['petro2v'],esp,votos))
rec.sort(key=lambda x:-x[4])
tbl(['Prioridad 2V · Departamento','Cepeda hoy %','Techo (Petro 2V) %','Espacio (pp)','Votos por recuperar (aprox.)'],
    [[k,f"{c:.1f}",f"{p2:.1f}",f"+{esp:.1f}",f"{votos:,.0f}".replace(',','.')] for k,c,p2,esp,votos in rec[:9]])
body('Estos nueve territorios concentran la mayor parte del voto recuperable del país. Tres de ellos —**Bogotá, Atlántico y Valle**— son, además, donde el techo histórico está más por encima del piso actual: la elección se reconquista en las urbes, no en el mapa rural que la izquierda ya domina.')
h('La ecuación de la segunda vuelta')
bullet('**Recuperar las grandes ciudades** (techo alto, masa grande): Bogotá, Barranquilla, Cali, Cartagena. Es el frente decisivo.')
bullet('**Movilizar los bastiones** que ya son de Cepeda pero votan poco (la Costa, el sur de Bogotá): no hay que convencer, hay que sacar a votar — es el Capítulo 7.')
bullet('**Disputar el centro anti-Abelardo** (Fajardo, Claudia, Oviedo): un universo urbano de ~180-250 mil votos disputables, sobre todo en Bogotá y Antioquia.')
body('La derecha llega a la segunda vuelta con la tarea más simple —que el antipetrismo no se vuelva a dividir— pero con menos margen de crecimiento: su techo y su piso casi coinciden. La izquierda llega con la tarea más difícil pero con un **techo histórico muy por encima de su piso actual.** Ahí está la partida.',after=10)

# CAP 7 — GEOGRAFÍA DE LA ABSTENCIÓN
cap(7,'La reserva dormida: la geografía de la abstención')
body('Si el Capítulo 6 dice dónde están los votos por convencer, este dice dónde están los votos por **despertar**. La participación de la 1ª vuelta fue del 58% — y no se repartió parejo: votó mucho más el bloque andino de la derecha que la Costa y el sur de la izquierda.')
body('El mapa de la abstención municipal lo muestra: más oscuro = más gente que se quedó en casa. La mancha intensa cubre justamente la **Costa Caribe y el Pacífico** —territorio de Cepeda— mientras Boyacá, Cundinamarca, Nariño y el Eje andino votaron más.')
img('m_abstencion_mun.png',4.8)
h('La asimetría que favorece a Cepeda')
body('Esa desigualdad es, paradójicamente, una buena noticia para la izquierda: **su voto duerme más que el de la derecha.** Donde Cepeda es fuerte, la participación fue baja; donde Abelardo es fuerte, ya votaron casi todos. El margen de Cepeda para crecer movilizando es estructuralmente mayor. Estos son los municipios donde esa reserva es más grande —alta votación por Cepeda y baja participación:')
mob=[v for v in BF['muni'].values() if v['dep']!='Exterior' and v.get('cep26') and v.get('part26') and v['cep26']>=50 and v['part26']<=53 and v.get('base',0)>=8000]
mob.sort(key=lambda x:-x['base'])
if len(mob)<6:
    mob=[v for v in BF['muni'].values() if v['dep']!='Exterior' and v.get('cep26') and v.get('part26') and v['cep26']>=48 and v['part26']<=56 and v.get('base',0)>=6000]
    mob.sort(key=lambda x:-x['base'])
tbl(['Municipio','Departamento','Cepeda %','Participación %','Votos 26'],
    [[v['muni'],v['dep'],f"{v['cep26']:.1f}",f"{v['part26']:.0f}",f"{v['base']:,}".replace(',','.')] for v in mob[:9]])
body('Cada punto de participación que suba en estos municipios y en el sur de Bogotá entra mayoritariamente a la cuenta de Cepeda. La segunda vuelta de la izquierda se juega tanto en la maquinaria de movilización de sus bastiones como en la persuasión de las ciudades. (Participación por municipio en las hojas del Excel.)',after=10)

# CAP 8 — CUÁNTOS VOTOS NECESITA CEPEDA
cap(8,'Tres campañas para ganar la segunda vuelta')
body('La pregunta de fondo, hecha número — y convertida en plan operativo. Modelamos la 2ª vuelta con **supuestos de trasvase explícitos** —y conservadores— por bloque ideológico: Paloma endosa a Abelardo (85%), los minoritarios de derecha van a Abelardo (78%), el centro se reparte (Fajardo 55% Cepeda / 30% Abelardo; Claudia 65% / 20%) y los minoritarios de izquierda van a Cepeda (85%). Son cifras movibles; lo que importa es el orden de magnitud.')
h('La cuenta de cobro: la derecha ya se consolidó')
body(f'Si la derecha se une detrás de Abelardo, su piso de 2ª vuelta —solo con votos ya emitidos— es de **{NF(STRAT["abe_floor"])} de votos**. Cepeda parte de sus {NF(M2V["votos"]["cepeda"])} de 1ª vuelta y, **capturando todo el centro disponible**, llega a un techo de **{NF(STRAT["cep_ceiling"])}**. La diferencia es lo que hay que construir:')
img('g_trasvase_2v.png',6.6)
tbl(['Concepto','Votos'],[
    ['Piso de Abelardo (derecha consolidada)',NF(STRAT['abe_floor'])],
    ['Cepeda con su base + todo el centro',NF(STRAT['cep_ceiling'])],
    ['Brecha por cerrar',NF(STRAT['gap'])],
    ['Cepeda debe sumar a su 1ª vuelta (para empatar el piso)',NF(STRAT['need_over_1v'])],
    ['  · campaña de CENTRO (Fajardo + Claudia, ya votaron)',NF(STRAT['centro_total'])],
    ['  · campaña de MOVILIZACIÓN (recuperación / abstención)',NF(STRAT['gap'])]])
body(f'En una frase: Cepeda necesita sumar **~{NF(STRAT["need_over_1v"])}** a su primera vuelta, y ese total viene de **dos bolsas distintas**: **~{NF(STRAT["centro_total"])} del centro** (persuadir a quien ya votó por Fajardo o Claudia) y **~{NF(STRAT["gap"])} de movilización** (sacar a votar o recuperar a quien hoy no está con Cepeda). De ahí salen **tres campañas**, cada una con su mapa, su Excel y su público. **La clave para no enredarse es cómo se suman:**')
bullet(f'**Centro (≈{NF(STRAT["centro_total"])}) SÍ se suma a movilización (≈{NF(STRAT["gap"])}).** Son personas distintas: una persuade a quien votó centro, la otra moviliza a quien no votó. Juntas dan los ~{NF(STRAT["need_over_1v"])}.')
bullet('**Recuperación y Abstención NO se suman entre sí.** Son el **mismo universo** —la coalición que votó a Petro en 2022 y hoy no está con Cepeda— mirado de dos formas: Recuperación lo **dimensiona** (el techo: cuánto y dónde), Abstención dice **cómo recuperarlo** (subiendo participación donde ya ganamos la 2ª vuelta). Sumar las dos sería contar el mismo voto dos veces.')

# ── helper: bloque por ciudad (mapa numerado + análisis que desarrolla los barrios numerados) ──
BLAB=json.load(open(f'{OUT}/barrio_labels.json')); C9=STRAT['city9']
_MF={'centro':('centro','voto de centro transferible','votos'),'rec':('recuperar','voto por recuperar','votos'),'abst':('abstencion','voto neto por movilización','netos')}
_MK={'centro':'centro','rec':'recuperar','abst':'abst_neto'}
GEO={'bogota':('','el sur y el suroccidente populares'),   # pista solo para movilización (big 5, geografía conocida); centro se deja a los barrios numerados
 'medellin':('','el centro-oriente y el occidente populares'),
 'cali':('','el oriente popular: el Distrito de Aguablanca'),
 'barranquilla':('','el suroccidente popular'),
 'cartagena':('','el suroriente popular'),
 'manizales':('',''),'pereira':('',''),'bucaramanga':('',''),'cucuta':('','')}
def _topref(top,k=4): return ', '.join(f"{it['barrio']} ({it['n']})" for it in top[:k] if it['barrio'])
def bcity(slug,code,metric,w=5.4,note=''):
    mf,noun,unit=_MF[metric]; lab=BLAB.get(f'm_{slug}_{mf}_barrio.png',{}); top=lab.get('top',[])
    muni=C9[code]; mt=muni[_MK[metric]]; gc=GEO.get(slug,('',''))[0 if metric=='centro' else 1]
    img(f'm_{slug}_{mf}_barrio.png',w)
    p1=f"**{muni['name']}** aporta **{NF(mt)}** de {noun} a nivel municipio. Dentro de la ciudad, las mayores bolsas —numeradas en el mapa— son **{_topref(top)}**"
    if top: p1+=f"; encabeza {top[0]['barrio']} (1) con cerca de {NF(top[0]['val'])} {unit}"
    p1+=(f", concentradas en {gc}." if gc else ".")
    body(p1)
    if metric=='centro': p2='Es voto de **persuasión**: gente que ya salió a votar, pero por Fajardo o Claudia; ganarlos no amplía el electorado, lo reordena.'
    elif metric=='rec': p2='Es el techo que la izquierda **ya alcanzó con Petro en 2022** y todavía no recupera: votos que existieron, no proyecciones.'
    else: p2=f"La apuesta aquí es **movilización**: subir la participación rinde neto porque la izquierda gana la 2ª vuelta de la ciudad (share {muni['abst_share']:.0f}%)."
    body(p2+((' '+note) if note else ''),after=9)
EXCELN='Para cruzar con el Excel: el mapa asigna cada puesto al barrio que lo contiene geográficamente; la hoja del Excel lo agrupa por el barrio declarado del puesto, así que las cifras por barrio difieren un poco entre sí — el número que manda es siempre el total municipal.'
FILLN='Los barrios sin puesto de votación propio se colorean con la tendencia del vecino más cercano para no dejar huecos; el texto solo numera y desarrolla los de dato directo.'

h('Campaña 1 · Centro — los ~700.000 votos que se suman')
body(f'Es una campaña de **persuasión** a gente que ya fue a votar, pero por el centro. Con los supuestos del modelo —55% de Fajardo y 65% de Claudia— el centro aporta **{NF(STRAT["centro_total"])} votos** transferibles. Es un voto **urbano**: las grandes ciudades concentran la mitad y Bogotá sola, casi un tercio.')
tbl(['Municipio','Depto','Fajardo','Claudia','Centro transferible'],
    [[r['muni'],r['dep'],NF(r['faj']),NF(r['cla']),NF(r['tr'])] for r in STRAT['centro_top']])
body('Ciudad por ciudad, dónde está ese voto de centro **hasta el barrio**. En cada mapa las mayores bolsas van **numeradas** y se desarrollan en el texto; el color es ámbar (centro). *Claudia se estima a nivel de barrio repartiéndola en proporción a Fajardo, ya que en el preconteo solo existe a municipio.*',after=8)
bcity('bogota','16001','centro',5.8,note=EXCELN+' '+FILLN)
bcity('medellin','01001','centro',4.6)
bcity('cali','31001','centro',5.0)
bcity('barranquilla','03001','centro',4.6)
bcity('cartagena','05001','centro',5.0)
bcity('manizales','09001','centro',5.8)
bcity('pereira','24001','centro',5.6)
bcity('bucaramanga','27001','centro',4.4)
bcity('cucuta','25001','centro',5.0)
body('El detalle completo va en **Estrategia_1_Centro.xlsx**: el voto de centro transferible por **municipio** en todo el país (con acumulado) y por **localidad, comuna y barrio** de las ciudades principales.',after=10)

h('Campaña 2 · Recuperación — el techo de la izquierda (~2,0 millones)')
body('Es la campaña para la gente que se fue **por desilusión** o por los acuerdos de la segunda vuelta pasada. La vara no es lo que sacó Cepeda el 31 de mayo, sino lo que la izquierda ya logró con **Petro en la 2ª vuelta de 2022**. La diferencia entre ese techo y el piso de hoy es el voto por recuperar.')
body(f'No es hipotético: **ya existió**. Sumado municipio por municipio da **~{NF(STRAT["recuperar_total"])} de votos**. Se concentra en las grandes ciudades:')
tbl(['Municipio','Depto','Cepeda 1V','Techo (Petro 2V)','Votos por recuperar'],
    [[r['muni'],r['dep'],NF(r['cep_now']),NF(r['techo']),NF(r['recuperar'])] for r in STRAT['rec_top']])
body('El mismo recorrido por barrio (color verde · recuperación). Bogotá concentra un tercio de todo el voto por recuperar:',after=8)
bcity('bogota','16001','rec',5.8,note=FILLN)
bcity('cali','31001','rec',5.0)
bcity('medellin','01001','rec',4.6)
bcity('barranquilla','03001','rec',4.6)
bcity('cartagena','05001','rec',5.0)
bcity('manizales','09001','rec',5.8)
bcity('pereira','24001','rec',5.6)
bcity('cucuta','25001','rec',5.0)
body(f'(Bucaramanga no aparece aquí: su techo de izquierda es tan bajo —Petro sacó 25% en 2ª vuelta 2022— que el voto por recuperar a nivel municipio es prácticamente cero; su fuerte es el centro.) El detalle va en **Estrategia_2_Recuperacion.xlsx**: el voto por recuperar en **{NF(len(TER["muni"]))} municipios** y **{NF(len(TER["puesto"]))} puestos**, y por **comuna y barrio** de las ciudades principales.',after=10)

h('Campaña 3 · Abstención — movilizar donde somos fuertes (~1,9 millones)')
body('Es **movilización** (sacar a votar), no persuasión — y **solo donde la izquierda gana de verdad**. En segunda vuelta cada nuevo votante alimenta a los dos bloques, así que movilizar solo rinde **neto** donde ya ganamos cabeza a cabeza. Medimos el voto neto de cada abstencionista: abstención × (2 × share − 1), con el share de Petro en la 2ª vuelta de 2022. Donde la izquierda perdió ese número es negativo y el municipio no entra — por eso esta campaña vive en los bastiones de la Costa, el Pacífico y el sur de Bogotá, no en el centro andino.')
body(f'Con esa regla hay **~{NF(STRAT["abst_total_neto"])} de voto neto disponible**. **Respuesta directa a la pregunta:** los **{NF(STRAT["abst_target"])}** de la brecha se juntan con los **primeros {STRAT["abst_n_muni"]} municipios** —bastiones de la Costa, el Pacífico y el sur de Bogotá—:')
tbl(['Municipio','Depto','Censo','Abstención','Share Petro 2V','Neto Cepeda'],
    [[r['muni'],r['dep'],NF(r['censo']),NF(r['ab']),f"{r['sh']:.0f}%",NF(r['neto'])] for r in STRAT['abst_top']])
body('Y dónde, dentro de cada bastión, por barrio (color cobre · abstención). Las mayores bolsas van numeradas y se desarrollan en el texto:',after=8)
for _slug,_code in [('bogota','16001'),('cali','31001'),('cartagena','05001'),('barranquilla','03001'),
                    ('buenaventura','31019'),('soledad','03052'),('pasto','23001'),('santamarta','21001'),
                    ('palmira','31079'),('quibdo','17001'),('tumaco','23139'),('sincelejo','28001')]:
    bcity(_slug,_code,'abst',5.2,note=(FILLN if _slug=='bogota' else ''))
body('**Las otras cinco ciudades del informe —Medellín, Manizales, Pereira, Bucaramanga y Cúcuta— no entran a esta campaña.** La izquierda no gana su 2ª vuelta en ninguna (share de Petro-2V por debajo de 50%: Bucaramanga 25%, Cúcuta 18%, Medellín 34%, Pereira/Manizales ~47%), así que su voto neto de movilización es cero o marginal — sacar abstención ahí le sumaría a Abelardo, no a Cepeda. Su aporte está en Centro y Recuperación, ya mapeados.',after=8)
body(f'El detalle va en **Estrategia_3_Abstencion.xlsx**: los municipios ordenados por voto neto, con la columna que marca los {STRAT["abst_n_muni"]} que suman el objetivo, más el desglose por localidad/comuna/barrio. **En síntesis: la 2ª vuelta de Cepeda no se gana en el mapa rural que ya domina, sino sumando el centro y movilizando, ciudad por ciudad y barrio por barrio, el voto que fue de la izquierda en 2022.**',after=10)

# MÉTODO + índice soportes
cap('M','Nota metodológica e índice de soportes')
body('**Fuentes:** preconteo Registraduría 2026 a nivel de mesa (~99,9% escrutado), agregado a puesto y unido al georreferenciado nacional de puestos (barrio, comuna/localidad, potencial); escrutinios 2022 de primera y segunda vuelta (GCS) a nivel de mesa; Gran Consulta del 8-mar-2026 a nivel de puesto. Proporciones sobre **votos válidos + voto en blanco** para comparar 2022 y 2026 (un cambio de base explica diferencias de ±1 punto frente a porcentajes publicados que excluyen el blanco). Voto de Claudia López a nivel de municipio (no disponible a mesa en el preconteo). Correlaciones de Pearson sobre proporciones por puesto. **Estratos de Bogotá:** cruce geográfico punto-en-polígono de los 1.038 puestos georreferenciados contra la capa oficial de estratificación por manzana de la Secretaría Distrital de Planeación / IDECA (44.260 manzanas, oct-2025); el diferencial 2026 vs 2022 se calcula sobre los 892 puestos apareados entre ambos años. "Techo" = mejor resultado reciente de la izquierda (Petro 2ª vuelta 2022); "disputable" no implica transferencia automática. **Cartografía:** coropletas propias sobre los GeoJSON oficiales de departamentos y municipios (límites electorales) y la capa de localidades de Bogotá; el "voto por recuperar" del Capítulo 6 estima el espacio (techo 2V − Cepeda hoy) por el tamaño del electorado de cada departamento, como orden de prioridad, no como pronóstico. Cifras de preconteo, sujetas a escrutinio definitivo.',size=9)
body('**El archivo Excel de soporte ("Soporte_Analisis_Pacto_1V_2026.xlsx") contiene el detalle completo en 15 hojas:** Resumen · Petro 2V-1V (depto y municipio) · Cepeda vs Petro (depto, municipio, Bogotá localidad, crece+abstención) · Centro · Derecha (depto y municipio) · Ciudades por comuna · Ciudades por barrio (Bogotá, Medellín, Cali y Barranquilla) · Bogotá por estrato socioeconómico · **Camino 2ª vuelta por municipio y por puesto de votación** (votos por recuperar hasta el techo). Todas las hojas con autofiltros y formato condicional.',size=9)
body('**Mapas por barrio (Capítulo 8):** cada puesto de votación georreferenciado se cruza por punto-en-polígono con la capa oficial de barrios de cada ciudad (catastro/planeación). Bogotá, Medellín, Cali, Barranquilla, Cartagena, Manizales, Pereira, Bucaramanga, Cúcuta, Buenaventura y Soledad usan la capa oficial de su catastro o planeación municipal (varias obtenidas de los GeoServer de las alcaldías); un grupo menor de ciudades sin capa pública usa polígonos aproximados (celdas de puesto disueltas por barrio), señalado en el pie de cada mapa. Las mayores bolsas van **numeradas** en el mapa y se desarrollan en el texto. Los barrios sin puesto propio heredan la tendencia del vecino más cercano. Por el agrupamiento geográfico, las cifras por barrio del mapa pueden diferir levemente de las del Excel (que agrupa por el barrio declarado del puesto); el ancla es siempre el total municipal. Centro y recuperación cubren las grandes ciudades; abstención, las ciudades donde la izquierda gana la 2ª vuelta (neto positivo).',size=9)
body('**Las tres campañas del Capítulo 8 tienen además su propio Excel operativo, uno por estrategia:** **"Estrategia_1_Centro.xlsx"** (voto de centro transferible por municipio, localidad/barrio de Bogotá y comuna/barrio de las 16 ciudades) · **"Estrategia_2_Recuperacion.xlsx"** (voto por recuperar hasta el techo, por municipio, puesto y barrio) · **"Estrategia_3_Abstencion.xlsx"** (voto neto por movilización de abstención, por municipio —con la columna que marca los que suman el objetivo— y por comuna/localidad). Cada archivo abre con una hoja "Léeme" que explica la métrica, cómo se suma y cómo encaja con las otras dos campañas.',size=9)
body('**Modelo de 2ª vuelta (Capítulo 8):** supuestos de trasvase declarados (Paloma 85% a Abelardo; Fajardo 55% Cepeda / 30% Abelardo; Claudia 65% / 20%; minoritarios por bloque). "Piso" de la derecha = suma de votos ya emitidos con esos trasvases (no incluye nueva movilización, que alimenta a ambos bloques). "Techo" y "voto por recuperar" = mejor marca de la izquierda (Petro 2ª vuelta 2022). Es un escenario para dimensionar el esfuerzo, no un pronóstico.',size=9)
DOCX=f'{OUT}/Analisis_Nacional_Electoral_Pacto_1V_2026.docx'
d.save(DOCX); print('✓ Word consolidado (capítulos)')

def embed_inter(docx_path):
    """Incrusta Inter (Regular/Bold/Italic) en el .docx con obfuscación OOXML
    para que se vea idéntico en cualquier computador (Windows/Mac), sin depender
    de que el cliente tenga la fuente instalada."""
    import zipfile, uuid, re, shutil
    FDIR='tools/pacto-1v-2026/fonts'
    FONTS=[('embedRegular',f'{FDIR}/Inter-Regular.ttf'),
           ('embedBold',   f'{FDIR}/Inter-Bold.ttf'),
           ('embedItalic', f'{FDIR}/Inter-Italic.ttf')]
    def obf(data,guid):
        key=bytes.fromhex(guid.strip('{}').replace('-',''))[::-1]   # 16 bytes, invertidos
        b=bytearray(data)
        for i in range(32): b[i]^=key[i%16]                          # XOR a los primeros 32 bytes
        return bytes(b)
    zin=zipfile.ZipFile(docx_path,'r'); items={n:zin.read(n) for n in zin.namelist()}; zin.close()
    embeds=[]
    for idx,(kind,path) in enumerate(FONTS,1):
        guid='{'+str(uuid.uuid4()).upper()+'}'
        items[f'word/fonts/font{idx}.odttf']=obf(open(path,'rb').read(),guid)
        embeds.append((kind,f'rIdFont{idx}',f'fonts/font{idx}.odttf',guid))
    childs=''.join(f'<w:{k} r:id="{rid}" w:fontKey="{g}"/>' for k,rid,t,g in embeds)
    node=('<w:font w:name="Inter"><w:charset w:val="00"/><w:family w:val="swiss"/>'
          '<w:pitch w:val="variable"/>'+childs+'</w:font>')
    items['word/fontTable.xml']=items['word/fontTable.xml'].decode('utf-8').replace('</w:fonts>',node+'</w:fonts>').encode('utf-8')
    rels=('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
          '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
          +''.join(f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="{t}"/>' for k,rid,t,g in embeds)
          +'</Relationships>')
    items['word/_rels/fontTable.xml.rels']=rels.encode('utf-8')
    ct=items['[Content_Types].xml'].decode('utf-8')
    if 'Extension="odttf"' not in ct:
        ct=ct.replace('</Types>','<Default Extension="odttf" ContentType="application/vnd.openxmlformats-officedocument.obfuscatedFont"/></Types>')
    items['[Content_Types].xml']=ct.encode('utf-8')
    st=items['word/settings.xml'].decode('utf-8')
    if '<w:embedTrueTypeFonts/>' not in st:
        st=st.replace('<w:proofState','<w:embedTrueTypeFonts/><w:proofState',1) if '<w:proofState' in st \
           else re.sub(r'(<w:settings[^>]*>)',r'\1<w:embedTrueTypeFonts/>',st,count=1)
    items['word/settings.xml']=st.encode('utf-8')
    tmp=docx_path+'.tmp'; zo=zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED)
    for n,data in items.items(): zo.writestr(n,data)
    zo.close(); shutil.move(tmp,docx_path)
    return embeds

_emb=embed_inter(DOCX); print(f'✓ Inter incrustada ({len(_emb)} variantes)')
