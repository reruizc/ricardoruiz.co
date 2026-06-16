# -*- coding: utf-8 -*-
"""Anexo PDF · zoom por ciudad (Bogotá · Barranquilla · Cali · Medellín).
Una página por ciudad con las 3 capas (movilización, % jóvenes, % hombres)
lado a lado + estrategia táctica por edad y sexo + top zonas y barrios.
Reusa los PNGs ya generados por build_barrios.py. Misma identidad del Word
principal (Inter incrustada via embed_inter helper)."""
import os, json, re
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
PACTO = os.path.join(ROOT, "Bases de datos", "output_pacto_1v_2026")
OUTD  = os.path.join(ROOT, "Bases de datos", "output_pauta_2v")

NF = lambda n: f"{int(round(n)):,}".replace(",", ".") if n else "0"

OX, INK, GR = RGBColor(0x8a,0x1e,0x16), RGBColor(0x1a,0x16,0x10), RGBColor(0x5a,0x54,0x48)

# ── 1 · datos ────────────────────────────────────────────────────────────────
terr = json.load(open(os.path.join(PACTO, "twov_territorial.json")))
muni = {m["cod"]: m for m in terr["muni"]}

CITIES = [
    ("bogota", "Bogotá D.C.", "16001",
     "Distrito Capital · más de 3,9 millones de votos en juego",
     "localidad", True),
    ("barranquilla", "Barranquilla", "03001",
     "Capital del Atlántico · puerta del Caribe",
     "localidad", False),
    ("cali", "Cali", "31001",
     "Capital del Valle · epicentro Pacífico de la izquierda",
     "comuna", False),
    ("medellin", "Medellín", "01001",
     "Capital de Antioquia · ciudad clave del rival",
     "comuna", False),
]

# top zonas por ciudad (desde los puestos del territorial)
def top_zonas(code, n=8):
    ag = defaultdict(lambda: {"rec":0, "cep":0, "base":0})
    for p in terr["puesto"]:
        pc = str(p.get("dep","")).zfill(2) + str(p.get("mun","")).zfill(3)
        if pc != code: continue
        c = (p.get("comuna") or "(sin comuna)").strip()
        ag[c]["rec"]  += max(0, p.get("recuperar",0) or 0)
        ag[c]["cep"]  += p.get("cep_now",0) or 0
        ag[c]["base"] += p.get("base",0) or 0
    out = sorted(ag.items(), key=lambda x:-x[1]["rec"])[:n]
    return [(_clean_zone(c), v["rec"], v["cep"], v["base"]) for c,v in out]

def _clean_zone(s):  # quita códigos numéricos al inicio
    return re.sub(r"^\d{2,3}\s*", "", s).replace("LOCALIDAD NO.","Localidad ").replace("LOCALIDAD ","Localidad ").title().replace("Comuna ","Comuna ").strip()

def top_barrios(slug, key, n=6):
    g = json.load(open(os.path.join(OUTD, "barrios", f"{slug}.json")))
    feats = [f["properties"] for f in g["features"]
             if f["properties"].get("rec") is not None and f["properties"].get("f",0)==0]
    if key=="rec": feats.sort(key=lambda x:-x["rec"])
    elif key=="young": feats=[f for f in feats if f.get("young") is not None]; feats.sort(key=lambda x:-x["young"])
    elif key=="men": feats=[f for f in feats if f.get("men") is not None]; feats.sort(key=lambda x:-x["men"])
    elif key=="women": feats=[f for f in feats if f.get("men") is not None]; feats.sort(key=lambda x:x["men"])
    return feats[:n]

# ── 2 · documento ────────────────────────────────────────────────────────────
d = Document()
# márgenes
for sec in d.sections:
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(2.0)
# fuente base Inter
style = d.styles["Normal"]
style.font.name = "Inter"; style.font.size = Pt(10.5); style.font.color.rgb = INK
# docDefaults para tablas
sp = d.styles.element.xpath("w:docDefaults/w:rPrDefault/w:rPr")
if sp:
    rPr = sp[0]
    for c in rPr.findall(qn("w:rFonts")): rPr.remove(c)
    rf = OxmlElement("w:rFonts")
    for a in ("ascii","hAnsi","cs","eastAsia"): rf.set(qn(f"w:{a}"), "Inter")
    rPr.append(rf)

def h_title(t, sz=18, color=OX):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.size = Pt(sz); r.font.bold = True; r.font.color.rgb = color
    return p

def eyebrow(t):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t.upper()); r.font.size = Pt(8.5); r.font.bold = True
    r.font.color.rgb = OX; r.font.name = "Inter"
    rpr = r._element.get_or_add_rPr()
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), "30"); rpr.append(sp)

def body(t, sz=10.5, after=6, color=None):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # *negrita* simple
    parts = re.split(r"(\*\*[^*]+\*\*)", t)
    for s in parts:
        if not s: continue
        if s.startswith("**"):
            r = p.add_run(s[2:-2]); r.font.bold = True
        else:
            r = p.add_run(s)
        r.font.size = Pt(sz)
        if color: r.font.color.rgb = color
    return p

def note(t, sz=8.5):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.size = Pt(sz); r.font.color.rgb = GR; r.italic = True
    return p

def bullet(t):
    p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = re.split(r"(\*\*[^*]+\*\*)", t)
    for s in parts:
        if not s: continue
        if s.startswith("**"):
            r = p.add_run(s[2:-2]); r.font.bold = True
        else:
            r = p.add_run(s)
        r.font.size = Pt(10.5)

def tbl(headers, rows, col_widths=None):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i,hd in enumerate(headers):
        c = t.cell(0,i); c.text = ""; p = c.paragraphs[0]
        r = p.add_run(hd); r.font.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0xff,0xff,0xff)
        # fondo oxblood
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "8a1e16"); tcPr.append(shd)
    for ri,row in enumerate(rows, start=1):
        for ci,val in enumerate(row):
            c = t.cell(ri,ci); c.text = ""; p = c.paragraphs[0]
            r = p.add_run(str(val)); r.font.size = Pt(9)
            if ri % 2 == 0:
                tcPr = c._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F4F0E7"); tcPr.append(shd)
    if col_widths:
        for ri in range(len(rows)+1):
            for ci,w in enumerate(col_widths):
                t.cell(ri,ci).width = Cm(w)
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(6)

def page_break():
    p = d.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

def maps_row(slug):
    """4 PNGs en 2x2: rec | young / men | women."""
    metrics = [("rec","Voto recuperable"),("young","Composición jóvenes (18-35)"),
               ("men","Composición hombres"),("women","Composición mujeres")]
    t = d.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(m,lab) in enumerate(metrics):
        path = os.path.join(OUTD,"png",f"m_{slug}_{m}_barrio.png")
        if not os.path.exists(path): continue
        cell = t.cell(i//2, i%2); p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(); run.add_picture(path, width=Inches(2.65))
        cp = cell.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(lab); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GR
    d.add_paragraph().paragraph_format.space_after = Pt(8)

# ── 3 · portada ──────────────────────────────────────────────────────────────
eyebrow("Anexo · campañas de zoom por ciudad")
h_title("Las 4 ciudades grandes, una a una", sz=22)
note("Documento de trabajo · uso interno de campaña · borrador automático.")
body("Acompaña al Plan de Pauta Digital de 2ª vuelta. Una página por ciudad: Bogotá, Barranquilla, Cali y Medellín. Para cada una, las tres capas por barrio (voto recuperable, composición de jóvenes y de hombres) y la estrategia de pauta diferenciada por edad y sexo. La idea: que el equipo creativo y el comprador de medios tengan, en un solo lugar, qué mensaje correr, a quién y dónde dentro de la ciudad.")
body("**Cómo se lee cada ciudad:** arriba van las tres capas lado a lado para comparar de un vistazo. Luego el panorama de la ciudad (techo y tamaño del centro), las zonas (localidades o comunas) ordenadas por voto a recuperar, los barrios más jóvenes (donde corre la movilización GOTV) y los con mayor proporción de hombres (donde la izquierda históricamente rinde más). Cierra con una tabla de jugadas concretas por edad × sexo.")
body("**Recordatorio del marco:** edad y sexo a nivel barrio son **composición del electorado** (qué proporción del censo del puesto tiene esa edad o ese sexo), no voto del candidato por demografía. Es justo lo que el administrador de Meta deja segmentar dentro del radio de cada zona. Glosario completo en el documento principal.")
page_break()

# ── 4 · una página por ciudad ────────────────────────────────────────────────
SUPUESTOS_TONO = {
 "izq_joven": ("Movilización · jóvenes (18-35)",
   "Mensaje emocional, identidad y pertenencia. CTA: 'salí a votar'. Formato Reels 9:16, Stories con cuenta regresiva al 21J, video ≤15s con creadores locales. **TikTok orgánico** es donde está la base — no se paga, se trabaja con creadores."),
 "izq_36_60": ("Persuasión · 36-60",
   "Mensaje de gestión y seguridad económica, contraste de costos. Video corto en Meta + skippable de YouTube. Búsqueda Google para defender el nombre del candidato."),
 "der_60mas": ("Contención · 61+",
   "No se gasta en convencer. Si acaso, video sobrio en Meta y prensa digital con mensaje de garantías institucionales. **No** se promete movilizar este segmento; se contiene la fuga."),
 "centro": ("Persuasión · centro disponible (Fajardo + Claudia)",
   "Mensaje moderado, riesgo de los extremos, garantías. Sobre todo en estratos medios y altos. Video skippable YouTube + Meta con segmentación por intereses."),
 "hombre": ("Hombres",
   "El patrón validado es que los **hombres se inclinan más a la izquierda** que las mujeres. En los barrios con alto % de hombres, la pauta de movilización va con fuerza. Creativos directos, llamada a la acción explícita."),
 "mujer": ("Mujeres",
   "Las mujeres se inclinan más a la derecha y la brecha es mayor entre jóvenes. **No** asumir el voto: correr un creativo aparte para mujeres con mensaje propio (gestión, derechos, futuro), no el mismo de los hombres. Evitar tonos confrontacionales."),
}

JUGADAS = [
 ("18-35", "Hombres", "Movilización fuerte. Creativos con creadores locales, identidad de barrio, CTA 'salí a votar'. Pin en barrios jóvenes + radio 2-4 km."),
 ("18-35", "Mujeres", "Movilización con mensaje diferenciado: gestión, derechos, agenda joven. Creativo aparte del de hombres jóvenes. Pin igual."),
 ("36-60", "Hombres", "Persuasión con foco en economía y costos. Video corto Meta + skippable YouTube."),
 ("36-60", "Mujeres", "Persuasión moderada: garantías, gestión, contraste con extremos. Cuidar el tono."),
 ("61+", "Mixto", "Contención. Mensaje sobrio de continuidad institucional. Bajo gasto."),
]

for slug, name, code, sub, divisor, has_loc in CITIES:
    m = muni[code]
    eyebrow(f"Ciudad · {name}")
    h_title(name, sz=20)
    note(sub)
    # 3 mapas lado a lado
    maps_row(slug)
    # panorama
    h_title("El panorama", sz=12, color=INK)
    body(f"**Voto recuperable (techo de izquierda 2V 2022 − candidato 1V 2026):** {NF(m['recuperar'])} votos. "
         f"**Centro disponible (Fajardo + Claudia transferibles):** {NF(m['centro'])}. "
         f"**Base electoral:** {NF(m['base'])} votos. "
         f"**Techo proyectado de izquierda en 2V:** {NF(m['techo'])} ({(m['techo']/m['base']*100 if m['base'] else 0):.1f}% de la base).")
    # top zonas
    h_title(f"Top {divisor}s por voto a recuperar", sz=12, color=INK)
    zonas = top_zonas(code, n=6)
    tbl([divisor.capitalize(), "Recuperar", "Cepeda 1V", "Base"],
        [[z[0], NF(z[1]), NF(z[2]), NF(z[3])] for z in zonas],
        col_widths=[6.5,3.0,3.0,3.0])
    # barrios por composición
    h_title("Barrios para movilización GOTV (más jóvenes)", sz=12, color=INK)
    yb = top_barrios(slug, "young", n=5)
    if yb:
        tbl(["Barrio","% jóvenes (18-35)","Voto recuperable"],
            [[b["nm"], f"{round(b['young']*100)}%", NF(b.get("rec",0))] for b in yb],
            col_widths=[8.5,3.5,3.5])
    h_title("Barrios donde la izquierda históricamente rinde (más hombres)", sz=12, color=INK)
    mb = top_barrios(slug, "men", n=5)
    if mb:
        tbl(["Barrio","% hombres","Voto recuperable"],
            [[b["nm"], f"{round(b['men']*100)}%", NF(b.get("rec",0))] for b in mb],
            col_widths=[8.5,3.5,3.5])
    h_title("Barrios donde hay que cerrar la brecha (más mujeres)", sz=12, color=INK)
    mw = top_barrios(slug, "women", n=5)
    if mw:
        tbl(["Barrio","% hombres","Voto recuperable"],
            [[b["nm"], f"{round(b['men']*100)}%", NF(b.get("rec",0))] for b in mw],
            col_widths=[8.5,3.5,3.5])
        body("Estos barrios tienen mayor proporción de mujeres. La brecha de género (mujeres → derecha, mayor entre jóvenes) sugiere correr un creativo **propio para mujeres** con mensaje diferenciado, no el mismo de los hombres.", sz=9.5, color=GR)
    # jugadas por edad × sexo
    h_title("Jugadas concretas · edad × sexo", sz=12, color=INK)
    tbl(["Edad","Sexo","Jugada de pauta"],
        JUGADAS, col_widths=[2.2,2.2,11.0])
    note("Pin = centroide de la zona o el barrio (en el Excel principal). Radio típico: 2-4 km en Meta. Frecuencia objetivo 3-4; si pasa de 5, rotar creativo.")
    page_break()

# ── 5 · cierre ───────────────────────────────────────────────────────────────
eyebrow("Cierre")
h_title("Cómo combinar las 4 ciudades", sz=18)
body("Si se reparte el presupuesto urbano entre estas 4, una guía razonable es **proporcional al universo a recuperar + centro disponible**. Bogotá pesa con mucho lo más; Cali y Medellín se reparten el resto en proporción similar; Barranquilla queda más chica pero con mayor densidad de neto por peso (mejor relación de movilización útil).")
totals = [(name, muni[code]["recuperar"], muni[code]["centro"]) for _,name,code,_,_,_ in CITIES]
total_universo = sum(r+c for _,r,c in totals) or 1
tbl(["Ciudad","Recuperar","Centro disponible","Universo total","% sugerido del urbano"],
    [[n, NF(r), NF(c), NF(r+c), f"{(r+c)/total_universo*100:.1f}%"] for n,r,c in totals],
    col_widths=[3.5,3.0,3.5,3.0,3.5])
body("Los porcentajes son una guía, no una receta — el comprador de medios ajusta según el comportamiento real (costo por resultado por ciudad) en las primeras 24-48 horas.")
note("Documento de trabajo · borrador automático · ricardoruiz.co · Anexo del Plan de Pauta Digital de 2ª vuelta.", sz=8)

# ── 6 · guardar + embed Inter (reusa helper del Word principal) ──────────────
out_path = os.path.join(OUTD, "Anexo_Ciudades_Pauta_2V.docx")
d.save(out_path)

# embed Inter para que el cliente la vea idéntica
try:
    from build_word import embed_inter
    embed_inter(out_path)
except Exception as e:
    print("  ⚠ embed_inter falló:", e)

print(f"→ {os.path.basename(out_path)} · {round(os.path.getsize(out_path)/1024,1)} KB")
