#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Plan de Pauta Digital · 2ª Vuelta 2026 → Word (Inter incrustada).
Lee pauta-data.json (lo genera build.py) + mapas PNG del informe Pacto."""
import json, os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
OUTD = os.path.join(ROOT, "Bases de datos", "output_pauta_2v")
PACTO = os.path.join(ROOT, "Bases de datos", "output_pacto_1v_2026")
FDIR = os.path.join(ROOT, "tools", "pacto-1v-2026", "fonts")
D = json.load(open(os.path.join(OUTD, "pauta-data.json"), encoding="utf-8"))
M = D["meta"]; MUN = D["munis"]; COM = D["comunas"]; CITY = D["city_names"]; EDAD = D["edad"]

NF = lambda n: f"{round(n):,}".replace(",", ".")
def NFm(n): return (f"{n/1e6:.2f}".replace(".", ",") + "M") if abs(n) >= 1e6 else NF(n)
OX = RGBColor(0x8a,0x1e,0x16); INK = RGBColor(0x1a,0x15,0x10); GR = RGBColor(0x5a,0x54,0x48)
IZQ = RGBColor(0xb3,0x26,0x1d); DER = RGBColor(0x1f,0x47,0xcc)
FONT = "Inter"

d = Document()
for s in d.styles:
    if s.name == "Normal": s.font.name = FONT; s.font.size = Pt(10.5)
_st = d.styles.element; _dd = _st.find(qn("w:docDefaults"))
if _dd is not None:
    _rprd = _dd.find(qn("w:rPrDefault"))
    if _rprd is None: _rprd = OxmlElement("w:rPrDefault"); _dd.insert(0, _rprd)
    _rpr = _rprd.find(qn("w:rPr"))
    if _rpr is None: _rpr = OxmlElement("w:rPr"); _rprd.append(_rpr)
    _rf = _rpr.find(qn("w:rFonts"))
    if _rf is None: _rf = OxmlElement("w:rFonts"); _rpr.insert(0, _rf)
    for a in ("w:ascii","w:hAnsi","w:cs"): _rf.set(qn(a), FONT)

def shade(c,h): e=OxmlElement("w:shd"); e.set(qn("w:fill"),h); c._tc.get_or_add_tcPr().append(e)
def cap(num,txt):
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.page_break_before=True
    r=p.add_run(f"Sección {num}"); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=GR
    p2=d.add_paragraph(); r=p2.add_run(txt); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=OX; p2.paragraph_format.space_after=Pt(5)
def h(txt,size=12.5,color=INK,before=8,after=4):
    p=d.add_paragraph(); r=p.add_run(txt); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
def body(txt,size=10.5,after=6):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in re.split(r"(\*\*.*?\*\*)",txt):
        r=p.add_run(seg[2:-2] if seg.startswith("**") else seg); r.font.size=Pt(size); r.bold=seg.startswith("**"); r.font.color.rgb=INK
def note(txt): body(txt,size=9,after=6)
def bullet(txt):
    p=d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg in re.split(r"(\*\*.*?\*\*)",txt):
        r=p.add_run(seg[2:-2] if seg.startswith("**") else seg); r.font.size=Pt(10.5); r.bold=seg.startswith("**"); r.font.color.rgb=INK
def tbl(headers,rows,sizes=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"
    for i,hh in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(255,255,255); shade(c,"8A1E16")
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row): cs[i].text=""; r=cs[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(8.5)
    return t
def img(name,w=6.3,cap_txt=None):
    path=os.path.join(PACTO,name)
    if not os.path.exists(path): return
    d.add_picture(path,width=Inches(w)); p=d.paragraphs[-1]
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(3)
    if cap_txt:
        pc=d.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=pc.add_run(cap_txt); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=GR; pc.paragraph_format.space_after=Pt(6)

# ── PORTADA ─────────────────────────────────────────────────────────────────
p=d.add_paragraph(); r=p.add_run("DOCUMENTO DE TRABAJO · USO INTERNO DE CAMPAÑA"); r.font.size=Pt(11); r.bold=True; r.font.color.rgb=GR
p.paragraph_format.space_before=Pt(60)
p=d.add_paragraph(); r=p.add_run("Plan de Pauta Digital"); r.font.size=Pt(30); r.bold=True; r.font.color.rgb=INK
p=d.add_paragraph(); r=p.add_run("Segunda Vuelta Presidencial · 21 de junio de 2026"); r.font.size=Pt(15); r.font.color.rgb=OX
p.paragraph_format.space_after=Pt(18)
body("Dónde poner la pauta, a quién hablarle y con qué mensaje en el tramo final, a partir del escrutinio de 1ª vuelta por mesa, el modelo de 2ª vuelta y el análisis de composición etaria del voto. Acompaña al Excel de targeting (lista completa de municipios con centroides) y al tablero interactivo privado.")
note("Borrador automático para planeación interna. Es un escenario para dimensionar la pauta, no un pronóstico. Las cifras de movilización suponen el techo de la izquierda en 2ª vuelta (Petro 2V 2022). La inferencia ecológica acota, no fija intención individual.")

# ── RESUMEN ─────────────────────────────────────────────────────────────────
cap("0","El escenario en una página")
body(f"Cepeda llega a la 2ª vuelta con **{NFm(M['cepeda_1v'])}** votos de 1ª; Abelardo con **{NFm(M['abelardo_1v'])}**. Consolidada la derecha, su piso en 2ª vuelta es del orden de **{NFm(M['abe_floor'])}** y el techo de Cepeda con todo el centro **{NFm(M['cep_ceiling'])}**: una brecha de **{NF(M['gap'])}** votos.")
body(f"Para darle la vuelta, Cepeda debe sumar **{NF(M['need_over_1v'])}** votos sobre su 1ª vuelta: cerca de **{NF(M['center_contrib'])}** vienen de persuadir al centro (Fajardo + Claudia) y alrededor de **1,9 millones** de movilizar a la izquierda que no salió a votar. Es la jugada de Petro en 2022, más difícil esta vez porque la derecha ya se unificó. La pauta digital de estos días debe servir, sobre todo, a esa movilización.")
img("g_brecha_2v.png", 5.6, "La cuenta de la 2ª vuelta: piso de la derecha, techo de Cepeda y la brecha a cerrar.")

# ── 1 · DOS CAMPAÑAS ────────────────────────────────────────────────────────
cap("1","Dos campañas, no una")
body("A seis días, la pauta no es convencer indecisos despacio. Son dos campañas con audiencia, mensaje, canal y métrica distintos. La grande es movilizar; la fina es persuadir.")
h("① Movilización · GOTV  (≈ "+NF(M['abst_target'])+" votos · "+str(M['abst_n_muni'])+" municipios)", color=IZQ)
bullet("**A quién:** simpatizantes de izquierda que se quedaron en casa en 1ª vuelta.")
bullet("**Dónde:** bolsones de izquierda con alta abstención. Regla de oro: la pauta abierta por geografía se concentra **solo donde la izquierda gana la 2ª vuelta**. Prender abstención donde gana Abelardo le suma a Abelardo; para llegar a simpatizantes en zona adversa se usan listas propias (custom audiences), no segmentación geográfica abierta.")
bullet("**Mensaje:** emocional, identidad, pertenencia, urgencia. \"Salí a votar\", recordatorio del puesto, cierre de campaña.")
bullet("**Canales:** Meta (geo + edad + género), TikTok orgánico para jóvenes, WhatsApp territorial.")
h("② Persuasión · Centro  (≈ "+NF(M['center_contrib'])+" votos transferibles)", color=DER)
bullet("**A quién:** votantes de Fajardo y Claudia, transferibles al centro-izquierda.")
bullet("**Dónde:** centros urbanos (Bogotá, Medellín, Cali, capitales).")
bullet("**Mensaje:** moderado, garantías institucionales, riesgo de los extremos, gestión.")
bullet("**Canales:** Meta + YouTube (video corto) + prensa digital.")
note("Cómo se suman: Centro (~0,7M) SE SUMA a movilización (~1,9M) = lo que Cepeda debe crecer sobre su 1ª vuelta. La \"recuperación\" (techo) y la movilización son el mismo universo visto distinto: no se suman entre sí.")

# ── 2 · DÓNDE MOVILIZAR ─────────────────────────────────────────────────────
cap("2","Dónde movilizar")
body("Ranking de municipios por **voto neto de movilización** (abstención × ventaja de la izquierda en 2ª vuelta, solo donde la izquierda gana). El peso de presupuesto sugerido es proporcional a este neto. Lista completa con centroides en el Excel adjunto.")
mob=sorted([m for m in MUN if m["net"]>0], key=lambda x:-x["net"])[:25]
rows=[[i+1, m["muni"], m["dep"], NF(m["net"]), NF(m["censo"]),
       (f"{m['izq']*100:.0f}%" if m['izq'] is not None else "—"),
       (f"{m['cy']*100:.0f}%" if m.get('cy') is not None else "—"),
       f"{m['mw']*100:.1f}%"] for i,m in enumerate(mob)]
tbl(["#","Municipio","Departamento","Neto","Censo","Izq 2V","Cepeda jóv.","% ppto"], rows)
img("m_abstencion_mun.png", 5.4, "Abstención por municipio: la materia prima de la movilización.")

# ── 3 · DÓNDE PERSUADIR ─────────────────────────────────────────────────────
cap("3","Dónde persuadir al centro")
body("Ranking por **voto de centro disponible** (0,55·Fajardo + 0,65·Claudia, supuestos del modelo). Concentrado en lo urbano: aquí la pauta es persuasión, no movilización.")
per=sorted([m for m in MUN if m["cen"]>0], key=lambda x:-x["cen"])[:18]
rows=[[i+1, m["muni"], m["dep"], NF(m["cen"]), NF(m["censo"]), f"{m['pw']*100:.1f}%"] for i,m in enumerate(per)]
tbl(["#","Municipio","Departamento","Centro disp.","Censo","% ppto"], rows)

# ── 4 · HIPERLOCAL ──────────────────────────────────────────────────────────
cap("4","Saturación hiperlocal (radio en Meta)")
body("En las ciudades grandes, un municipio entero es demasiado amplio. Estas comunas/localidades concentran voto recuperable: usar el centroide + un radio de 2-4 km en Meta para saturarlas. El Excel trae todas las comunas con coordenadas.")
for cod in ["16001","31001","03001","05001"]:
    if cod not in COM: continue
    h(CITY.get(cod, cod), size=11.5, color=OX, before=10, after=3)
    rows=[[r["comuna"], NF(r["recuperar"]), NF(r["censo"]), (f"{r['lat']}, {r['lon']}" if r.get("lat") else "—")] for r in COM[cod][:8]]
    tbl(["Comuna / Localidad","Recuperar","Censo","Lat, Lon (radio Meta)"], rows)

# ── 4b · MAPAS DE CALOR POR BARRIO ──────────────────────────────────────────
from docx.shared import Inches as _In
def city_barrio_row(slug, nm):
    have = [m for m in ("rec","young","men") if os.path.exists(os.path.join(OUTD,"png",f"m_{slug}_{m}_barrio.png"))]
    if not have: return
    h(nm, size=11.5, color=OX, before=12, after=2)
    labs = {"rec":"Voto recuperable","young":"% jóvenes (18-35)","men":"% hombres"}
    t = d.add_table(rows=1, cols=len(have))
    for i,m in enumerate(have):
        cell = t.cell(0,i); p = cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(); run.add_picture(os.path.join(OUTD,"png",f"m_{slug}_{m}_barrio.png"), width=_In(2.05))
        cp = cell.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(labs[m]); r.italic=True; r.font.size=Pt(7.5); r.font.color.rgb=GR

cap("5","Mapas de calor por barrio")
body("Para las ciudades grandes, el detalle por barrio: dónde está el voto recuperable y cómo se compone el electorado por edad y sexo. **Edad y sexo van como composición del censo** (dónde se concentran jóvenes / hombres por puesto), que es justo lo que se segmenta en Meta dentro de un radio — no es voto por demografía a nivel barrio (eso sería inferencia demasiado ruidosa). El tablero interactivo trae estas capas para 14 ciudades; aquí van las principales.")
for slug,nm in [("bogota","Bogotá"),("cali","Cali"),("cartagena","Cartagena"),
                ("barranquilla","Barranquilla"),("buenaventura","Buenaventura"),("medellin","Medellín")]:
    city_barrio_row(slug, nm)

# ── 6 · EDAD Y GÉNERO ───────────────────────────────────────────────────────
cap("6","Edad y género: la palanca del mensaje")
nac=EDAD.get("Nacional",{})
def eline(g):
    gd=nac.get(g,{});
    return f"{g}: Cepeda {gd.get('cepeda',0)*100:.0f}% · Abelardo {gd.get('abelardo',0)*100:.0f}%" if gd else g
body("**Choque de generaciones.** Cepeda arrasa entre los jóvenes y casi desaparece entre los mayores; con Abelardo es al revés, y pasa en todo el país.")
bullet(eline("18-35")); bullet(eline("36-60")); bullet(eline("61+"))
body("Implicación de pauta: la **movilización va a jóvenes** — ahí el voto ya es de Cepeda, solo falta sacarlo. Con los mayores no se gasta en convencer, se contiene. En Meta se segmenta por edad: a 18-35 en zonas de izquierda con abstención, GOTV directo.")
body("**Brecha de género.** El patrón regional (validado por mesa con efectos fijos): los hombres se inclinan más a la izquierda y las mujeres más a la derecha, y la brecha es mayor entre los jóvenes. No cambia mucho el \"dónde\" (el censo hombre/mujer es parejo) pero sí el \"qué\": correr creativos separados por género — a los hombres jóvenes, movilización directa; a las mujeres jóvenes, un mensaje que cierre la brecha en vez de dar por hecho el voto.")

# ── 6 · PLATAFORMAS Y PRESUPUESTO ───────────────────────────────────────────
cap("7","Plataformas y presupuesto")
tbl(["Plataforma","Rol","Targeting","Split sugerido"],
    [["Meta (FB/IG)","Movilización + persuasión","Geo (municipio + pin/radio) · edad · género","45%"],
     ["YouTube / Google","Alcance + persuasión","Depto/ciudad · edad · afinidad","20%"],
     ["TikTok (orgánico)","Jóvenes","Creadores + orgánico (NO pauta pagada)","—"],
     ["WhatsApp / territorial","GOTV","Listas propias por municipio","15%"],
     ["X (Twitter)","Conversación","Orgánico + amplificación","—"],
     ["Reserva / testing","A/B + contingencia","Reasignar a lo que convierta","20%"]])
body("**Reparto entre municipios:** proporcional al universo de cada campaña (neto para movilización, centro para persuasión) — ver columna \"% ppto\" de las tablas y el tablero interactivo, que calcula el reparto y el alcance estimado con el presupuesto y el CPM reales.")
note("CPM de referencia para Meta en Colombia: $2-4 USD por 1.000 impresiones; sube en 2ª vuelta por la competencia. Alcance ≈ impresiones / frecuencia (objetivo 3-4). Son estimados para planear, no garantías.")

# ── 7 · MEDICIÓN ────────────────────────────────────────────────────────────
cap("8","Medición")
body("Instalar el **Meta Pixel + Conversions API (CAPI)** en la landing y marcar cada anuncio con **UTMs** (campaña = mov/centro, contenido = creativo). Mirar tres cosas por municipio: **costo por resultado** (registro de puesto, clic-a-WhatsApp, vista del 75% del video), **frecuencia** (si pasa de 4-5 satura: rotar creativo o ampliar geo) y **CTR**. Reasignar el presupuesto cada 24-48 h a lo que convierte más barato. El KPI final no son likes: son **votantes movilizados**, así que priorizar acciones (registro, compartir, recordatorio) sobre alcance puro.")

# ── 8 · ALERTAS LEGALES ─────────────────────────────────────────────────────
cap("9","Tres alertas antes de poner un peso")
bullet("**Autorización de anuncios políticos en Meta.** Correr ads electorales en Colombia exige autorización de la cuenta (verificación de identidad + disclaimer \"pagado por\"). Tarda días. Si la cuenta no está autorizada, no se podrá lanzar pauta nueva antes del 21. Verificar el estatus hoy.")
bullet("**Topes de campaña (CNE) y Cuentas Claras.** El gasto digital cuenta para el tope y se reporta. El plan debe caber en el tope restante; cruzar con el contador/jurídica.")
bullet("**TikTok no acepta pauta pagada** (solo orgánico + creadores). Veda y propaganda: confirmar las fechas vigentes de propaganda electoral digital con la jurídica antes de programar.")

# ── METODOLOGÍA ─────────────────────────────────────────────────────────────
cap("10","Metodología y fuentes")
note("Escrutinio de 1ª vuelta 2026 por mesa (Registraduría). Modelo de 2ª vuelta con supuestos de trasvase explícitos: Paloma 85% → Abelardo, minoritarios de derecha 78% → Abelardo, Fajardo 55% Cepeda / 30% Abelardo, Claudia 65% / 20%, minoritarios de izquierda 85% → Cepeda. Techo de la izquierda = Petro 2ª vuelta 2022 por municipio. Censo electoral Divipole y georreferenciación de puestos (centroides ponderados por censo). Composición etaria por inferencia ecológica a nivel puesto (IC95 ±3-6 pp típico). Neto de movilización = abstención × (2·share_izquierda_2V − 1), solo donde la izquierda gana la 2ª vuelta. Voto blando de centro = 0,55·Fajardo + 0,65·Claudia. El bloque de movilización es un escenario para dimensionar pauta, no un pronóstico; la inferencia ecológica acota, no fija intención individual.")

DOCX = os.path.join(OUTD, "Plan_Pauta_Digital_2V.docx")
d.save(DOCX)

# ── INCRUSTAR INTER ─────────────────────────────────────────────────────────
def embed_inter(docx_path):
    import zipfile, uuid, shutil
    FONTS=[("embedRegular",os.path.join(FDIR,"Inter-Regular.ttf")),
           ("embedBold",   os.path.join(FDIR,"Inter-Bold.ttf")),
           ("embedItalic", os.path.join(FDIR,"Inter-Italic.ttf"))]
    def obf(data,guid):
        key=bytes.fromhex(guid.strip("{}").replace("-",""))[::-1]
        b=bytearray(data)
        for i in range(32): b[i]^=key[i%16]
        return bytes(b)
    zin=zipfile.ZipFile(docx_path,"r"); items={n:zin.read(n) for n in zin.namelist()}; zin.close()
    embeds=[]
    for idx,(kind,path) in enumerate(FONTS,1):
        guid="{"+str(uuid.uuid4()).upper()+"}"
        items[f"word/fonts/font{idx}.odttf"]=obf(open(path,"rb").read(),guid)
        embeds.append((kind,f"rIdFont{idx}",f"fonts/font{idx}.odttf",guid))
    childs="".join(f'<w:{k} r:id="{rid}" w:fontKey="{g}"/>' for k,rid,t,g in embeds)
    node=('<w:font w:name="Inter"><w:charset w:val="00"/><w:family w:val="swiss"/>'
          '<w:pitch w:val="variable"/>'+childs+'</w:font>')
    items["word/fontTable.xml"]=items["word/fontTable.xml"].decode("utf-8").replace("</w:fonts>",node+"</w:fonts>").encode("utf-8")
    rels=('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
          '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
          +"".join(f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="{t}"/>' for k,rid,t,g in embeds)
          +"</Relationships>")
    items["word/_rels/fontTable.xml.rels"]=rels.encode("utf-8")
    ct=items["[Content_Types].xml"].decode("utf-8")
    if 'Extension="odttf"' not in ct:
        ct=ct.replace("</Types>",'<Default Extension="odttf" ContentType="application/vnd.openxmlformats-officedocument.obfuscatedFont"/></Types>')
    items["[Content_Types].xml"]=ct.encode("utf-8")
    st=items["word/settings.xml"].decode("utf-8")
    if "<w:embedTrueTypeFonts/>" not in st:
        st=st.replace("<w:proofState","<w:embedTrueTypeFonts/><w:proofState",1) if "<w:proofState" in st \
           else re.sub(r"(<w:settings[^>]*>)",r"\1<w:embedTrueTypeFonts/>",st,count=1)
    items["word/settings.xml"]=st.encode("utf-8")
    tmp=docx_path+".tmp"; zo=zipfile.ZipFile(tmp,"w",zipfile.ZIP_DEFLATED)
    for n,data in items.items(): zo.writestr(n,data)
    zo.close(); shutil.move(tmp,docx_path)
    return embeds
emb=embed_inter(DOCX)
print(f"→ Plan_Pauta_Digital_2V.docx ({round(os.path.getsize(DOCX)/1024,1)} KB · Inter incrustada, {len(emb)} variantes)")
