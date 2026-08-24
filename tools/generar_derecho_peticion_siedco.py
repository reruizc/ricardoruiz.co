from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('Bases de datos/Seguridad/derecho-peticion/Derecho_de_peticion_SIEDCO_datos_2026.docx')

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tc_pr.append(shd)

def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar = OxmlElement('w:tcMar'); tcPr.append(mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = mar.find(qn(f'w:{side}'))
        if node is None: node = OxmlElement(f'w:{side}'); mar.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'), 'dxa')

def set_fixed_table(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
            tcW.set(qn('w:w'), str(round(widths[i]*1440))); tcW.set(qn('w:type'), 'dxa')
            set_cell_margin(cell)

def text(p, value, bold=False, size=11, color=None, italic=False):
    r = p.add_run(value); r.bold = bold; r.italic = italic; r.font.name = 'Calibri'; r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    return r

def para(doc, value='', bold=False, size=11, after=6, before=0, align=None, italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.10
    if align is not None: p.alignment=align
    text(p,value,bold,size,italic=italic); return p

def heading(doc, value, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); p.paragraph_format.keep_with_next=True; text(p,value,True,16 if level==1 else 13,(46,116,181)); return p

def bullet(doc, value):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.10; text(p,value); return p

doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for n,size,color,before,after in [('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6),('Heading 3',12,'1F4D78',8,4)]:
    s=styles[n]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; text(header,'DERECHO DE PETICIÓN · SOLICITUD DE INFORMACIÓN PÚBLICA',True,8,(89,89,89))
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; text(footer,'Solicitud de información estadística anonimizada · SIEDCO',False,8,(89,89,89))

p=para(doc,'DERECHO DE PETICIÓN',True,20,after=2)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
para(doc,'Solicitud de información pública estadística anonimizada — SIEDCO',False,11,after=16,align=WD_ALIGN_PARAGRAPH.CENTER,italic=True)
meta=[('Destinatario','Policía Nacional de Colombia — Dirección de Investigación Criminal e INTERPOL (DIJIN) / dependencia responsable del SIEDCO'),('Asunto','Solicitud de datos abiertos o copia en formato reutilizable: 1 de enero de 2026 hasta la fecha de corte más reciente disponible'),('Peticionario','[NOMBRE COMPLETO] · CC [NÚMERO DE DOCUMENTO]'),('Notificaciones','[CORREO ELECTRÓNICO] · [DIRECCIÓN Y MUNICIPIO, SI SE DESEA]'),('Fecha','[CIUDAD], [FECHA DE RADICACIÓN]')]
for a,b in meta:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); text(p,a+': ',True); text(p,b)
para(doc,'Respetados señores:',after=8)
para(doc,'En ejercicio del derecho fundamental de petición, solicito información pública estadística contenida en el Sistema de Información Estadístico, Delincuencial, Contravencional y Operativo de la Policía Nacional (SIEDCO). La solicitud se limita a información agregada o anonimizada; no se solicitan nombres, documentos de identidad, direcciones exactas, números de noticia criminal ni otros datos personales o reservados.',after=8)

heading(doc,'1. Fundamento y alcance')
para(doc,'Esta es una petición de información y de documentos. Se formula al amparo del artículo 23 de la Constitución Política y de la Ley 1755 de 2015, en particular sus artículos 13, 14 y 16. Solicito respuesta de fondo y entrega electrónica en un plazo de diez (10) días hábiles, conforme al artículo 14 para peticiones de documentos e información. Si alguna parte estuviera sometida a reserva, solicito que se indique la norma y la motivación concreta de la reserva, y se entregue el resto de la información de manera parcial o anonimizada.')

heading(doc,'2. Información solicitada')
para(doc,'Solicito una exportación nacional de los registros estadísticos de hechos registrados entre el 1 de enero de 2026 y la fecha de corte más reciente disponible al momento de procesar esta petición. De ser posible, agradeceré un archivo CSV o CSV comprimido, codificado en UTF-8, con diccionario de datos y metadatos de extracción. Cada fila puede corresponder a un hecho o a una celda agregada; en este último caso, debe incluir el campo de conteo.',after=8)
bullet(doc,'Todas las tipologías disponibles en SIEDCO, incluyendo como mínimo: homicidio; homicidio en accidente de tránsito; lesiones personales; lesiones en accidente de tránsito; amenazas; extorsión; secuestro; terrorismo; delitos sexuales; violencia intrafamiliar; delitos informáticos; piratería terrestre; y hurtos a personas, comercio, residencias, automotores, motocicletas, bicicletas y celulares.')
bullet(doc,'Identificador técnico anonimizado del registro o, si no existe o no es entregable, una llave de agregación claramente documentada; campo CANTIDAD o equivalente; fecha del hecho y fecha de registro/denuncia, cuando esta última exista y sea publicable.')
bullet(doc,'Ubicación del hecho: departamento y municipio, con códigos DIVIPOLA/DANE; zona urbana/rural; barrio, vereda, comuna, localidad o equivalente cuando se encuentre disponible. Para preservar datos personales o sensibles, no se solicitan coordenadas ni direcciones exactas.')
bullet(doc,'Contexto del hecho: hora, día de semana, mes, clase de sitio, arma o medio empleado, móvil del agresor y móvil de la víctima, y modalidad o descripción de conducta/artículo penal, cuando aplique.')
bullet(doc,'Características de la víctima únicamente en forma anonimizada y categórica: sexo o género, edad o rango de edad, estado civil, nivel educativo, nacionalidad o país de nacimiento, ocupación/clase de empleo, cuando estos campos existan y su entrega sea procedente.')
bullet(doc,'Para hurtos de bienes, las variables técnicas disponibles y no personales: clase y tipo de bien, marca, modelo, color, línea y demás atributos del objeto o vehículo, cuando apliquen.')
bullet(doc,'Catálogos de valores, diccionario de variables, reglas de codificación, fecha/hora de extracción, cobertura temporal por tipología, definición de cada indicador y nota sobre cambios metodológicos, correcciones, duplicados, datos sin fecha y campos con valores por defecto.')

heading(doc,'3. Precisiones sobre el formato de entrega')
para(doc,'Solicito que los datos se entreguen sin conversiones a PDF ni imágenes, en formato abierto, estructurado y reutilizable (preferiblemente CSV UTF-8 y diccionario CSV/XLSX o PDF). Si el volumen requiere dividir la entrega, puede remitirse por enlaces de descarga o archivos comprimidos. Agradezco conservar los códigos originales junto con sus etiquetas, para facilitar la trazabilidad y el cruce con DIVIPOLA.')
para(doc,'Si no es posible entregar el nivel de desagregación solicitado, solicito indicar cuál es el mayor nivel de detalle disponible y entregar esa versión. En particular, si una variable puede facilitar la identificación de personas, solicito su supresión o generalización antes que negar la base completa.')

heading(doc,'4. Información sobre disponibilidad y calidad')
para(doc,'Adicionalmente, solicito informar: (i) cuál es la fecha máxima de actualización disponible para cada tipología; (ii) si existen rezagos de cargue o procesos de validación pendientes para 2026; (iii) qué campos se suprimieron o agregaron frente a los archivos difundidos para 2025; y (iv) si existen criterios de reserva, anonimización o umbrales de publicación que afecten la entrega.')

heading(doc,'5. Notificación')
para(doc,'Solicito que la respuesta y los archivos sean enviados al correo electrónico indicado arriba. Autorizo que las notificaciones se surtan por ese medio.')
para(doc,'Atentamente,',after=24)
para(doc,'[NOMBRE COMPLETO]',True,after=1)
para(doc,'CC [NÚMERO DE DOCUMENTO]',after=0)

heading(doc,'Anexo técnico — diagnóstico de la base actualmente disponible',level=1)
para(doc,'Este anexo facilita la interpretación de la solicitud; no limita el alcance de los datos requeridos.',italic=True)
para(doc,'La carpeta local revisada contiene 20 archivos de enero de 2026 y aproximadamente 79.700 filas agregadas. Las variables recurrentes son: fecha, mes, departamento, municipio, código DANE, barrio, zona, comuna/zona, clase de sitio, hora, día de semana, arma/medio, móvil de agresor y víctima, cantidad y —en delitos con víctima persona— edad, sexo/género, estado civil, país de nacimiento, profesión, escolaridad y, en algunos casos, descripción de conducta. Los hurtos de vehículos incluyen atributos del bien.',after=8)
para(doc,'Vacíos principales identificados:',bold=True,after=4)
for item in [
    'Cobertura temporal: los archivos examinados corresponden solo a enero de 2026; faltan febrero y los meses posteriores hasta el corte vigente.',
    'Trazabilidad: no hay identificador anonimizado de hecho ni fecha de registro/denuncia; tampoco se aporta diccionario de datos, catálogos ni metadatos de extracción.',
    'Coherencia de esquema: los nombres de campos no son uniformes (p. ej., DEPTO_HECHO/DEPARTAMENTO, FECHA_HECHO/FECHA HECHO, CLASE_SITIO/CLASE SITIO) y algunas tipologías omiten MES, HORA_HECHO, MOVIL_AGRESOR, MOVIL_VICTIMA o variables de víctima.',
    'Dimensión de víctima: en hurtos a comercio, vehículos, bicicletas y celulares faltan de forma estructural edad, género, escolaridad, país y otras características; debe aclararse si no aplican o si no se recolectan.',
    'Dimensión jurídica y operativa: no hay, de forma consistente, modalidad/conducta, artículo penal, estado de validación, fuente de registro, fecha de actualización ni marca de corrección/duplicidad.',
    'Georreferenciación: hay barrio y códigos territoriales, pero no una especificación de calidad/cobertura ni una clave geográfica estable que permita distinguir barrios homónimos o relacionar de forma verificable comuna/localidad/vereda.',
    'Calidad: no se documentan valores faltantes, valores por defecto, rezagos ni cambios de metodología. Esto impide separar un cero real de un dato no reportado.'
]: bullet(doc,item)
para(doc,'Canal sugerido de radicación: Sistema PQR2S de la Policía Nacional. Conserve el número de radicado y adjunte este documento, completando antes los campos entre corchetes.',italic=True,after=0)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
