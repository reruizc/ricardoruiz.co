#!/usr/bin/env python3
"""Brújula Asunción 2026 · construye los JSON de datos territoriales.

Lee (insumos en 'Proyecto BL Paraguay/', fuera de git; la app vive en 'brujula-asuncion/'):
  Resultados-por-mesa-2021.csv          municipales 2021 por mesa (TSJE) · ⚠ 2015.csv es una copia byte a byte
  Resultados-por-mesa-2018.csv          generales 2018 por mesa (TSJE)
  ASUNCIÓN/asuncion_unificado_arquetipos_votos_por_barrio_2026.xlsx   padrón 2026 + arquetipos por puesto y zona
  ASUNCIÓN/asuncion_matriz_comportamiento_politico_arquetipos_2026.xlsx  68 barrios Atlas 2012 → zona + comportamiento
  ~/Downloads/BRUJULA_POLITICA_ASUNCION_2026_PARAGUAYO_REVISADO.docx      preocupaciones por barrio (tablas)

Escribe brujula-asuncion/datos/{zonas,puestos,barrios}.json + datos.js + geo.js
"""
import csv, json, re, sys, unicodedata, collections, pathlib
from docx import Document
import openpyxl, warnings
warnings.filterwarnings('ignore')

APP  = pathlib.Path(__file__).resolve().parents[1]            # brujula-asuncion/ (en el repo)
ROOT = APP.parent / 'Proyecto BL Paraguay'                    # insumos pesados, FUERA de git
OUT  = APP / 'datos'
DOCX = pathlib.Path(sys.argv[1]) if len(sys.argv) > 1 else pathlib.Path.home()/'Downloads'/'BRUJULA_POLITICA_ASUNCION_2026_PARAGUAYO_REVISADO.docx'

def norm(s):
    s = unicodedata.normalize('NFKD', str(s or '')).encode('ascii','ignore').decode().lower()
    return re.sub(r'[^a-z0-9]+', ' ', s).strip()

ZONAS = {1:'La Encarnación',2:'La Catedral',3:'San Roque',4:'La Recoleta',5:'Santísima Trinidad',6:'Zeballos Cué'}

# ── resultados TSJE · N elecciones ─────────────────────────────────────────
# Cada elección: archivo, tipo, y las listas que se muestran con nombre. Lo que no esté en
# `listas` cae a 'otros'. Los nombres de candidato están verificados contra prensa/TSJE
# (ver DATOS-ELECTORALES.md); una lista sin nombre verificado va solo con su sigla.
ELECCIONES = [
  {'k':'mun2010','anio':2010,'tipo':'municipal','titulo':'Intendente 2010','file':'Resultados-por-mesa-2010.csv',
   'listas':[('anr','lista1anr','ANR','Arnaldo Samaniego','#c8102e'),('ppq','lista8ppq','Patria Querida','Miguel Carrizosa','#7a3b8f'),('unace','lista7unace','UNACE','','#2f6b3f')]},
  {'k':'pres2013','anio':2013,'tipo':'presidencial','titulo':'Presidente 2013','file':'Resultados-por-mesa-2013.csv',
   'listas':[('anr','lista1anr','ANR','Horacio Cartes','#c8102e'),('pa','lista4pa','Alianza Paraguay Alegre','Efraín Alegre','#1f5fbf'),('ap','lista3ap','Avanza País','Mario Ferreiro','#d96a1a'),('ppq','lista8ppq','Patria Querida','Miguel Carrizosa','#7a3b8f')]},
  {'k':'mun2015','anio':2015,'tipo':'municipal','titulo':'Intendente 2015','file':'Resultados-por-mesa-2015-ASUNCION-real.csv',
   'listas':[('prf','lista3prf','PRF · alianza','Mario Ferreiro','#d96a1a'),('anr','lista1anr','ANR','Arnaldo Samaniego','#c8102e'),('afi','lista4afi','Lista 4 · AFI','','#7a3b8f')]},
  {'k':'pres2018','anio':2018,'tipo':'presidencial','titulo':'Presidente 2018','file':'Resultados-por-mesa-2018.csv',
   'listas':[('anr','lista1anr','ANR','Mario Abdo Benítez','#c8102e'),('ganar','lista5ganar','GANAR','Efraín Alegre','#1f5fbf'),('pvp','lista23pvp','Lista 23 · PVP','','#d96a1a')]},
  {'k':'mun2021','anio':2021,'tipo':'municipal','titulo':'Intendente 2021','file':'Resultados-por-mesa-2021.csv',
   'listas':[('anr','lista1anr','ANR','Óscar “Nenecho” Rodríguez','#c8102e'),('ja','lista5ja','Juntos por Asunción','Eduardo Nakayama','#1f5fbf'),('apt','lista4apt','Asunción para Todos','Johanna Ortega','#7a3b8f')]},
  {'k':'pres2023','anio':2023,'tipo':'presidencial','titulo':'Presidente 2023','file':'Resultados-por-mesa-2023.csv',
   'listas':[('anr','lista1ANR','ANR','Santiago Peña','#c8102e'),('cn','lista3CN20232','Concertación','Efraín Alegre','#1f5fbf'),('pcn','lista911PCN','Cruzada Nacional','Paraguayo “Payo” Cubas','#d96a1a')]},
]
# Locales que cambiaron de SEDE conservando el código (medido comparando nombres entre años):
# el resultado de otro año ahí es de OTRO colegio. (5,15) es dudoso —“Col. Nacional de las
# Américas” vs “Col. Téc. de las Américas”— y se deja fuera por prudencia. NO incluye las
# renombradas que son el mismo sitio (Corazón de Jesús, Nac. de Niñas→Asunción Escalada, Haití).
SEDE_CAMBIADA = {(2,8):{2010,2013,2015,2018}, (3,9):{2010,2013,2015,2018}, (4,17):{2010,2013,2015},
                 (5,15):{2018}, (3,24):{2023}}   # (z,local) -> años cuyo resultado es de otra sede

def load(f):
    with open(f, encoding='utf-8-sig', errors='replace') as fh:
        return [r for r in csv.DictReader(fh, delimiter=';')
                if r['candidatura'].startswith('1-') and r['distrito'].strip()=='0-ASUNCION']

def agg(rows, listas):
    """Agrega por zona y por (zona,local). Devuelve {(z,l|None): Counter}."""
    out = collections.defaultdict(collections.Counter)
    cols = {c for _,c,_,_,_ in listas}
    for r in rows:
        z = int(r['zona'].split('-')[0]); l = int(r['local'].split('-')[0])
        for key in ((z,None),(z,l)):
            c = out[key]
            for k,col,_,_,_ in listas:
                v = r.get(col); c[k] += int(v) if v not in ('',None) else 0
            c['otros'] += sum(int(v) for k,v in r.items() if k.lower().startswith('lista') and k not in cols and v not in ('',None))
            for k in ('blancos','nulos','total'):
                v = r.get(k); c[k] += int(v) if v not in ('',None) else 0
            c['mesas'] += 1
    return out

def bloque(c, e):
    """Compacto: {t: total, m: mesas, p: {lista: pct}}. Nombres/colores viven en meta.elecciones."""
    v = c['total']
    resto = v - sum(c[k] for k,_,_,_,_ in e['listas'])
    return {'t':v,'m':c['mesas'],'p':{**{k:pct(c[k],v) for k,_,_,_,_ in e['listas']}, 'otros':pct(resto,v)}}

def pct(a,b): return round(100*a/b,1) if b else None

RES = {}
for e in ELECCIONES:
    rows = load(ROOT/e['file']); RES[e['k']] = agg(rows, e['listas'])
    ciudad = collections.Counter()
    for (z,l),c in RES[e['k']].items():
        if l is None: ciudad.update(c)
    top = ', '.join(f"{nom} {pct(ciudad[k],ciudad['total'])}%" for k,_,nom,_,_ in e['listas'][:2])
    print(f"{e['titulo']:>16}: {len(rows):>5} mesas · {ciudad['total']:>7} votos · {top}")

# ── padrón + arquetipos ─────────────────────────────────────────────────────
wb = openpyxl.load_workbook(ROOT/'ASUNCIÓN'/'asuncion_unificado_arquetipos_votos_por_barrio_2026.xlsx', read_only=True)
def sheet(ws):
    rows = list(ws.iter_rows(values_only=True)); h = rows[0]
    return [dict(zip(h,r)) for r in rows[1:] if r[0] is not None]
ARQ = {'Influencer':'A1','Profesional':'A2','Familia':'A3','Joven':'A4','Resiliente':'A5'}
ARQ_LONG = {'Influencer aspiracional urbana':'A1','Profesional conectado de gestión':'A2','Familia viral comunitaria':'A3','Joven viral popular':'A4','Residente resiliente del Bañado':'A5'}
puestos_x = sheet(wb['Puesto_unificado_2026'])
zonas_x   = sheet(wb['Barrio_proyeccion_2026'])

def pct(a,b): return round(100*a/b,1) if b else None
def bloque21(c): 
    v = c['total']
    return {'anr':c['anr'],'ja':c['ja'],'apt':c['apt'],'otros':c['otros'],'blancos':c['blancos'],'nulos':c['nulos'],'total':v,
            'p_anr':pct(c['anr'],v),'p_ja':pct(c['ja'],v),'p_apt':pct(c['apt'],v),'mesas':c['mesas']}
def bloque18(c):
    v = c['total']
    return {'anr':c['anr'],'ganar':c['ganar'],'pvp':c['pvp'],'otros':c['otros'],'blancos':c['blancos'],'nulos':c['nulos'],'total':v,
            'p_anr':pct(c['anr'],v),'p_ganar':pct(c['ganar'],v),'p_pvp':pct(c['pvp'],v),'mesas':c['mesas']}

zonas = []
for zx in zonas_x:
    z = [k for k,v in ZONAS.items() if norm(v)==norm(zx['barrio_operativo'])][0]
    zonas.append({
        'cod': z, 'nombre': ZONAS[z], 'puestos': zx['puestos'], 'mesas': zx['mesas_padron'],
        'electores_2026': zx['electores_2026'], 'p_mujeres': round(100*zx['pct_mujeres'],1),
        'participacion_2021': round(100*zx['participacion_historica'],1),
        'proy_votantes_2026': {'baja':zx['proy_votantes_baja_2026'],'base':zx['proy_votantes_base_2026'],'alta':zx['proy_votantes_alta_2026']},
        'edad': {k: round(100*zx[f'pct_edad_{k}'],1) for k in ['18_24','25_34','35_44','45_59','60_mas']},
        'arquetipos': {ARQ_LONG[k]: round(100*zx[f'pct_votos_{k}'],1) for k in ARQ_LONG},
        'elecciones': {e['k']: bloque(RES[e['k']][(z,None)], e) for e in ELECCIONES if (z,None) in RES[e['k']]},
    })
zonas.sort(key=lambda d:d['cod'])

puestos = []
for p in puestos_x:
    z, l = p['zona_cod'], p['local_cod']
    el = {}
    for e in ELECCIONES:
        if (z,l) in RES[e['k']] and e['anio'] not in SEDE_CAMBIADA.get((z,l), set()):
            el[e['k']] = bloque(RES[e['k']][(z,l)], e)
    puestos.append({
        'zona': z, 'cod': l, 'nombre': p['puesto_votacion'], 'segmento': p['segmento'], 'segmento_nombre': p['segmento_nombre'],
        'mesas': p['mesas'], 'electores_2026': p['electores_2026'], 'participacion_2021': round(100*p['participacion_historica_bruta'],1),
        'proy_votantes_2026': p['proy_votantes_base_2026'],
        'arquetipos': {ARQ[k]: round(100*p[f'pct_votos_{k}'],1) for k in ARQ},
        'sede_cambiada': (z,l) in SEDE_CAMBIADA,
        'elecciones': el,
    })
print('puestos', len(puestos), '· sin 2021:', [p['nombre'] for p in puestos if 'mun2021' not in p['elecciones']],
      '· con algún año excluido por cambio de sede:', [p['nombre'] for p in puestos if p['sede_cambiada']])
cob = collections.Counter(k for p in puestos for k in p['elecciones']); print('cobertura por elección (puestos):', dict(cob))

# ── 68 barrios: matriz xlsx + tablas del docx ──────────────────────────────
wb2 = openpyxl.load_workbook(ROOT/'ASUNCIÓN'/'asuncion_matriz_comportamiento_politico_arquetipos_2026.xlsx', read_only=True)
mat = sheet(wb2['Matriz_68_barrios_2026'])
doc = Document(DOCX)
TEMAS = {}; pre = {}
for t in doc.tables:
    h = [c.text.strip() for c in t.rows[0].cells]
    if h[:2]==['Código','Tema']:
        for r in t.rows[1:]: TEMAS[r.cells[0].text.strip()] = r.cells[1].text.strip()
    if h[:2]==['Barrio','Perfil']:
        for r in t.rows[1:]:
            c = [x.text.strip() for x in r.cells]
            codes = [(m.group(1), int(m.group(2))) for m in re.finditer(r'([A-Z]{2})(\d)', c[2])]
            pre[norm(c[0])] = {'perfil':c[1], 'preocupaciones':codes, 'arquetipo_preocupacion':c[3], 'evidencia':c[4]}
ARQP = {}
for t in doc.tables:
    h = [c.text.strip() for c in t.rows[0].cells]
    if h[0].startswith('Arquetipo de preocupación'):
        for r in t.rows[1:]:
            c=[x.text.strip() for x in r.cells]; ARQP[c[0]] = {'barrios':c[1],'emocion':c[2]}
ALIAS = {'ita enramada':'ita enramada','ita pyta punta':'ita pyta punta','ytay':'ytay itay','murucuya':'mburucuya murucuya','loma pyta':'loma pyta',
         'zeballos cue':'zeballos cue','banado cara cara':'banado cara cara','obrero intendente guggiari':'barrio obrero intendente guggiari','jukyty':'jukyty'}
barrios = []; faltan = []
for m in mat:
    n = norm(m['Barrio Atlas 2012']); key = ALIAS.get(n, n)
    d = pre.get(key) or next((v for k,v in pre.items() if k.startswith(key) or key.startswith(k)), None)
    if not d: faltan.append(m['Barrio Atlas 2012']); d = {'perfil':None,'preocupaciones':[],'arquetipo_preocupacion':None,'evidencia':None}
    zc = [k for k,v in ZONAS.items() if norm(v)==norm(m['Zona operativa base'])][0]
    barrios.append({
        'nombre': m['Barrio Atlas 2012'], 'zona': zc, 'zona_nombre': ZONAS[zc], 'zona_confianza': m['Confianza asignación zona'],
        'segmento': m['Segmento'], 'segmento_nombre': m['Tipo de segmento'], 'tipo': m['Tipo socio-territorial principal'],
        'sesgos_emocionales': m['Sesgos emocionales predominantes'], 'sesgos_cognitivos': m['Sesgos cognitivos predominantes'],
        'comportamiento': m['Comportamiento político dominante'], 'riesgo': m['Riesgo electoral'], 'prioridad': m['Prioridad territorial'],
        'electores_2026': m['Electores 2026 estimados'], 'votantes_base_2026': m['Votantes escenario base 2026'],
        'arquetipos': {'A1':round(100*m['% Influencer aspiracional urbana'],1),'A2':round(100*m['% Profesional conectado de gestión'],1),
                       'A3':round(100*m['% Familia viral comunitaria'],1),'A4':round(100*m['% Joven viral popular'],1),'A5':round(100*m['% Residente resiliente del Bañado'],1)},
        'mensaje': m['Mensaje principal 2026'],
        'perfil_2026': d['perfil'], 'preocupaciones': d['preocupaciones'], 'arquetipo_preocupacion': d['arquetipo_preocupacion'], 'evidencia': d['evidencia'],
    })
print('barrios', len(barrios), '· sin match docx:', faltan)

# ── cruce barrio → polígono INE (por código BAR_LOC, nunca por texto en el frontend) ──
GEO = APP/'geo'/'ASUNCION-BARRIOS.geojson'
if GEO.exists():
    gj = json.load(open(GEO))
    geo_by_norm = {norm(f['properties'].get('nombre') or f['properties']['BARLO_DESC']): f['properties']['BAR_LOC'] for f in gj['features']}
    # Las abreviaturas del Atlas/docx contra el nombre completo del INE. Cada par se revisó a mano.
    ALIAS_GEO = {'presidente carlos a lopez':'presidente carlos antonio lopez','mariscal jose f estigarribia':'mariscal jose felix estigarribia',
                 'mcal francisco s lopez':'mariscal francisco solano lopez','obrero intendente guggiari':'obrero intendente b guggiari',
                 'gral jose eduvigis diaz':'general jose eduvigis diaz','gral bernardino caballero':'general bernardino caballero',
                 'madame elisa a lynch':'madame elisa alicia linch','zeballos cue':'zeballos kue','murucuya':'mburucuya'}
    sin_geo = []
    for b in barrios:
        n = norm(b['nombre']); k = ALIAS_GEO.get(n, n)
        cod = geo_by_norm.get(k) or next((c for kk,c in geo_by_norm.items() if kk==k), None)
        b['geo_cod'] = cod
        if not cod: sin_geo.append(b['nombre'])
    usados = {b['geo_cod'] for b in barrios if b['geo_cod']}
    print('geo: barrios con polígono', len(usados), 'de 68 · sin polígono:', sin_geo,
          '· polígonos sin barrio:', [f['properties']['nombre'] for f in gj['features'] if f['properties']['BAR_LOC'] not in usados])
    # copia liviana del GeoJSON para la página (solo código + nombre, coordenadas a 5 decimales ≈ 1 m)
    def rnd(c): return [rnd(x) for x in c] if isinstance(c[0], list) else [round(c[0],5), round(c[1],5)]
    slim = {'type':'FeatureCollection','features':[{'type':'Feature','properties':{'cod':f['properties']['BAR_LOC'],'nombre':f['properties']['nombre']},
             'geometry':{'type':f['geometry']['type'],'coordinates':rnd(f['geometry']['coordinates'])}} for f in gj['features']]}
    (APP/'datos'/'barrios.geojson').write_text(json.dumps(slim, ensure_ascii=False, separators=(',',':')), encoding='utf-8')
    (APP/'datos'/'geo.js').write_text('window.BRUJULA_GEO=' + json.dumps(slim, ensure_ascii=False, separators=(',',':')) + ';\n', encoding='utf-8')
    print('geo.js', (APP/'datos'/'geo.js').stat().st_size//1024, 'KB')
else:
    print('geo: sin GeoJSON en geo/, la página no tendrá mapa')

OUT.mkdir(parents=True, exist_ok=True)
meta = {'v':'2026-08-21','fuentes':{'resultados':'TSJE vía datos.gov.py · resultados por mesa, distrito Asunción: municipales 2010·2015·2021 y presidenciales 2013·2018·2023',
        'padron':'Padrón CAPITAL agregado por mesa (sin datos personales) · electores 2026',
        'arquetipos':'Modelo de 5 arquetipos sociodigitales por puesto (inferencial, no encuesta)',
        'preocupaciones':'Rastreo de prensa y reclamos vecinales feb-ago 2026 (ABC, Última Hora, La Nación) · 68 barrios Atlas 2012'},
        'temas': TEMAS, 'arquetipos_preocupacion': ARQP}
meta['elecciones'] = [{'k':e['k'],'anio':e['anio'],'tipo':e['tipo'],'titulo':e['titulo'],'listas':[{'k':k,'n':nom,'s':cand,'c':col} for k,_,nom,cand,col in e['listas']]} for e in ELECCIONES]
json.dump({'meta':meta,'zonas':zonas}, open(OUT/'zonas.json','w'), ensure_ascii=False, separators=(',',':'))
json.dump({'meta':{'v':meta['v']},'puestos':puestos}, open(OUT/'puestos.json','w'), ensure_ascii=False, separators=(',',':'))
json.dump({'meta':{'v':meta['v'],'temas':TEMAS,'arquetipos_preocupacion':ARQP},'barrios':barrios}, open(OUT/'barrios.json','w'), ensure_ascii=False, separators=(',',':'))
for f in ['zonas','puestos','barrios']: print(f, (OUT/f'{f}.json').stat().st_size//1024, 'KB')

# bundle para abrir el HTML sin servidor (file://): un solo script con los 3 JSON
bundle = {'zonas': zonas, 'puestos': puestos, 'barrios': barrios, 'meta': meta}
(OUT/'datos.js').write_text('window.BRUJULA_DATOS=' + json.dumps(bundle, ensure_ascii=False, separators=(',',':')) + ';\n', encoding='utf-8')
print('datos.js', (OUT/'datos.js').stat().st_size//1024, 'KB')
