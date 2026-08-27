"""Genera el banco de preguntas bilingüe desde el documento editorial."""
import json
import re
import sys
from pathlib import Path
from docx import Document

src = Path(sys.argv[1])
out = Path(sys.argv[2])
doc = Document(src)
paras = [p.text.strip() for p in doc.paragraphs]

def parse_version(start, end):
    questions, topic = {}, None
    i = start
    while i < end:
        text = paras[i]
        match_topic = re.match(r"Tema (\d+)\. (.+)", text)
        match_q = re.match(r"(\d+\.\d+)\. (.+)", text)
        if match_topic:
            topic = {"id": int(match_topic.group(1)), "nombre": match_topic.group(2)}
        elif match_q and topic:
            code, question = match_q.groups()
            scale = [paras[i + j].removeprefix(f"{j}. ") for j in range(1, 6)]
            questions[code] = {"tema": topic.copy(), "texto": question, "escala": scale}
            i += 5
        i += 1
    return questions

v1_start = next(i for i, text in enumerate(paras) if text.startswith("VERSIÓN 1."))
v2_start = next(i for i, text in enumerate(paras) if text.startswith("VERSIÓN 2."))
matrix_start = next(i for i, text in enumerate(paras) if text.startswith("MATRIZ MAESTRA"))
v1 = parse_version(v1_start, v2_start)
v2 = parse_version(v2_start, matrix_start)
scores = {}
for row in doc.tables[0].rows[1:]:
    cells = [c.text.strip() for c in row.cells]
    code = cells[0]
    positions = {}
    for candidate, raw in zip(("camilo", "soledad"), cells[3:5]):
        m = re.match(r"(\d)\s*[·(]\s*([EIAMB])", raw)
        if not m:
            raise ValueError(f"Posición inválida: {code} {candidate} {raw}")
        # E (explícita) entra como evidencia alta; I (inferencia sustentada), media.
        evidence = m.group(2)
        confidence = {"E": "A", "I": "M"}.get(evidence, evidence)
        positions[candidate] = {"pos": int(m.group(1)), "conf": confidence, "evidencia": evidence}
    scores[code] = positions

bank = []
for code, informed in v1.items():
    if code not in v2 or code not in scores:
        raise ValueError(f"Falta versión o posiciones para {code}")
    bank.append({"codigo": code, **informed, "popular": {"texto": v2[code]["texto"], "escala": v2[code]["escala"]}, "pos": scores[code]})
out.write_text("window.BRUJULA_BANCO = " + json.dumps(bank, ensure_ascii=False, separators=(",", ":")) + ";\n", encoding="utf-8")
print(f"{len(bank)} preguntas → {out}")
