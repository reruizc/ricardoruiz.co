from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/Users/ricardoruiz/ricardoruiz.co')
SOURCE = ROOT / 'Bases de datos/DNP/informe de actividades/Informe_Ricardo_Ruiz_JUL_2026_FINAL.docx'
JUNE_REFERENCE = ROOT / 'Bases de datos/DNP/informe de actividades/Informe_Ricardo_Ruiz_JUN_2026V2.docx'
OUTPUT = ROOT / 'Bases de datos/DNP/informe de actividades/Informe_Ricardo_Ruiz_AGO_2026_BORRADOR.docx'

def set_paragraph_text(paragraph, text, blue_link=False):
    # `paragraph.clear()` does not remove hyperlink XML. Retain paragraph
    # formatting only, then rebuild the inline content deterministically.
    for child in list(paragraph._p):
        if child.tag.endswith('}pPr'):
            continue
        paragraph._p.remove(child)
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    if blue_link:
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.underline = True

def set_cell(cell, lines, evidence=False):
    paragraphs = cell.paragraphs
    # Clone paragraph formatting from the reference while rebuilding only slots.
    for index, text in enumerate(lines):
        if index < len(paragraphs):
            p = paragraphs[index]
        else:
            p = cell.add_paragraph()
            p._p.get_or_add_pPr().append(deepcopy(paragraphs[0]._p.pPr))
        set_paragraph_text(p, text, blue_link=evidence)
    for p in paragraphs[len(lines):]:
        set_paragraph_text(p, '')

def set_row_no_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    trPr.append(cant_split)

def set_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'true')
    trPr.append(header)

doc = Document(SOURCE)
doc.paragraphs[4].clear()
r = doc.paragraphs[4].add_run('Período: agosto de 2026')
r.bold = True
doc.paragraphs[10].clear()
r = doc.paragraphs[10].add_run('ENTRE EL 01 - 08 - 2026 al 31 - 08 - 2026')
r.bold = True

data = {
1: (
  ['1.1. Realicé el análisis del estado del arte, el análisis estadístico y territorial del sector TIC y la evaluación preliminar de la Política Nacional de Inteligencia Artificial, con corte estadístico al 4 de agosto de 2026. El ejercicio integró el procesamiento de fuentes sectoriales, el análisis de avances del CONPES 4144 y hallazgos preliminares para el seguimiento de la política.'],
  ['Evidencia 1.1.1. - Informe técnico Entrega 3']
),
2: (
  ['2.1. Analicé la información de ideación y las respuestas recopiladas para la definición del Observatorio Nacional de Inteligencia Artificial (ONIAC), identificando brechas de información, necesidades de los actores y oportunidades de articulación para orientar el seguimiento y la toma de decisiones sobre las políticas del sector TIC.'],
  ['Evidencia 2.1.1. - Matriz de ideación y proyección del ONIAC', 'Evidencia 2.1.2. - Consolidado de respuestas para la definición del ONIAC']
),
3: (
  ['3.1. Ajusté y consolidé la propuesta de indicadores de resultado e impacto para las políticas del sector TIC, alineada con los ejes del CONPES 4144, las fuentes verificadas y los sistemas de seguimiento institucionales. La actualización incorporó los resultados del análisis preliminar y la revisión de la información de ejecución disponible.'],
  ['Evidencia 3.1.1. - Informe técnico Entrega 2', 'Evidencia 3.1.2. - Informe técnico Entrega 3']
),
4: (
  ['4.1. Elaboré la versión 2 de la propuesta de Mesas Temáticas y Diálogos Bilaterales de RENADIA, con el diseño operativo, criterios de priorización, productos esperados y ruta de implementación para fortalecer el relacionamiento de la red con sus miembros.'],
  ['Evidencia 4.1.1. - RENADIA Mesas Temáticas y Diálogos Bilaterales']
),
5: (
  ['5.1. Revisé la versión 3 de la solicitud de ajuste de ficha sin trámite presupuestal del proyecto BPIN 202500000015922 (FUTIC - MinTIC). Verifiqué la respuesta a las observaciones, la consistencia entre el documento y sus anexos, los costos, metas, recursos y cronograma, e identifiqué hallazgos pendientes y nuevas observaciones para el proceso de evaluación ex ante.'],
  ['Evidencia 5.1.1. - Revisión v3 de observaciones BPIN 202500000015922']
),
6: (
  ['6.1. Apliqué herramientas de analítica avanzada y analítica de texto al análisis de la Política Nacional de Inteligencia Artificial, como insumo para el procesamiento de fuentes, la identificación de hallazgos y la evaluación preliminar de los avances de implementación de la política.'],
  ['Evidencia 6.1.1. - Informe técnico Entrega 2', 'Evidencia 6.1.2. - Informe técnico Entrega 3']
),
7: (
  ['7.1. Participé en los seguimientos semanales de RENADIA (4, 11 y 20 de agosto), en la revisión final del webinar RENADIA (4 de agosto) y en el espacio de avances del Observatorio (13 de agosto).',
   '7.2. Participé en la revisión del proyecto de inversión y de sus conclusiones (3 y 11 de agosto), así como en los espacios técnicos sobre el trámite de inversión de la Dirección de Desarrollo Digital (12 de agosto) y sobre corresponsabilidad, cuidado y capacitación (14 de agosto).'],
  ['Evidencia 7.1.1. - Calendario de actividades de agosto']
),
8: (
  ['8.1. Elaboré y presenté el informe periódico de actividades correspondiente al período comprendido entre el 1 y el 31 de agosto de 2026, en el marco del Contrato No. 1025-2026.'],
  ['Evidencia 8.1.1. - Presentar informes - Informe de Actividades Agosto']
),
9: (
  ['9.1. Verifiqué de manera periódica la bandeja del Sistema de Gestión Documental ORFEO durante el período, sin que se registrara correspondencia asignada al usuario que requiriera respuesta o gestión documental.'],
  ['Evidencia 9.1.1. - Verificación de bandeja ORFEO']
),
}

table = doc.tables[0]
for row_index, (description, evidence) in data.items():
    set_cell(table.rows[row_index].cells[1], description)
    set_cell(table.rows[row_index].cells[2], evidence, evidence=True)

# June is the authority for the product-delivery table. Clone its heading and
# table just before the signature block, then update only the two delivered
# product records for August.
june = Document(JUNE_REFERENCE)
signature = doc.paragraphs[14]._p
signature.addprevious(deepcopy(june.paragraphs[12]._p))
signature.addprevious(deepcopy(june.tables[1]._tbl))
products = doc.tables[-1]
products._tbl.append(deepcopy(products.rows[1]._tr))
product_heading = doc.paragraphs[14]
product_heading.paragraph_format.keep_with_next = True
set_header_repeat(products.rows[0])
set_row_no_split(products.rows[1])
set_row_no_split(products.rows[2])

# The June product table relies on a Word table style, while the activity
# table carries explicit borders. Apply the latter so both grids render alike
# in Word and LibreOffice.
activity_pr = table._tbl.tblPr
product_pr = products._tbl.tblPr
for tag in ('w:tblBorders', 'w:tblCellMar'):
    prior = product_pr.find(qn(tag))
    if prior is not None:
        product_pr.remove(prior)
    source_part = activity_pr.find(qn(tag))
    if source_part is not None:
        product_pr.append(deepcopy(source_part))

set_cell(products.rows[1].cells[0], [
    'Producto 2.',
    'Elaboración de un (1) documento que presente el balance de la ejecución de los proyectos de inversión del sector TIC durante la vigencia 2026, así como un análisis preliminar de gestión administrativa del sector TIC, articulado a los mecanismos de implementación y evaluación de políticas digitales de conformidad con el objeto contractual y el Plan Nacional de Desarrollo.'
])
set_cell(products.rows[1].cells[1], [
    '1.1. Entregué un (1) documento técnico correspondiente a la Entrega 2, que desarrolla el marco conceptual y normativo para el análisis de datos y la evaluación de impacto de las políticas del sector TIC; incluye revisión bibliográfica, inventario de fuentes verificadas, propuesta de indicadores, análisis de texto de los documentos CONPES y el análisis inicial del Plan de Acción y Seguimiento.'
])
set_cell(products.rows[1].cells[2], [
    '1.1.1. Documento correspondiente a la Entrega 2 del producto.',
    '1.1.2. Anexo A - código de reproducción del análisis de corpus.'
], evidence=True)

set_cell(products.rows[2].cells[0], [
    'Producto 3.',
    'Presentación del estado del arte del producto, los análisis estadísticos, comparativos y de tendencias sobre la información del sector TIC, la identificación de brechas territoriales, la aplicación de modelos de evaluación ex ante e intermedia y un reporte parcial con resultados y visualizaciones preliminares.'
])
set_cell(products.rows[2].cells[1], [
    '2.1. Entregué un (1) documento técnico correspondiente a la Entrega 3, que presenta el estado del arte, el análisis estadístico y territorial del sector TIC y la evaluación preliminar de la Política Nacional de Inteligencia Artificial. Incluye hallazgos sobre el avance de las acciones, brechas territoriales, riesgos de implementación y visualizaciones para el análisis de la política.'
])
set_cell(products.rows[2].cells[2], [
    '2.1.1. Documento correspondiente a la Entrega 3 del producto.',
    '2.1.2. Anexo A - consultas y procesamiento de información del sector TIC.'
], evidence=True)

doc.core_properties.title = 'Informe de actividades - agosto de 2026'
doc.core_properties.subject = 'Contrato No. 1025-2026'
doc.save(OUTPUT)
print(OUTPUT)
