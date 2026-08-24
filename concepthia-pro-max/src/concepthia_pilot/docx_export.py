"""Generate a Word-compatible institutional response document."""
from __future__ import annotations

from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HEADING_PATTERN = re.compile(r"^\d+\.\s+[A-ZÁÉÍÓÚÜÑ ]+$")


def set_run_font(run, *, bold: bool | None = None, size: float = 11) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text.startswith("Asunto:"):
        label, value = text.split(":", 1)
        set_run_font(paragraph.add_run(f"{label}:"), bold=True)
        set_run_font(paragraph.add_run(value))
        return
    set_run_font(paragraph.add_run(text))


def create_docx(content: str) -> bytes:
    """Return an Arial, Spanish-formatted DOCX from the editable draft text."""
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    for raw_line in content.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if HEADING_PATTERN.match(line):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(9)
            set_run_font(paragraph.add_run(line), bold=True, size=11)
            continue
        add_body_paragraph(document, line)

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
